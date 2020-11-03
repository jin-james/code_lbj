import json
import os
import re
import time
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.xmlchemy import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import RGBColor
import lxml.etree as ET


# MML公式转office MathML的XSL文件
MML2OMML = r'C:\Program Files (x86)\Microsoft Office\Office15\MML2OMML.XSL'

def writeword(gutter, secrecy, main_title, sub_title, info_bar, testee, score_bar, note, winding, group, score_area, fmt, size, usage):
	model_path = r'C:\Users\j20687\Desktop\model.docx'
	# t0 = int(round(time.time() * 1000))
	# tmp_path = '/tmp/%d.docx' % t0
	# os.system('cp %s %s' % (model_path, tmp_path))

	document = Document(model_path)
	# 设置整个文档的默认字体
	microsoft_font = u'宋体'  # u 表示后面的字符串以 Unicode 格式进行编码
	black_font = u'黑体'
	number_font = 'Times New Roman'
	area = qn('w:eastAsia')
	document.styles['Normal'].font.name = microsoft_font
	document.styles['Normal'].font.size = Pt(10.5)
	document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
	document.styles['Normal']._element.rPr.rFonts.set(area, number_font)
	# 指定段落样式
	styles = document.styles
	s = styles.add_style('paragraph', WD_STYLE_TYPE.PARAGRAPH)
	s.font.name = microsoft_font
	s.font.size = Pt(10.5)
	s.font.color.rgb = RGBColor(0, 0, 0)
	s.paragraph_format.line_spacing = Pt(0)  # 行距值
	s.paragraph_format.space_before = Pt(0)  # 段前距
	s.paragraph_format.space_after = Pt(0)  # 段后距
	s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE  # 行间距,单倍行距
	s._element.rPr.rFonts.set(area, number_font)  # 除中文外其它文字 使用的字体 ，备选项
	# 设置页眉
	# header = document.sections[0].header
	# h_para = header.paragraphs[0]
	# h_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
	# run = h_para.add_run("本卷由系统自动生成，请仔细校对后使用")
	# run.font.size = Pt(9)

	types = ['一、', '二、', '三、', '四、', '五、', '六、', '七、', '八、', '九、', '十、']
	if gutter == 1:
		pass
	if secrecy:
		run = document.add_paragraph()
		run.style = s
		p = run.add_run(secrecy)
		p.bold = True
		font = p.font
		font.name = black_font
		font.color.rgb = RGBColor(0, 0, 0)
		font.size = Pt(10.5)
	if main_title:
		run = document.add_paragraph()
		run.style = s
		p = run.add_run(main_title)
		p.bold = True
		font = p.font
		font.name = black_font
		font.color.rgb = RGBColor(0, 0, 0)
		font.size = Pt(15)
		run.alignment = WD_ALIGN_PARAGRAPH.CENTER
	if sub_title:
		run = document.add_paragraph()
		run.style = s
		p1 = run.add_run(sub_title)
		p1.bold = True
		font1 = p1.font
		font1.name = black_font
		font1.color.rgb = RGBColor(0, 0, 0)
		font1.size = Pt(18)
		run.alignment = WD_ALIGN_PARAGRAPH.CENTER
	if info_bar:
		p = document.add_paragraph(info_bar)
		p.style = s
		p.alignment = WD_ALIGN_PARAGRAPH.CENTER
	if testee:
		p = document.add_paragraph(testee)
		p.style = s
		p.alignment = WD_ALIGN_PARAGRAPH.CENTER
	if score_bar:
		rows = 2
		cols = len(group) + 2
		table = document.add_table(rows=rows, cols=cols, style='Table Grid')
		width = OxmlElement('w:tblW')
		width.set(qn('w:type'), 'pct')
		width.set(qn('w:w'), '{}'.format(cols * 500))
		table._tblPr.append(width)  # 指定表格宽度
		table.alignment = WD_TABLE_ALIGNMENT.CENTER
		for r in table.rows:
			for c in r.cells:
				for para in c.paragraphs:
					para.alignment = WD_ALIGN_PARAGRAPH.CENTER
		cell1 = table.rows[0].cells
		cell1[0].paragraphs[0].add_run('题号')
		cell1[-1].paragraphs[0].add_run('总分')
		for i in range(1, cols - 1):
			cell1[i].paragraphs[0].add_run(types[i - 1])
		cell2 = table.rows[1].cells
		cell2[0].paragraphs[0].add_run('得分')

	if note:
		p = document.add_paragraph()
		p.style = s
		run = p.add_run('注意事项：')
		run.add_break()
		re_split = re.compile("<.*?>", re.S | re.I)
		line = re.sub(re_split, '**', note)
		lines = line.split('**')
		for ll in lines:
			run = p.add_run(ll)
			run.add_break()
	if winding:
		i = 0
		for corner in winding:
			p = document.add_paragraph()
			p.style = s
			run = p.add_run(corner)
			run.bold = True
			run.font.size = Pt(12)
			p.alignment = WD_ALIGN_PARAGRAPH.CENTER
			if group:
				for type in group:
					if type['exam_winding']:
						name = type['exam_winding']['name'] or ""
						if name == corner:
							ques_type = type['name']  # 每一种题型
							ques_type = types[i] + str(ques_type)
							i += 1
							questiongroup(type, score_area, usage, ques_type, document, s)  # 试题主体
	else:
		i = 0
		for type in group:
			ques_type = type['name']  # 每一种题型
			ques_type = types[i] + str(ques_type)
			i += 1
			questiongroup(type, score_area, usage, ques_type, document, s)  # 试题主体
	# if load_code == 2:
	# 	document.add_page_break()  # 换页
	# 	p = document.add_paragraph()
	# 	p.style = s
	# 	p.add_run('参考答案').bold = True
	# 	p.alignment = WD_ALIGN_PARAGRAPH.CENTER
	# 	answergroup(group, document, s)
	t1 = int(round(time.time() * 1000))
	# docx = '/tmp/%d.%s' % (t1, 'docx')
	docx = r'C:\Users\j20687\Desktop\%d.%s' % (t1, 'docx')
	document.save(docx)
	# if fmt == 1:
	# 	doc_path = '/tmp/%d.%s' % (t1, 'doc')
	# 	os.system('soffice --headless --invisible --convert-to doc %s --outdir /tmp' % docx)
	# 	return doc_path
	# elif fmt == 2:
	# 	return docx
	# else:
	# 	pdf_path = '/tmp/%d.%s' % (t1, 'pdf')
	# 	os.system('soffice --headless --invisible --convert-to pdf %s --outdir /tmp' % docx)
	# 	return pdf_path
	return docx


def questiongroup(type, score_area, usage, ques_type, document, s):
	if score_area == 1:
		table = document.add_table(rows=2, cols=3, style='Table Grid')
		width = OxmlElement('w:tblW')
		width.set(qn('w:type'), 'pct')
		width.set(qn('w:w'), '{}'.format(3 * 500))
		table._tblPr.append(width)
		for r in table.rows:
			for c in r.cells:
				for para in c.paragraphs:
					para.alignment = WD_ALIGN_PARAGRAPH.CENTER
		set_cell_border(
			table.cell(0, 2),
			top={"sz": 0, "val": "single", "color": "#FFFFFF"},
			bottom={"sz": 0, "color": "#FFFFFF", "val": "single"},
			left={"sz": 0, "val": "single", "color": "#000000"},
			right={"sz": 0, "val": "single", "color": "#FFFFFF", },
		)
		set_cell_border(
			table.cell(1, 2),
			top={"sz": 0, "val": "single", "color": "#FFFFFF"},
			bottom={"sz": 0, "color": "#FFFFFF", "val": "single"},
			left={"sz": 0, "val": "single", "color": "#000000"},
			right={"sz": 0, "val": "single", "color": "#FFFFFF", },
		)
		p = table.cell(0, 2).merge(table.cell(1, 2)).paragraphs[0]
		p.add_run().add_break()
		p.add_run('%s' % str(ques_type))
		table.cell(0, 0).paragraphs[0].add_run('评卷人')
		table.cell(0, 1).paragraphs[0].add_run('得分')
	else:
		p = document.add_paragraph()
		p.add_run('%s' % str(ques_type))

	questions = type['exam_questions']  # 一个题型的所有题
	for question in questions:  # 每一道题
		content = question['question']['content']['content'].replace('<p>', '').replace('&nbsp;', ' ') \
			if question['question']['content']['content'] else ""
		option = question['question']['options'] or ""
		opt_line = question['question']['option_line'] or 0
		answer = question['question']['answer'].replace('<p>', '').replace('&nbsp;', ' ') \
			if question['question']['answer'] else ""
		analysis = question['question']['explain'].replace('<p>', '').replace('&nbsp;', ' ') \
			if question['question']['explain'] else ""
		if usage == 1:
			getline(content, document, s)
			if option or opt_line == 0:
				getoption(option, opt_line, document, s)
			p1 = document.add_paragraph()
			p1.style = s
			p1.add_run("【答案】").font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
			getline(answer, document, s)
			p2 = document.add_paragraph()
			p2.style = s
			p2.add_run("【解析】").font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
			getline(analysis, document, s)
		else:
			getline(content, document, s)
			if option or opt_line == 0:
				getoption(option, opt_line, document, s)


def set_cell_border(cell, **kwargs):
	"""
	Set cell`s border
	Usage:
	set_cell_border(
		cell,
		top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
		bottom={"sz": 12, "color": "#00FF00", "val": "single"},
		left={"sz": 24, "val": "dashed", "shadow": "true"},
		right={"sz": 12, "val": "dashed"},
	)
	"""
	tc = cell._tc
	tcPr = tc.get_or_add_tcPr()

	# 添加标签，如果不存在新增
	tcBorders = tcPr.first_child_found_in("w:tcBorders")
	if tcBorders is None:
		tcBorders = OxmlElement('w:tcBorders')
		tcPr.append(tcBorders)

	# 所有所需的标签
	for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
		edge_data = kwargs.get(edge)
		if edge_data:
			tag = 'w:{}'.format(edge)

			# 添加单元格的子标签
			element = tcBorders.find(qn(tag))
			if element is None:
				element = OxmlElement(tag)
				tcBorders.append(element)
			# looks like order of attributes is important
			for key in ["sz", "val", "color", "space", "shadow"]:
				if key in edge_data:
					element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def getline(string, document, s):
	lines = string.split('</p>')
	for i in range(len(lines)):
		line = lines[i]
		re_table = re.compile("<table.*?>.*?</table>", re.S | re.I)
		if re_table.findall(line):
			get_table(document, line, re_table)
		p = document.add_paragraph()
		p.style = s
		sub_getline(line, p)


def sub_getline(line, p):
	re_math = re.compile("<math.*?>.*?</math>", re.S | re.I)
	re_image = re.compile("<img.*?>", re.S | re.I)
	re_src = re.compile("src=(\".*?\")", re.S | re.I)
	image = re_image.findall(line)
	line = re.sub(re_image, '%$', line)
	src = []
	for img in image:
		src.append(re_src.match(img).replace('http://', ''))
	img_split = line.split('%$')
	for i in range(len(img_split)):
		li = img_split[i]
		if li:
			math = re_math.findall(li)
			if math:
				li = re.sub(re_math, '#$', li)
				sub_line = li.split('#$')
				for j in range(len(math)):
					p.add_run(sub_line[j])
					omml = mathml2omml_lbj(math[j])
					root = ET.fromstring(str(omml))
					para = p._p
					para.extend(root)  # 将omml公式按xml节点的形式添加到当前段落的标签下
				if len(sub_line) > 1:
					p.add_run(sub_line[-1])
			else:
				p.add_run(li)
		if i <= len(img_split) - 2:
			p.add_run().add_picture(src[i])


def mathml2omml_lbj(equation):
	xslt_file = MML2OMML
	dom = ET.fromstring(equation)
	xslt = ET.parse(xslt_file)
	transform = ET.XSLT(xslt)
	newdom = transform(dom)
	return newdom


def get_table(document, line, re_table):
	re_tr = re.compile("<tr>.*?</tr>", re.S | re.I)
	re_td = re.compile("<td.*?>.*?</td>", re.S | re.I)
	re_math = re.compile("<math.*?>.*?</math>", re.S | re.I)
	re_image = re.compile("<img.*?>", re.S | re.I)
	re_src = re.compile("src=(\".*?\")", re.S | re.I)
	table = re_table.findall(line)
	for tbl in table:
		rows = 0
		cols = 0
		td_array = []
		trs = re_tr.findall(tbl)
		rows = len(trs)
		for tr in trs:
			tds = re_td.findall(tr)
			cols = len(tds)
			for td in tds:
				td = re.sub('<td.*?>', '', td)
				td = td.replace('</td>', '')
				td_array.append(td)
		i = 0
		tbl_para = document.add_table(rows=rows, cols=cols, style='Table Grid')
		width = OxmlElement('w:tblW')
		width.set(qn('w:type'), 'pct')
		width.set(qn('w:w'), '{}'.format(cols * 500))
		tbl_para._tblPr.append(width)
		tbl_para.alignment = WD_TABLE_ALIGNMENT.CENTER
		for r in tbl_para.rows:
			for c in r.cells:
				for para in c.paragraphs:
					para.alignment = WD_ALIGN_PARAGRAPH.CENTER
					if re_image.match(td_array[i]):
						img = re_image.match(td_array[i])
						src = re_src.match(str(img)).replace('http://', '')
						para.add_run().add_picture(src)
					elif re_math.match(td_array[i]):
						math = re_math.match(td_array[i])
						omml = mathml2omml_lbj(math)
						root = ET.fromstring(str(omml))
						p = para._p
						p.extend(root)  # 将omml公式按xml节点的形式添加到当前段落的标签下
					else:
						para.add_run(td_array[i])
					i += 1


def getoption(option, opt_line, document, s):
	opt_str = []
	for opt in option:
		key = opt['key']
		content = opt['content']
		string = key + '. ' + content
		opt_str.append(string)
	if opt_line == 1:
		table = document.add_table(rows=1, cols=len(opt_str), style='Table Grid')
		get_opt_table(table, opt_str, s)
	elif opt_line == 2:
		table = document.add_table(rows=2, cols=divmod(len(opt_str), 2)[0], style='Table Grid')
		get_opt_table(table, opt_str, s)
	else:  # 答案超过四行
		table = document.add_table(rows=len(opt_str), cols=1, style='Table Grid')
		get_opt_table(table, opt_str, s)


def get_opt_table(table, opt_str, s):
	i = 0
	for r in table.rows:
		for c in r.cells:
			set_cell_border(
				c,
				top={"sz": 0, "val": "single", "color": "#FFFFFF"},
				bottom={"sz": 0, "color": "#FFFFFF", "val": "single"},
				left={"sz": 0, "val": "single", "color": "#FFFFFF"},
				right={"sz": 0, "val": "single", "color": "#FFFFFF", },
			)
			for para in c.paragraphs:
				para.style = s
				# para.alignment = WD_ALIGN_PARAGRAPH.CENTER
				string = opt_str[i].replace('<p>', '').replace('</p>', '')
				sub_getline(string, para)
				i += 1


def get_paper(path, fmt, size, usage):
	'''

	:param path: json串路径
	:param fmt: 文件下载的版本; 1:'word2003', 2:'word2007', 3:'PDF'
	:param size: 版面大小; 1:'A3', 2:'A4', 3:'8开', 4:'16开'
	:param usage: 试卷类型; 1:教师用卷；2:学生用卷
	:return:
	'''
	alias = []
	file = open(path, 'r', encoding='utf-8')
	exam_paper = json.load(file)
	exam_paper = exam_paper['data']
	exam_style = exam_paper['exam_style']
	question_group = exam_paper['question_group']
	exam_winding = exam_paper['exam_winding']
	winding = []
	for wind in exam_winding:
		winding.append(wind['name'])
	for body in exam_style:
		alias.append(exam_style[str(body)]['alias'])

	gutter = 1 if 'gutter' in alias else 0
	secrecy = exam_paper['secrecy'] if exam_paper['secrecy'] else ''
	main_title = exam_paper['main_title'] if exam_paper['main_title'] else ''
	sub_title = exam_paper['sub_title'] if exam_paper['sub_title'] else ''
	info_bar = exam_paper['info_bar'] if exam_paper['info_bar'] else ''
	testee = exam_paper['testee'] if exam_paper['testee'] else ''
	score_bar = 1 if 'score_bar' in alias else 0
	note = exam_paper['note']
	group = question_group
	# group = question_group if 'group' in alias else 0
	score_area = 1 if 'score_area' in alias else 0
	file_path = \
		writeword(gutter, secrecy, main_title, sub_title, info_bar, testee, score_bar, note, winding, group, score_area, fmt, size, usage)
	return file_path


if __name__ == '__main__':
	path = r'C:\Users\j20687\Desktop\校本资源相关\11.json'
	fmt = 2  # 文件下载的版本; 1:'word2003', 2:'word2007', 3:'PDF'
	size = 2  # 版面大小; 1:'A3', 2:'A4', 3:'8开', 4:'16开'
	usage = 2  # 试卷类型; 1:教师用卷；2:学生用卷
	file_path = get_paper(path, fmt, size, usage)
	print(file_path)

