import barcode
from barcode.writer import ImageWriter
import os
import time
import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.xmlchemy import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL


def ichoice_sub_write(table, r, c, n, ques_no, isfillin=False):
	ques_no += 1
	p = table.cell(r, c).paragraphs[0]
	p.style = s
	if isfillin:
		p.add_run().add_break()
		p.add_run(str(ques_no) + '题、________________________________________')
	else:
		p.add_run().add_break()
		p.add_run(str(ques_no) + '题、')
	for i in range(n):
		p.add_run().add_break()
	return ques_no


def set_col_widths(table, widths):
	for row in table.rows:
		for idx, width in enumerate(widths):
			row.cells[idx].width = width


def hide_frame(table, rows, cols, ischoice=True):
	if ischoice:
		for r in table.rows:
			for c in r.cells:
				set_cell_border(
					c,
					top={"sz": 0, "val": "single", "color": "#FFFFFF"},
					left={"sz": 0, "val": "single", "color": "#FFFFFF"},
					bottom={"sz": 0, "color": "#FFFFFF", "val": "single"},
					right={"sz": 0, "val": "single", "color": "#FFFFFF", },
				)
		set_cell_border(
			table.cell(0, 1),
			left={"sz": 5, "val": "single", "color": "#000000"},
			top={"sz": 5, "color": "#000000", "val": "single"},
		)
		set_cell_border(
			table.cell(0, 2),
			top={"sz": 5, "val": "single", "color": "#000000"},
		)
		set_cell_border(
			table.cell(0, 3),
			top={"sz": 5, "color": "#000000", "val": "single"},
			right={"sz": 5, "val": "single", "color": "#000000", },
		)
		set_cell_border(
			table.cell(rows-1, 1),
			left={"sz": 5, "val": "single", "color": "#000000"},
			bottom={"sz": 5, "color": "#000000", "val": "single"},
		)
		set_cell_border(
			table.cell(rows-1, 2),
			bottom={"sz": 5, "val": "single", "color": "#000000"},
		)
		set_cell_border(
			table.cell(rows-1, 3),
			bottom={"sz": 5, "color": "#000000", "val": "single"},
			right={"sz": 5, "val": "single", "color": "#000000", },
		)
		for r in range(1, rows-1):
			set_cell_border(
				table.cell(r, 1),
				left={"sz": 5, "val": "single", "color": "#000000", },
			)
			set_cell_border(
				table.cell(r, 3),
				right={"sz": 5, "val": "single", "color": "#000000", },
			)
	else:
		for r in range(rows):
			set_cell_border(
				table.cell(r, 0),
				top={"sz": 0, "val": "single", "color": "#FFFFFF"},
				bottom={"sz": 0, "color": "#FFFFFF", "val": "single"},
				left={"sz": 0, "val": "single", "color": "#FFFFFF"},
			)
			set_cell_border(
				table.cell(r, cols - 1),
				top={"sz": 0, "val": "single", "color": "#FFFFFF"},
				bottom={"sz": 0, "color": "#FFFFFF", "val": "single"},
				right={"sz": 0, "val": "single", "color": "#FFFFFF", },
			)


def add_options(p, option_no):
	p.add_run().add_picture(r'C:\Users\j20687\Desktop\校本资源相关\mark_pic\sel-a.pdf', width=Pt(13.8))
	p.add_run(' ').add_picture(r'C:\Users\j20687\Desktop\校本资源相关\mark_pic\sel-b.pdf', width=Pt(13.8))
	p.add_run(' ').add_picture(r'C:\Users\j20687\Desktop\校本资源相关\mark_pic\sel-c.pdf', width=Pt(13.8))
	p.add_run(' ').add_picture(r'C:\Users\j20687\Desktop\校本资源相关\mark_pic\sel-d.pdf', width=Pt(13.8))
	if int(option_no) == 5:
		p.add_run(' ').add_picture(r'C:\Users\j20687\Desktop\校本资源相关\mark_pic\sel-e.pdf', width=Pt(13.8))
	if int(option_no) == 6:
		p.add_run(' ').add_picture(r'C:\Users\j20687\Desktop\校本资源相关\mark_pic\sel-e.pdf', width=Pt(13.8))
		p.add_run(' ').add_picture(r'C:\Users\j20687\Desktop\校本资源相关\mark_pic\sel-f.pdf', width=Pt(13.8))


def add_mark(cols, table, rows):
	cell1 = table.cell(0, 0)
	p1 = cell1.paragraphs[0]
	p1.add_run().add_picture(r'C:\Users\j20687\Desktop\校本资源相关\mark_pic\Rectangle-c.png', width=Pt(6))
	p1.alignment = WD_TABLE_ALIGNMENT.LEFT
	cell1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
	cell2 = table.cell(0, cols + 1)
	p2 = cell2.paragraphs[0]
	p2.add_run().add_picture(r'C:\Users\j20687\Desktop\校本资源相关\mark_pic\Rectangle-c.png', width=Pt(6))
	p2.alignment = WD_TABLE_ALIGNMENT.RIGHT
	cell2.vertical_alignment = WD_ALIGN_VERTICAL.TOP
	if rows > 1:
		cell3 = table.cell(rows - 1, 0)
		cell3.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
		p3 = cell3.paragraphs[0]
		p3.add_run().add_picture(r'C:\Users\j20687\Desktop\校本资源相关\mark_pic\Rectangle-c.png', width=Pt(6))
		p3.alignment = WD_TABLE_ALIGNMENT.LEFT
		cell4 = table.cell(rows - 1, cols + 1)
		p4 = cell4.paragraphs[0]
		p4.add_run().add_picture(r'C:\Users\j20687\Desktop\校本资源相关\mark_pic\Rectangle-c.png', width=Pt(6))
		p4.alignment = WD_TABLE_ALIGNMENT.RIGHT
		cell4.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

def set_cell_border(cell, **kwargs):
	"""
	Set cell`s border
	Usage:
	set_cell_border(
		cell,
		top={"sz": 0, "val": "single", "color": "#FF0000", "space": "0"},
		bottom={"sz": 0, "color": "#00FF00", "val": "single"},
		left={"sz": 24, "val": "dashed", "shadow": "true"},
		right={"sz": 0, "val": "dashed"},
	)
	"""
	tc = cell._tc
	tcPr = tc.get_or_add_tcPr()

	# check for tag existnace, if none found, then create one
	tcBorders = tcPr.first_child_found_in("w:tcBorders")
	if tcBorders is None:
		tcBorders = OxmlElement('w:tcBorders')
		tcPr.append(tcBorders)

	# list over all available tags
	for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
		edge_data = kwargs.get(edge)
		if edge_data:
			tag = 'w:{}'.format(edge)

			# check for tag existnace, if none found, then create one
			element = tcBorders.find(qn(tag))
			if element is None:
				element = OxmlElement(tag)
				tcBorders.append(element)
			# looks like order of attributes is important
			for key in ["sz", "val", "color", "space", "shadow"]:
				if key in edge_data:
					element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def write_option(ques_no, choice_len, document):
	if choice_len:
		if choice_len > 30:
			if divmod(choice_len, 15)[1] != 0:
				rows = divmod(choice_len, 15)[0] + 1
			else:
				rows = divmod(choice_len, 15)[0]
		else:
			rows = 2
		cols = 3
		table_choice = document.add_table(rows=rows, cols=cols + 2, style='Table Grid')
		set_col_widths(table_choice, width5)
		hide_frame(table_choice, rows, cols + 2)
		c = 1
		r = 0
		p = table_choice.cell(r, c).paragraphs[0]
		p.add_run().add_break()
		for n in range(choice_len):
			# option_no = len(choice_group[n].get('options',[]))
			p.style = s
			if ques_no + 1 < 10:
				p.add_run(' ' + str(ques_no + 1) + '、')
			else:
				p.add_run(str(ques_no + 1) + '、')
			# with open(r'C:\Users\j20687\Desktop\校本资源相关\mark_pic\sel-a.tiff', 'rb') as f:
			# p.add_run('[A]').bold = True
			# p.add_run(' [B]').bold = True
			# p.add_run(' [C]').bold = True
			# p.add_run(' [D]').bold = True
			p.add_run().add_picture(r'C:\Users\j20687\Desktop\校本资源相关\mark_pic\sel-a.bmp', width=Pt(13.8))
			p.add_run(' ').add_picture(r'C:\Users\j20687\Desktop\校本资源相关\mark_pic\sel-b.bmp', width=Pt(13.8))
			p.add_run(' ').add_picture(r'C:\Users\j20687\Desktop\校本资源相关\mark_pic\sel-c.bmp', width=Pt(13.8))
			p.add_run(' ').add_picture(r'C:\Users\j20687\Desktop\校本资源相关\mark_pic\sel-d.bmp', width=Pt(13.8))
			# p.add_run(' ').add_picture(r'C:\Users\j20687\Desktop\校本资源相关\mark_pic\sel-e.tiff', width=Pt(13.8))
			# p.add_run(' ').add_picture(r'C:\Users\j20687\Desktop\校本资源相关\mark_pic\sel-f.tiff', width=Pt(13.8))

			# add_options(p, option_no)
			ques_no += 1
			if divmod(ques_no, 15)[1] == 0 and (r+1)*15 < choice_len:
				r += 1
				c = 1
				p.add_run().add_break()  # 另起一行的时候在前一单元格末加一空行
				p = table_choice.cell(r, c).paragraphs[0]
				p.add_run().add_break()  # 另起一行的时候在前加一空行
			elif divmod(ques_no, 5)[1] == 0:
				if ques_no != 15:
					c += 1
					p.add_run().add_break()  # 另起一列的时候在前一列末加一空行
					p = table_choice.cell(r, c).paragraphs[0]
					p.add_run().add_break()  # 另起一列的时候在前加一空行
			else:
				p.add_run().add_break()
		add_mark(cols, table_choice, rows)
	return ques_no


def sub_subjective(n, ques_no, page_num, document):
	width3 = (Cm(0.59), Cm(16.02), Cm(0.59))
	rows = 2
	cols = 1
	div1, div2 = divmod(n, 2)
	if div1 != 0 and div2 == 0:
		document.add_page_break()
		page_num += 1
		add_barcode(document, page_num)
	table_sub = document.add_table(rows=rows, cols=cols + 2, style='Table Grid')
	set_col_widths(table_sub, width3)
	add_mark(cols, table_sub, rows)
	table_sub.cell(1, 1).merge(table_sub.cell(0, 1))
	hide_frame(table_sub, rows, cols + 2, ischoice=False)
	ques_no = ichoice_sub_write(table_sub, 0, 1, 18, ques_no)
	return ques_no, page_num


def sub_fillin(fillin_len, ques_no, document):
	width3 = (Cm(0.59), Cm(16.02), Cm(0.59))
	if fillin_len:
		rows = fillin_len+1 if fillin_len == 1 else fillin_len
		cols = 1
		table_fillin = document.add_table(rows=rows, cols=cols + 2, style='Table Grid')
		set_col_widths(table_fillin, width3)
		hide_frame(table_fillin, rows, cols + 2, ischoice=False)
		for r in range(1 if fillin_len == 1 else rows):
			ques_no = ichoice_sub_write(table_fillin, r, 1, 1, ques_no, isfillin=True)
		add_mark(cols, table_fillin, rows)
		if fillin_len == 1:
			table_fillin.cell(0, 1).merge(table_fillin.cell(1, 1))
	return ques_no


def choice0_60(fillin_len, choice_len, ques_no, document, page_num):
	num = 1
	if 45 >= choice_len >= 31:
		number = 2
		if fillin_len > number:
			ques_no = sub_fillin(number, ques_no, document)
			fillin_len = fillin_len - number
			num += 1
			add_barcode(document, page_num + num)
			ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document)
		else:
			ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document, number=number)
	elif 30 >= choice_len >= 16:
		number = 4
		if fillin_len > number:
			ques_no = sub_fillin(number, ques_no, document)
			fillin_len = fillin_len - number
			num += 1
			document.add_page_break()
			add_barcode(document, page_num + num)
			ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document)
		else:
			ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document,
																				number=number)
	elif 15 >= choice_len >= 1:
		number = 6
		if fillin_len > number:
			ques_no = sub_fillin(number, ques_no, document)
			fillin_len = fillin_len - number
			num += 1
			add_barcode(document, page_num + num)
			ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document)
		else:
			ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document,
																				number=number)
	else:
		ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document, number=fillin_len)
	return ques_no, fill_no, num, page_num, is_fillin_breakpage


def choice60_(fillin_len, choice_len, ques_no, document, page_num):
	num = 1
	if 75 >= choice_len >= 61:
		number = 2
		if fillin_len > number:
			ques_no = sub_fillin(number, ques_no, document)
			fillin_len = fillin_len - number
			document.add_page_break()
			num += 1
			add_barcode(document, page_num + num)
			ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document)
		else:
			ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document,
																				number=number)
	elif 60 >= choice_len >= 46 and fillin_len > 4:
		number = 4
		if fillin_len > number:
			ques_no = sub_fillin(number, ques_no, document)
			fillin_len = fillin_len - number
			document.add_page_break()
			num += 1
			add_barcode(document, page_num + num)
			ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document)
		else:
			ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document,
																				number=number)
	elif 45 >= choice_len >= 31:
		number = 7
		if fillin_len > number:
			ques_no = sub_fillin(number, ques_no, document)
			fillin_len = fillin_len - number
			num += 1
			add_barcode(document, page_num + num)
			ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document)
		else:
			ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document,
																				number=number)
	elif 30 >= choice_len >= 16:
		number = 9
		if fillin_len > number:
			ques_no = sub_fillin(number, ques_no, document)
			fillin_len = fillin_len - number
			document.add_page_break()
			num += 1
			add_barcode(document, page_num + num)
			ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document)
		else:
			ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document,
																				number=number)
	elif 15 >= choice_len >= 1 and fillin_len > 11:
		number = 11
		if fillin_len > number:
			ques_no = sub_fillin(number, ques_no, document)
			fillin_len = fillin_len - number
			num += 1
			add_barcode(document, page_num + num)
			ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document)
		else:
			ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document, number=number)
	else:
		ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document, number=fillin_len)
	return ques_no, fill_no, num, page_num, is_fillin_breakpage


def fillin_in_choice_judgement(fillin_len, ques_no, page_num, document, number=0):
	if not number:
		d0, d1 = divmod(fillin_len, 14)
		for n in range(d0):
			ques_no = sub_fillin(14, ques_no, document)
			page_num += 1
		ques_no = sub_fillin(d1, ques_no, document)
		is_fillin_breakpage = True
		fill_no = d1
	else:
		ques_no = sub_fillin(fillin_len, ques_no, document)
		is_fillin_breakpage = False
		fill_no = 0
	return ques_no, fill_no, page_num, is_fillin_breakpage


def writeline(document, black_font):
	p = document.add_paragraph()
	# p.add_run().add_break()
	run = p.add_run('非选择题（请在各试题的答题区内作答）')
	run.bold = True
	run.font.size = Pt(12)
	run.font.name = black_font


def add_barcode(document, page_num):
	# print(barcode.PROVIDED_BARCODES)
	ITF = barcode.get_barcode_class('itf')
	itf = ITF('{}'.format(page_num), writer=ImageWriter())
	itf_path = r'C:\Users\j20687\Desktop\itf{}_barcode'.format(page_num)
	itf.save(itf_path)
	p = document.add_paragraph()
	p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	run = p.add_run()
	run.add_picture(itf_path+'.png', height=Cm(0.5), width=Cm(4))


if __name__ == '__main__':
	document = docx.Document(r'C:\Users\j20687\Desktop\校本资源相关\答题卡\answersheet.docx')
	# 设置整个文档的默认字体
	microsoft_font = u'宋体'  # u 表示后面的字符串以 Unicode 格式进行编码
	black_font = u'黑体'
	number_font = 'Times New Roman'
	area = qn('w:eastAsia')
	document.styles['Normal'].font.name = microsoft_font
	document.styles['Normal'].font.size = Pt(10.5)
	document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
	document.styles['Normal']._element.rPr.rFonts.set(area, number_font)
	# 指定段落样式
	styles = document.styles
	# fls = time.time()
	# strr = 'lbj_Style%s' %fls  #自定义的样式的名称
	# strr = strr.replace('.', '')
	# strr = strr + ''.join(random.sample('zyxwvutsrqponmlkjihgfedcbaABCDEFGHIJKLMNOPQRST', 5))
	s = styles.add_style('lbj_Style', WD_STYLE_TYPE.PARAGRAPH)
	s.font.name = microsoft_font
	s.font.size = Pt(10.5)
	s.font.color.rgb = RGBColor(0, 0, 0)
	s.paragraph_format.line_spacing = Pt(0)  # 行距值
	s.paragraph_format.space_before = Pt(0)  # 段前距
	s.paragraph_format.space_after = Pt(0)  # 段后距
	s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE  # 行间距
	s._element.rPr.rFonts.set(area, number_font)  # 除中文外其它文字 使用的字体 ，备选项

	add_barcode(document, 1)
	run = document.add_paragraph()
	run.style = s
	p = run.add_run('2019学校12月月考卷试卷副标题答题卡')
	font = p.font
	font.name = black_font
	font.color.rgb = RGBColor(0, 0, 0)
	font.size = Pt(15)
	run.alignment = WD_ALIGN_PARAGRAPH.CENTER
	p = document.add_paragraph()
	p.add_run().add_picture(r'C:\Users\j20687\Desktop\校本资源相关\mark_pic\prefix.png')

	p = document.add_paragraph()
	# p.add_run().add_break()
	run = p.add_run('选择题（请用2B铅笔填涂）')
	run.bold = True
	run.font.size = Pt(12)
	run.font.name = black_font
	width5 = (Cm(0.59), Cm(5.34), Cm(5.34), Cm(5.34), Cm(0.59))
	width3 = (Cm(0.59), Cm(16.02), Cm(0.59))
	# choice_group = sheet.get('choice_question',[])
	sub_mark_num_vec = []  # 每页竖长条表示块数量
	sub_mark_num = 0  # 竖长条表示块总数量
	sub_mark_choice = []  # 选择题黑块区域
	sub_mark_fillin = []  # 填空题黑块区域
	sub_mark_eassy = []  # 简答题黑块区域
	choice_page = 0  # 选择题所在页数
	fillin_page = 0  # 填空题所在页数
	eassy_page = 0  # 简答题所在页数
	page_num = 0  # 页数
	ques_no = 0  # 题号
	choice_answer_cnt = 4  # 选项列数
	choice_answer_pattern = []  # 每列选择题数目
	choice_col_num = 0  # 选择题每行最多列数
	# choice_len = len(choice_group)
	# fillin_group = sheet.get('fillin_question',[])
	# fillin_len = len(fillin_group)
	fillin_len = 10
	choice_len = 30
	fill_no = 0
	is_fillin_breakpage = False
	if 0 <= choice_len <= 5:
		choice_col_num = 1
		choice_answer_pattern.append(choice_len)
	elif 6 <= choice_len <= 10:
		choice_col_num = 2
		choice_answer_pattern.append(5)
		choice_answer_pattern.append(choice_len-5)
	else:
		choice_col_num = 3
		for n in range(divmod(choice_len, 5)[0]):
			choice_answer_pattern.append(5)
		if divmod(choice_len, 5)[1]:
			choice_answer_pattern.append(divmod(choice_len, 5)[1])

	if choice_len > 150:
		ques_no = write_option(ques_no, 60, document)
		document.add_page_break()
		page_num += 1
		add_barcode(document, page_num + 1)
		else_choice = choice_len - 60
		pre_no, rest = divmod(else_choice, 90)
		for n in range(pre_no):
			page_num += 1
			add_barcode(document, page_num + 1)
			ques_no = write_option(ques_no, 90, document)
			document.add_page_break()
		ques_no = write_option(ques_no, rest, document)
		writeline(document, black_font)
		choice_page = page_num + 1 if rest else page_num
		ques_no, fill_no, num, page_num, is_fillin_breakpage = choice60_(fillin_len, rest, ques_no, document, page_num)
		page_num += num
		fillin_page = page_num
	else:
		if choice_len > 60:
			else_choice = choice_len - 60
			ques_no = write_option(ques_no, 60, document)
			document.add_page_break()
			page_num += 1
			add_barcode(document, page_num + 1)
			ques_no = write_option(ques_no, else_choice, document)
			writeline(document, black_font)
			choice_page = page_num + 1 if else_choice else page_num
			ques_no, fill_no, num, page_num, is_fillin_breakpage = choice60_(fillin_len, else_choice, ques_no, document, page_num)
			page_num += num
			fillin_page = page_num
		else:
			ques_no = write_option(ques_no, choice_len, document)
			writeline(document, black_font)
			choice_page = page_num + 1
			ques_no, fill_no, num, page_num, is_fillin_breakpage = choice0_60(fillin_len, choice_len, ques_no, document, page_num)
			page_num += num
			fillin_page = page_num
	# ques_no = sub_fillin(fillin_len, ques_no, document)
	# p = document.add_paragraph()
	# p.add_run().add_break()
	# run = p.add_run('非选择题（请在各试题的答题区内作答）')
	# run.bold = True
	# run.font.size = Pt(12)
	# run.font.name = black_font
	# run.underline = True

	print(ques_no, fill_no, is_fillin_breakpage)
	print(page_num, choice_page, fillin_page)

	# is_fillin_breakpage = False
	# if 0 < choice_len <= 60:
	# 	if 60 >= choice_len >= 46:
	# 		# document.add_page_break()
	# 		ques_no = sub_fillin(fillin_len, ques_no, document)
	# 		fill_no = divmod(fillin_len, 14)[1] if fillin_len > 14 else fillin_len
	# 		is_fillin_breakpage = True
	# 	elif 45 >= choice_len >= 31 and fillin_len > 2:
	# 		ques_no = sub_fillin(2, ques_no, document)
	# 		document.add_page_break()
	# 		ques_no = sub_fillin(fillin_len-2, ques_no, document)
	# 		fill_no = divmod(fillin_len-2, 14)[1] if fillin_len-2 > 14 else fillin_len
	# 		is_fillin_breakpage = True
	# 	elif 30 >= choice_len >= 16 and fillin_len > 4:
	# 		ques_no = sub_fillin(4, ques_no, document)
	# 		document.add_page_break()
	# 		ques_no = sub_fillin(fillin_len - 4, ques_no, document)
	# 		fill_no = divmod(fillin_len-4, 14)[1] if fillin_len-4 > 14 else fillin_len
	# 		is_fillin_breakpage = True
	# 	elif 15 >= choice_len >= 1 and fillin_len > 6:
	# 		ques_no = sub_fillin(6, ques_no, document)
	# 		document.add_page_break()
	# 		ques_no = sub_fillin(fillin_len - 6, ques_no, document)
	# 		fill_no = divmod(fillin_len-6, 14)[1] if fillin_len-6 > 14 else fillin_len
	# 		is_fillin_breakpage = True
	# else:
	# 	choice_len = choice_len - 60

	# if fillin_len:
	# 	rows = fillin_len+1 if fillin_len == 1 else fillin_len
	# 	cols = 1
	# 	table_fillin = document.add_table(rows=rows, cols=cols + 2, style='Table Grid')
	# 	set_col_widths(table_fillin, width3)
	# 	hide_frame(table_fillin, rows, cols + 2, ischoice=False)
	# 	for r in range(1 if fillin_len == 1 else rows):
	# 		ques_no = ichoice_sub_write(table_fillin, r, 1, 1, ques_no, isfillin=True)
	# 	add_mark(cols, table_fillin, rows)
	# 	if fillin_len == 1:
	# 		table_fillin.cell(0, 1).merge(table_fillin.cell(1, 1))
	# # document.add_page_break()

	# subjective_question = sheet.get('subjective_question',[])
	# subjec_len = len(subjective_question)
	subjec_len = 4
	rest_subjec_len = subjec_len
	nn = 0
	if is_fillin_breakpage and fillin_len:
		if fill_no == 0:
			nn = 2
			rest_subjec_len = subjec_len - nn
			for n in range(nn):
				ques_no, page_num = sub_subjective(n, ques_no, page_num, document)
			eassy_page = page_num
			document.add_page_break()
			page_num += 1
			add_barcode(document, page_num)
		elif 7 >= fill_no >= 1:
			nn = 1
			rest_subjec_len = subjec_len - nn
			for n in range(nn):
				ques_no, page_num = sub_subjective(n, ques_no, page_num, document)
			eassy_page = page_num
			document.add_page_break()
			page_num += 1
			add_barcode(document, page_num)
		else:
			nn = 0
			document.add_page_break()
			page_num += 1
			add_barcode(document, page_num)
			eassy_page = page_num
	else:
		document.add_page_break()
		page_num += 1
		add_barcode(document, page_num)
		eassy_page = page_num
	div1, div2 = divmod(rest_subjec_len, 2)
	for n in range(rest_subjec_len):
		ques_no, page_num = sub_subjective(n, ques_no, page_num, document)
	sub_page_num = 1 if div2 else 0
	page_num += sub_page_num
	#  return sub_mark_num_vec except eassy
	if choice_len:
		if fillin_page >= choice_page:
			d1, d2 = divmod(choice_len, 90)
			if d2 > 45:
				for n in range(d1):
					sub_mark_num_vec.append(4)
			else:
				if fill_no >= 0:
					sub_mark_num_vec.append(8)
					if divmod(fill_no, 14)[1] > 7:
						for n in range(divmod(fill_no, 14)[0]):
							sub_mark_num_vec.append(4)
					else:
						if fill_no:
							sub_mark_num_vec.append(4+nn*4)
		else:
			if fillin_len:
				sub_mark_num_vec.append(8)
			else:
				sub_mark_num_vec.append(6)
	else:
		if fillin_len:
			sub_mark_num_vec.append(4+nn*4 if fill_no else nn*4)
		else:
			sub_mark_num_vec.append(2)
	sub_mark_num_vec[0] = sub_mark_num_vec[0]+2
	#  return sub_mark_num_vec with eassy
	for n in range(div1):
		sub_mark_num_vec.append(8)
	if div2 != 0:
		sub_mark_num_vec.append(4)
	#  return sub_mark for choice, fillin, eassy
	block_num = 0
	start = 3
	step = 3
	if choice_len and choice_page:
		for n in range(choice_page):
			block_num += 1
			sub_mark_choice.append(start + step * (block_num-1) + block_num-1)
			sub_mark_choice.append(start + step * block_num + block_num - 1)
	if fillin_page >= choice_page and divmod(choice_len, 90)[1] <= 45:
		for n in range(fillin_page):
			block_num += 1
			sub_mark_fillin.append(start + step * (block_num - 1) + block_num - 1)
			sub_mark_fillin.append(start + step * block_num + block_num - 1)
	else:
		for n in range(fillin_page-choice_page):
			block_num += 1
			sub_mark_fillin.append(start + step * (block_num - 1) + block_num - 1)
			sub_mark_fillin.append(start + step * block_num + block_num - 1)
	for n in range(subjec_len):
		block_num += 1
		sub_mark_eassy.append(start + step * (block_num - 1) + block_num - 1)
		sub_mark_eassy.append(start + step * block_num + block_num - 1)



	#  return sub_mark_num
	for num in sub_mark_num_vec:
		sub_mark_num += num

	print(sub_mark_choice)
	print(sub_mark_fillin)
	print(sub_mark_eassy)
	print(sub_mark_num)
	print(sub_mark_num_vec)
	print(page_num)
	print(choice_answer_cnt, choice_answer_pattern, choice_col_num)





	document.save(r'C:\Users\j20687\Desktop\answersheet.docx')





