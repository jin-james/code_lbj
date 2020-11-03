import uuid

import barcode
import docx
from MyQR import myqr
from barcode.writer import ImageWriter
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.oxml.xmlchemy import OxmlElement
from docx.shared import Pt, RGBColor, Cm


def ichoice_sub_write(table, r, c, n, ques_no, isfillin=False):
    ques_no += 1
    p = table.cell(r, c).paragraphs[0]
    p.style = s
    if isfillin:
        # p.add_run().add_break()
        p.add_run(str(ques_no) + '题、_____________' + ' ').font.size = Pt(9)
    # p.add_run(str(ques_no) + '题、_____________' + ' ')
    # p.add_run(str(ques_no) + '题、_____________' + ' ')
    # p.add_run(str(ques_no) + '题、_____________' + ' ')
    else:
        p.add_run().add_break()
        p.add_run(str(ques_no) + '题、').font.size = Pt(9)
    for i in range(n):
        p.add_run().add_break()
    return ques_no


def set_col_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def hide_frame(table, rows, cols, ischoice=True):
    if ischoice:
        for r in table.rows:
            for c in r.cells:
                set_cell_border(
                    c,
                    top={"sz": 0, "val": "single", "color": "#FFFFFF"},
                    left={"sz": 0, "val": "single", "color": "#FFFFFF"},
                    bottom={"sz": 0, "color": "#FFFFFF", "val": "single"},
                    right={"sz": 0, "val": "single", "color": "#FFFFFF", },
                )
        set_cell_border(
            table.cell(0, 1),
            left={"sz": 5, "val": "single", "color": "#000000"},
            top={"sz": 5, "color": "#000000", "val": "single"},
        )
        set_cell_border(
            table.cell(0, 2),
            top={"sz": 5, "val": "single", "color": "#000000"},
        )
        set_cell_border(
            table.cell(0, 3),
            top={"sz": 5, "color": "#000000", "val": "single"},
            right={"sz": 5, "val": "single", "color": "#000000", },
        )
        set_cell_border(
            table.cell(rows - 1, 1),
            left={"sz": 5, "val": "single", "color": "#000000"},
            bottom={"sz": 5, "color": "#000000", "val": "single"},
        )
        set_cell_border(
            table.cell(rows - 1, 2),
            bottom={"sz": 5, "val": "single", "color": "#000000"},
        )
        set_cell_border(
            table.cell(rows - 1, 3),
            bottom={"sz": 5, "color": "#000000", "val": "single"},
            right={"sz": 5, "val": "single", "color": "#000000", },
        )
        for r in range(1, rows - 1):
            set_cell_border(
                table.cell(r, 1),
                left={"sz": 5, "val": "single", "color": "#000000", },
            )
            set_cell_border(
                table.cell(r, 3),
                right={"sz": 5, "val": "single", "color": "#000000", },
            )
    else:
        for r in range(rows):
            set_cell_border(
                table.cell(r, 0),
                top={"sz": 0, "val": "single", "color": "#FFFFFF"},
                bottom={"sz": 0, "color": "#FFFFFF", "val": "single"},
                left={"sz": 0, "val": "single", "color": "#FFFFFF"},
            )
            set_cell_border(
                table.cell(r, cols - 1),
                top={"sz": 0, "val": "single", "color": "#FFFFFF"},
                bottom={"sz": 0, "color": "#FFFFFF", "val": "single"},
                right={"sz": 0, "val": "single", "color": "#FFFFFF", },
            )


def set_frame(table):
    for r in table.rows:
        for c in r.cells:
            set_cell_border(
                c,
                top={"sz": 5, "val": "single", "color": "#000000"},
                left={"sz": 5, "val": "single", "color": "#000000"},
                bottom={"sz": 5, "color": "#000000", "val": "single"},
                right={"sz": 5, "val": "single", "color": "#000000", },
            )


def hide_frame_single(table):
    for r in table.rows:
        for c in r.cells:
            set_cell_border(
                c,
                top={"sz": 0, "val": "single", "color": "#FFFFFF"},
                left={"sz": 0, "val": "single", "color": "#FFFFFF"},
                bottom={"sz": 0, "color": "#FFFFFF", "val": "single"},
                right={"sz": 0, "val": "single", "color": "#FFFFFF", },
            )


def add_options(p, option_no):
    p.add_run().add_picture(r'C:\Users\j20687\Desktop\校本资源相关\mark_pic\sel-a.pdf', width=Pt(13.8))
    p.add_run(' ').add_picture(r'C:\Users\j20687\Desktop\校本资源相关\mark_pic\sel-b.pdf', width=Pt(13.8))
    p.add_run(' ').add_picture(r'C:\Users\j20687\Desktop\校本资源相关\mark_pic\sel-c.pdf', width=Pt(13.8))
    p.add_run(' ').add_picture(r'C:\Users\j20687\Desktop\校本资源相关\mark_pic\sel-d.pdf', width=Pt(13.8))
    if int(option_no) == 5:
        p.add_run(' ').add_picture(r'C:\Users\j20687\Desktop\校本资源相关\mark_pic\sel-e.pdf', width=Pt(13.8))
    if int(option_no) == 6:
        p.add_run(' ').add_picture(r'C:\Users\j20687\Desktop\校本资源相关\mark_pic\sel-e.pdf', width=Pt(13.8))
        p.add_run(' ').add_picture(r'C:\Users\j20687\Desktop\校本资源相关\mark_pic\sel-f.pdf', width=Pt(13.8))


def add_mark(cols, table, rows=None):
    cell1 = table.cell(0, 0)
    p1 = cell1.paragraphs[0]
    p1.add_run().add_picture(r'D:\校本资源相关\mark_pic\Rectangle-c.png', width=Pt(6))
    p1.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    cell2 = table.cell(0, cols + 1)
    p2 = cell2.paragraphs[0]
    p2.add_run().add_picture(r'D:\校本资源相关\mark_pic\Rectangle-c.png', width=Pt(6))
    p2.alignment = WD_TABLE_ALIGNMENT.RIGHT
    cell2.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    if rows:
        cell3 = table.cell(rows - 1, 0)
        cell3.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        p3 = cell3.paragraphs[0]
        p3.add_run().add_picture(r'D:\校本资源相关\mark_pic\Rectangle-c.png', width=Pt(6))
        p3.alignment = WD_TABLE_ALIGNMENT.LEFT
        cell4 = table.cell(rows - 1, cols + 1)
        p4 = cell4.paragraphs[0]
        p4.add_run().add_picture(r'D:\校本资源相关\mark_pic\Rectangle-c.png', width=Pt(6))
        p4.alignment = WD_TABLE_ALIGNMENT.RIGHT
        cell4.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM


def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 0, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 0, "color": "#00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 0, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def write_option(ques_no, choice_len, document):
    if choice_len:
        if choice_len > 30:
            if divmod(choice_len, 15)[1] != 0:
                rows = divmod(choice_len, 15)[0] + 1
            else:
                rows = divmod(choice_len, 15)[0]
        else:
            rows = 2
        cols = 3
        table_choice = document.add_table(rows=rows, cols=cols + 2)
        table_choice.autofit = False
        set_col_widths(table_choice, width5)
        hide_frame(table_choice, rows, cols + 2)
        c = 1
        r = 0
        p = table_choice.cell(r, c).paragraphs[0]
        p.add_run().add_break()
        for n in range(choice_len):
            # option_no = len(choice_group[n].get('options',[]))
            p.style = s
            if ques_no + 1 >= 100:
                p.add_run(str(ques_no + 1) + " ").font.size = Pt(9)
            elif ques_no + 1 >= 10:
                p.add_run(str(ques_no + 1) + "  ").font.size = Pt(9)
            else:
                p.add_run(str(ques_no + 1) + "   ").font.size = Pt(9)
            # with open(r'C:\Users\j20687\Desktop\校本资源相关\mark_pic\sel-a.tiff', 'rb') as f:
            p.add_run('[A]').font.size = Pt(9)
            p.add_run(' [B]').font.size = Pt(9)
            p.add_run(' [C]').font.size = Pt(9)
            p.add_run(' [D]').font.size = Pt(9)
            # p.add_run(' [E]').font.size = Pt(9)
            # p.add_run(' [F]').font.size = Pt(9)
            # p.add_run(' [G]').font.size = Pt(9)
            # p.add_run().add_picture(r'D:\校本资源相关\mark_pic\bmp小尺寸\sel-a_.bmp', width=Pt(13))
            # p.add_run(' ').add_picture(r'D:\校本资源相关\mark_pic\bmp小尺寸\sel-b_.bmp', width=Pt(13))
            # p.add_run(' ').add_picture(r'D:\校本资源相关\mark_pic\bmp小尺寸\sel-c_.bmp', width=Pt(13))
            # p.add_run(' ').add_picture(r'D:\校本资源相关\mark_pic\bmp小尺寸\sel-d_.bmp', width=Pt(13))
            # p.add_run(' ').add_picture(r'D:\校本资源相关\mark_pic\bmp小尺寸\sel-e_.bmp', width=Pt(13))
            # p.add_run(' ').add_picture(r'D:\校本资源相关\mark_pic\bmp小尺寸\sel-f_.bmp', width=Pt(13))
            # p.add_run(' ').add_picture(r'D:\校本资源相关\mark_pic\bmp小尺寸\sel-g_.bmp', width=Pt(13))

            # add_options(p, option_no)
            ques_no += 1
            if divmod(ques_no, 15)[1] == 0 and (r + 1) * 15 < choice_len:
                r += 1
                c = 1
                p.add_run().add_break()  # 另起一行的时候在前一单元格末加一空行
                p = table_choice.cell(r, c).paragraphs[0]
                p.add_run().add_break()  # 另起一行的时候在前加一空行
            elif divmod(ques_no, 5)[1] == 0:
                if ques_no != 15:
                    c += 1
                    p.add_run().add_break()  # 另起一列的时候在前一列末加一空行
                    p = table_choice.cell(r, c).paragraphs[0]
                    p.add_run().add_break()  # 另起一列的时候在前加一空行
            else:
                p.add_run().add_break()
        add_mark(cols, table_choice)
    return ques_no


def sub_subjective(n, ques_no, page_num, document):
    width3 = (Cm(0.59), Cm(16.02), Cm(0.59))
    rows = 2
    cols = 1
    div1, div2 = divmod(n, 2)
    if div1 != 0 and div2 == 0:
        document.add_page_break()
        document.add_section()
        page_num += 1
    # add_barcode(document, page_num)
    table_sub = document.add_table(rows=rows, cols=cols + 2, style='Table Grid')
    set_col_widths(table_sub, width3)
    add_mark(cols, table_sub)
    table_sub.cell(1, 1).merge(table_sub.cell(0, 1))

    table_sub.rows[0].height = Cm(4.8)

    # cell = table_sub.cell(0, 1)
    # for i in range(40):
    # 	tbl = cell.add_table(rows=1, cols=20)
    # 	set_frame(tbl)
    # 	tbl.rows[0].height = Cm(0.8)
    # for para in cell.paragraphs:
    # 	paragraph_format = para.paragraph_format
    # 	paragraph_format.line_spacing = Pt(5)  # 0.5倍行间距

    hide_frame(table_sub, rows, cols + 2, ischoice=False)
    ques_no = ichoice_sub_write(table_sub, 0, 1, 18, ques_no)
    return ques_no, page_num


def write_eng(page_num, document):
    width3 = (Cm(0.59), Cm(16.02), Cm(0.59))
    rows = 1
    cols = 1
    table_sub = document.add_table(rows=rows, cols=cols + 2, style='Table Grid')
    set_col_widths(table_sub, width3)
    add_mark(cols, table_sub)
    hide_frame(table_sub, rows, cols + 2, ischoice=False)
    cell = table_sub.cell(0, 1)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p = cell.paragraphs[0]
    p.add_run().add_break()
    p.add_run('此处开始作答________________________________________________________________________')
    for n_line in range(10):
        p.add_run().add_break()
        p.add_run('__________________________________________________________________________________')


def write_forbidden(document):
    table = document.add_table(rows=2, cols=1)
    set_cell_border(
        table.cell(0, 0),
        top={"sz": 5, "val": "single", "color": "#000000"},
        left={"sz": 5, "val": "single", "color": "#000000"},
        # bottom={"sz": 5, "color": "#000000", "val": "single"},
        right={"sz": 5, "val": "single", "color": "#000000", },
    )
    set_cell_border(
        table.cell(1, 0),
        # top={"sz": 5, "val": "single", "color": "#000000"},
        left={"sz": 5, "val": "single", "color": "#000000"},
        bottom={"sz": 5, "color": "#000000", "val": "single"},
        right={"sz": 5, "val": "single", "color": "#000000", },
    )
    cell1 = table.cell(0, 0)
    set_cell_background_color(cell1)
    p = cell1.paragraphs[0]
    p.add_run().add_picture(r'D:\校本资源相关\答题卡\forbidden.png')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_break()
    cell2 = table.cell(1, 0)
    set_cell_background_color(cell2)
    p = cell2.paragraphs[0]
    run = p.add_run("请勿在此区域作答或者做任何标记")
    run.font.size = Pt(26)
    run.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def write_chn(ques_no, document):
    width3 = (Cm(0.59), Cm(16.02), Cm(0.59))
    rows = 2
    cols = 1
    table_sub = document.add_table(rows=rows, cols=cols + 2, style='Table Grid')
    set_col_widths(table_sub, width3)
    add_mark(cols, table_sub)
    table_sub.cell(1, 1).merge(table_sub.cell(0, 1))
    table_sub.rows[0].height = Cm(4.8)

    cell = table_sub.cell(0, 1)
    for i in range(40):
        tbl = cell.add_table(rows=1, cols=20)
        set_frame(tbl)
        tbl.rows[0].height = Cm(0.8)
    for para in cell.paragraphs[1:]:
        paragraph_format = para.paragraph_format
        paragraph_format.line_spacing = Pt(5)  # 0.5倍行间距
    hide_frame(table_sub, rows, cols + 2, ischoice=False)


def sub_fillin(fillin_len, ques_no, document):
    width3 = (Cm(0.59), Cm(16.02), Cm(0.59))
    if fillin_len:
        rows = fillin_len + 1 if fillin_len == 1 else fillin_len
        cols = 1
        table_fillin = document.add_table(rows=rows, cols=cols + 2, style='Table Grid')
        set_col_widths(table_fillin, width3)
        hide_frame(table_fillin, rows, cols + 2, ischoice=False)

        for r in range(4):
            ques_no = ichoice_sub_write(table_fillin, r, 1, 0, ques_no, isfillin=True)
        add_mark(cols, table_fillin)
        if fillin_len == 1:
            table_fillin.cell(0, 1).merge(table_fillin.cell(1, 1))
    return ques_no


def choice0_60(fillin_len, choice_len, ques_no, document, page_num):
    num = 1
    if 45 >= choice_len >= 31:
        number = 2
        if fillin_len > number:
            ques_no = sub_fillin(number, ques_no, document)
            fillin_len = fillin_len - number
            num += 1
            # add_barcode(document, page_num + num)
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document)
        else:
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document,
                                                                                    number=number)
    elif 30 >= choice_len >= 16:
        number = 4
        if fillin_len > number:
            ques_no = sub_fillin(number, ques_no, document)
            fillin_len = fillin_len - number
            num += 1
            document.add_page_break()
            # add_barcode(document, page_num + num)
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document)
        else:
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document,
                                                                                    number=number)
    elif 15 >= choice_len >= 1:
        number = 6
        if fillin_len > number:
            ques_no = sub_fillin(number, ques_no, document)
            fillin_len = fillin_len - number
            num += 1
            # add_barcode(document, page_num + num)
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document)
        else:
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document,
                                                                                    number=number)
    else:
        ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document,
                                                                                number=fillin_len)
    return ques_no, fill_no, num, page_num, is_fillin_breakpage


def choice60_(fillin_len, choice_len, ques_no, document, page_num):
    num = 1
    if 75 >= choice_len >= 61:
        number = 2
        if fillin_len > number:
            ques_no = sub_fillin(number, ques_no, document)
            fillin_len = fillin_len - number
            document.add_page_break()
            num += 1
            # add_barcode(document, page_num + num)
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document)
        else:
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document,
                                                                                    number=number)
    elif 60 >= choice_len >= 46 and fillin_len > 4:
        number = 4
        if fillin_len > number:
            ques_no = sub_fillin(number, ques_no, document)
            fillin_len = fillin_len - number
            document.add_page_break()
            num += 1
            # add_barcode(document, page_num + num)
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document)
        else:
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document,
                                                                                    number=number)
    elif 45 >= choice_len >= 31:
        number = 7
        if fillin_len > number:
            ques_no = sub_fillin(number, ques_no, document)
            fillin_len = fillin_len - number
            num += 1
            # add_barcode(document, page_num + num)
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document)
        else:
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document,
                                                                                    number=number)
    elif 30 >= choice_len >= 16:
        number = 9
        if fillin_len > number:
            ques_no = sub_fillin(number, ques_no, document)
            fillin_len = fillin_len - number
            document.add_page_break()
            num += 1
            # add_barcode(document, page_num + num)
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document)
        else:
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document,
                                                                                    number=number)
    elif 15 >= choice_len >= 1 and fillin_len > 11:
        number = 11
        if fillin_len > number:
            ques_no = sub_fillin(number, ques_no, document)
            fillin_len = fillin_len - number
            num += 1
            # add_barcode(document, page_num + num)
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document)
        else:
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document,
                                                                                    number=number)
    else:
        ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document,
                                                                                number=fillin_len)
    return ques_no, fill_no, num, page_num, is_fillin_breakpage


def fillin_in_choice_judgement(fillin_len, ques_no, page_num, document, number=0):
    if not number:
        d0, d1 = divmod(fillin_len, 14)
        for n in range(d0):
            ques_no = sub_fillin(14, ques_no, document)
            page_num += 1
        ques_no = sub_fillin(d1, ques_no, document)
        is_fillin_breakpage = True
        fill_no = d1
    else:
        ques_no = sub_fillin(fillin_len, ques_no, document)
        is_fillin_breakpage = False
        fill_no = 0
    return ques_no, fill_no, page_num, is_fillin_breakpage


def writeline(document, black_font):
    p = document.add_paragraph()
    # p.add_run().add_break()
    run = p.add_run('非选择题（请在各试题的答题区内作答）')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = black_font


def add_barcode(document, page_num, flag=False):
    if flag:
        # print(barcode.PROVIDED_BARCODES)
        ITF = barcode.get_barcode_class('itf')
        itf = ITF('{}'.format(page_num), writer=ImageWriter())
        itf_path = r'C:\Users\j20687\Desktop\itf{}_barcode'.format(page_num)
        itf.save(itf_path)
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        run.add_picture(itf_path + '.png', height=Cm(0.5), width=Cm(4))


def add_prefix_info(header, type_test_no, type_sheet, subject_id, main_title, stu_id_count=8):
    '''
    :param header:
    :param type_test_no: 0:条形码，1：准考证，2：学籍号，3：短考号
    :param main_title: 标题
    :param type_sheet: 0:网阅，1：手阅
    :param stu_id_count: 学号长度，默认为8位
    :param subject_id: 学科识别号
    :return:
    '''
    if type_sheet == 0:
        if type_test_no == 0 or type_test_no == 1:
            table_pre = header.add_table(rows=2, cols=2, width=Cm(22))
            table_pre.alignment = WD_TABLE_ALIGNMENT.CENTER
            # for i in range(2):
            #     for cell in table_pre.column_cells(i):
            #         cell.width = Cm(0.86)
            #     for cell in table_pre.row_cells(i):
            #         cell.width = Cm(0.86)
            prefix_title(table_pre, main_title)
            # p.add_run().add_break()
            table_pre.cell(0, 0).merge(table_pre.cell(0, 1))
            cell10 = table_pre.cell(1, 0)
            p = cell10.paragraphs[0]
            run = p.add_run("姓名：____________  班级：____________")
            run.font.bold = True
            run.font.name = u'黑体'
            p.add_run().add_break()
            run = p.add_run("考场/座位号：_______________")
            run.font.bold = True
            run.font.name = u'黑体'
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run.add_break()
            if type_test_no == 0:
                prefix_notice_1(cell10)
                cell11 = table_pre.cell(1, 1)
                prefix_codebar(cell11)
                prefix_absent(cell11)
                p = cell10.paragraphs[-1]
                p.add_run().add_break()
            else:
                prefix_notice_2(cell10)
                prefix_absent(cell10)
                cell11 = table_pre.cell(1, 1)
                prefix_admission_card(cell11, stu_id_count)
                p = cell11.paragraphs[-1]
                p.add_run().add_break()
        else:
            table_pre = header.add_table(rows=3, cols=2, width=Cm(22))
            table_pre.alignment = WD_TABLE_ALIGNMENT.CENTER
            prefix_title(table_pre, main_title)
            table_pre.cell(0, 0).merge(table_pre.cell(0, 1))
            cell10 = table_pre.cell(1, 0)
            p = cell10.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("姓名：____________  班级：____________  考场/座位号：_______________")
            run.font.bold = True
            run.font.name = u'黑体'
            table_pre.cell(1, 0).merge(table_pre.cell(1, 1))
            cell20 = table_pre.cell(2, 0)
            prefix_notice_3(cell20)
            prefix_absent(cell20)
            cell21 = table_pre.cell(2, 1)
            prefix_admission_card(cell21, stu_id_count)
            p = cell21.paragraphs[-1]
            p.add_run().add_break()

    else:
        path = get_2dcode(subject_id)
        if type_test_no == 0 or type_test_no == 3:
            table_pre = header.add_table(rows=2, cols=2, width=Cm(22))
            table_pre.alignment = WD_TABLE_ALIGNMENT.CENTER
            prefix_title(table_pre, main_title)
            table_pre.cell(0, 0).merge(table_pre.cell(0, 1))
            cell10 = table_pre.cell(1, 0)
            prefix_2dcode(cell10, path)
            cell11 = table_pre.cell(1, 1)
            if type_test_no == 0:
                prefix_codebar(cell11)
            else:
                prefix_short_card(cell11)
        if type_test_no == 1:
            table_pre = header.add_table(rows=2, cols=2, width=Cm(22))
            table_pre.alignment = WD_TABLE_ALIGNMENT.CENTER
            prefix_title(table_pre, main_title)
            cell10 = table_pre.cell(1, 0)
            prefix_2dcode(cell10, path)
            table_pre.cell(0, 1).merge(table_pre.cell(1, 1))
            cell01 = table_pre.cell(0, 1)
            prefix_admission_card(cell01, stu_id_count)
            p = cell01.paragraphs[-1]
            p.add_run().add_break()
        if type_test_no == 2:
            table_pre = header.add_table(rows=3, cols=2, width=Cm(22))
            table_pre.alignment = WD_TABLE_ALIGNMENT.CENTER
            prefix_title(table_pre, main_title)
            table_pre.cell(0, 0).merge(table_pre.cell(0, 1))
            cell10 = table_pre.cell(1, 0)
            p = cell10.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("姓名：____________  班级：____________  考场/座位号：_______________")
            run.font.bold = True
            run.font.name = u'黑体'
            table_pre.cell(1, 0).merge(table_pre.cell(1, 1))
            cell20 = table_pre.cell(2, 0)
            prefix_2dcode_(cell20, path)
            cell21 = table_pre.cell(2, 1)
            prefix_admission_card(cell21, stu_id_count)
            p = cell21.paragraphs[-1]
            p.add_run().add_break()


def get_2dcode(stu_id):
    uid = uuid.uuid4()
    path = '{}.png'.format(uid)
    myqr.run(
        words=stu_id,
        # 扫描二维码后，显示的内容，或是跳转的链接
        version=5,  # 设置容错率
        # level='H',  # 控制纠错水平，范围是L、M、Q、H，从左到右依次升高
        # picture='we.png',  # 图片所在目录，可以是动图
        # colorized=False,  # 黑白(False)还是彩色(True)
        # contrast=1.0,  # 用以调节图片的对比度，1.0 表示原始图片。默认为1.0。
        # brightness=1.0,  # 用来调节图片的亮度，用法同上。
        save_name=path,  # 控制输出文件名，格式可以是 .jpg， .png ，.bmp ，.gif
    )
    return path


def prefix_title(table_pre, main_title):
    hide_frame_single(table_pre)
    cell00 = table_pre.cell(0, 0)
    p = cell00.paragraphs[0]
    run = p.add_run(main_title)
    run.font.size = Pt(14)
    run.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def prefix_notice_1(cell):
    tbl_sub = cell.add_table(rows=1, cols=1)
    tbl_sub.cell(0, 0).width = Pt(350)
    set_frame(tbl_sub)
    set_cell_border(
        tbl_sub.cell(0, 0),
        bottom={"sz": 5, "color": "#000000", "val": "single"},
    )
    p = tbl_sub.cell(0, 0).paragraphs[0]
    paragraph_format = p.paragraph_format
    paragraph_format.line_spacing = Pt(12)  # 12磅行间距
    run = p.add_run("注意事项")
    run.font.bold = True
    p.add_run().add_break()
    run = p.add_run("1．答题前，考生先将自己的姓名、班级、考场填写清楚，并认真核对条形码上")
    run.font.size = Pt(9)
    p.add_run().add_break()
    run = p.add_run("的姓名和准考证号。")
    run.font.size = Pt(9)
    p.add_run().add_break()
    run = p.add_run("2．选择题部分请按题号用2B铅笔填涂方框，修改时用橡皮擦干净，不留痕迹。")
    run.font.size = Pt(9)
    p.add_run().add_break()
    run = p.add_run("3．非选择题部分请按题号用0.5毫米黑色墨水签字笔书写，否则作答无效。要")
    run.font.size = Pt(9)
    p.add_run().add_break()
    run = p.add_run("求字体工整、笔迹清晰。作图时，必须用2B铅笔，并描浓。")
    run.font.size = Pt(9)
    p.add_run().add_break()
    run = p.add_run("4．在草稿纸、试题卷上答题无效。")
    run.font.size = Pt(9)
    p.add_run().add_break()
    run = p.add_run("5．请勿折叠答题卡,保持字体工整、笔迹清晰、卡面清洁。")
    run.font.size = Pt(9)


def prefix_notice_2(cell):
    tbl_sub = cell.add_table(rows=1, cols=1)
    tbl_sub.cell(0, 0).width = Pt(350)
    set_frame(tbl_sub)
    set_cell_border(
        tbl_sub.cell(0, 0),
        bottom={"sz": 5, "color": "#000000", "val": "single"},
    )
    p = tbl_sub.cell(0, 0).paragraphs[0]
    paragraph_format = p.paragraph_format
    paragraph_format.line_spacing = Pt(12)  # 12磅行间距
    run = p.add_run("注意事项")
    run.font.bold = True
    p.add_run().add_break()
    run = p.add_run("1．答题前请将姓名、班级、考场、准考证号填写清楚。")
    run.font.size = Pt(9)
    p.add_run().add_break()
    run = p.add_run("2．客观题答题，必须使用2B铅笔填涂，修改时用橡皮擦干净。")
    run.font.size = Pt(9)
    p.add_run().add_break()
    run = p.add_run("3．主观题答题，必须使用黑色签字笔书写。")
    run.font.size = Pt(9)
    p.add_run().add_break()
    run = p.add_run("4．必须在题号对应的答题区域内作答，超出答题区域书写无效。")
    run.font.size = Pt(9)
    p.add_run().add_break()
    run = p.add_run("5．保持答卷清洁、完整。")
    run.font.size = Pt(9)


def prefix_notice_3(cell):
    tbl_sub = cell.add_table(rows=1, cols=1)
    tbl_sub.cell(0, 0).width = Cm(6)
    set_frame(tbl_sub)
    set_cell_border(
        tbl_sub.cell(0, 0),
        bottom={"sz": 5, "color": "#000000", "val": "single"},
    )
    p = tbl_sub.cell(0, 0).paragraphs[0]
    paragraph_format = p.paragraph_format
    paragraph_format.line_spacing = Pt(12)  # 12磅行间距
    run = p.add_run("注意事项")
    run.font.bold = True
    p.add_run().add_break()
    run = p.add_run("1．答题前请将姓名、班级、考场、准考证号")
    run.font.size = Pt(9)
    run.add_break()
    run = p.add_run("填写清楚。")
    run.font.size = Pt(9)
    run.add_break()
    run = p.add_run("2．客观题答题，必须使用2B铅笔填涂，修改")
    run.font.size = Pt(9)
    run.add_break()
    run = p.add_run("用橡皮擦干净。")
    run.font.size = Pt(9)
    run.add_break()
    run = p.add_run("3.必须在题号对应的答题区域内作答，超出")
    run.font.size = Pt(9)
    run.add_break()
    run = p.add_run("答题区域书写无效。")
    run.font.size = Pt(9)
    run.add_break()


def prefix_codebar(cell):
    tbl1 = cell.add_table(rows=1, cols=1)
    tbl1.cell(0, 0).width = Cm(6.32)
    tbl1.cell(0, 0).height = Cm(2.2)
    set_cell_border(
        tbl1.cell(0, 0),
        top={"sz": 5, "val": "dashed", "color": "#000000"},
        left={"sz": 5, "val": "dashed", "color": "#000000"},
        bottom={"sz": 5, "color": "#000000", "val": "dashed"},
        right={"sz": 5, "val": "dashed", "color": "#000000", },
    )
    p = tbl1.cell(0, 0).paragraphs[0]
    run = p.add_run("贴条形码区")
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(200, 200, 200)
    p.add_run().add_break()
    p.add_run().add_break()
    run = p.add_run("(正面朝上，切勿贴出虚线方框)")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(200, 200, 200)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def prefix_absent(cell):
    tbl2 = cell.add_table(rows=1, cols=1)
    tbl2.cell(0, 0).width = Cm(6.57)
    tbl2.cell(0, 0).height = Cm(2)
    set_frame(tbl2)
    p = tbl2.cell(0, 0).paragraphs[0]
    run = p.add_run("正确填涂    ")
    run.add_picture(r'D:\校本资源相关\mark_pic\bmp小尺寸\black.png', width=Pt(14))
    run.font.bold = True
    run.font.size = Pt(9)
    run = p.add_run("    缺考标记    ")
    run.add_picture(r'D:\校本资源相关\mark_pic\bmp小尺寸\white.png', width=Pt(14))
    run.font.bold = True
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def prefix_admission_card(cell, stu_id_count):
    tbl = cell.add_table(rows=3, cols=stu_id_count)
    for i in range(stu_id_count):
        for cell in tbl.column_cells(i):
            cell.width = Cm(0.86)
    set_frame(tbl)
    p = tbl.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("准考证号")
    for i in range(stu_id_count - 1):
        tbl.cell(0, i).merge(tbl.cell(0, i + 1))
    for j in range(stu_id_count):
        p = tbl.cell(2, j).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for k in range(9):
            run = p.add_run("[{}]".format(k))
            run.font.size = Pt(9)
            run.add_break()
        run = p.add_run("[{}]".format(9))
        run.font.size = Pt(9)


def prefix_2dcode(cell, path):
    tbl2 = cell.add_table(rows=1, cols=2)
    # set_frame(tbl2)
    for cell in tbl2.column_cells(0):
        cell.width = Cm(4)
    for cell in tbl2.column_cells(1):
        cell.width = Cm(4)
    p = tbl2.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Pt(90))
    p = tbl2.cell(0, 1).paragraphs[0]
    # p.add_run().add_break()
    run = p.add_run("姓名：_________")
    run.font.bold = True
    run.add_break()
    run.add_break()
    run = p.add_run("班级：_________")
    run.font.bold = True
    run.add_break()
    run.add_break()
    run = p.add_run("考号：_________")
    run.font.bold = True
    run.add_break()
    run.add_break()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def prefix_2dcode_(cell, path):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    tbl2 = cell.add_table(rows=2, cols=1)
    set_frame(tbl2)
    tbl2.cell(0, 0).width = Cm(6)
    tbl2.cell(1, 0).width = Cm(6)
    tbl2.cell(0, 0).height = Cm(6)
    tbl2.cell(1, 0).height = Cm(2)
    p = tbl2.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Pt(90))
    p = tbl2.cell(1, 0).paragraphs[0]
    run = p.add_run("正确填涂    ")
    run.add_picture(r'D:\校本资源相关\mark_pic\bmp小尺寸\black.png', width=Pt(14))
    run.font.bold = True
    run.font.size = Pt(9)


def prefix_short_card(cell):
    cell.vertical = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("准考证号后三位数字")
    run.font.size = Pt(9)
    tbl = cell.add_table(rows=3, cols=3)
    set_frame(tbl)
    for cell in tbl.column_cells(0):
        cell.width = Cm(1)
    for cell in tbl.column_cells(1):
        cell.width = Cm(1)
    for cell in tbl.column_cells(2):
        cell.width = Cm(6.8)
    for cell in tbl.row_cells(0):
        cell.height = Cm(1)
    for cell in tbl.row_cells(1):
        cell.height = Cm(1)
    p = tbl.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("短")
    run.font.size = Pt(9)
    p = tbl.cell(1, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("考")
    run.font.size = Pt(9)
    p = tbl.cell(2, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("号")
    run.font.size = Pt(9)
    for i in range(2):
        tbl.cell(i, 0).merge(tbl.cell(i + 1, 0))
    for j in range(3):
        cell2 = tbl.cell(j, 2)
        cell2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for k in range(9):
            p = cell2.paragraphs[0]
            run = p.add_run("[{}]".format(k) + " ")
            run.font.size = Pt(9)
        run = p.add_run("[{}]".format(9))
        run.font.size = Pt(9)
    # tbl.cell(0, 2).merge(tbl.cell(1, 2)).merge(tbl.cell(2, 2))


def set_cell_background_color(cell):
    # if not isinstance(rgb_color, RGBValue):
    #     print('rgbColor is not RGBValue...', type(rgb_color))
    #     return
    # hr = str(hex(int(rgb_color.r)))[-2:]
    # hg = str(hex(int(rgb_color.g)))[-2:]
    # hb = str(hex(int(rgb_color.b)))[-2:]
    # color_str = hr + hg + hb
    # shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value=color_str))
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="CCCAC9"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading_elm_1)


def add_postfix_info(document, page_count):
    page_binary = get_binary_code(page_count)
    for i in range(page_count):
        section = document.add_section()
        footer = section.footer
        p = footer.paragraphs[0]
        p.add_run('{}'.format(page_binary[i]))


def get_binary_code(page_count):
    sum = []
    for i in range(1, page_count+1):
        binary = []
        m, n = divmod(i, 2)
        recurrent(m, n, binary)
        sum.append(binary)
    blank_count = max(max([len(ls) for ls in sum]), 3)
    for ls in sum:
        gap = blank_count - len(ls)
        for i in range(gap):
            ls.insert(0, 0)
    return sum


def recurrent(m, n, binary):
    if m == 0:
        binary.insert(0, n)
    else:
        binary.insert(0, n)
        m2, n2 = divmod(m, 2)
        recurrent(m2, n2, binary)


if __name__ == '__main__':
    document = docx.Document(r'C:\Users\j20687\Desktop\answersheet.docx')
    # document = docx.Document(r'C:\Users\j20687\Desktop\answersheet_dp.docx')
    # 设置整个文档的默认字体
    microsoft_font = u'宋体'  # u 表示后面的字符串以 Unicode 格式进行编码
    black_font = u'黑体'
    number_font = 'Times New Roman'
    area = qn('w:eastAsia')
    document.styles['Normal'].font.name = microsoft_font
    document.styles['Normal'].font.size = Pt(10.5)
    document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
    document.styles['Normal']._element.rPr.rFonts.set(area, number_font)
    # 指定段落样式
    styles = document.styles
    s = styles.add_style('lbj_Style', WD_STYLE_TYPE.PARAGRAPH)
    s.font.name = microsoft_font
    s.font.size = Pt(10.5)
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.line_spacing = Pt(0)  # 行距值
    s.paragraph_format.space_before = Pt(0)  # 段前距
    s.paragraph_format.space_after = Pt(0)  # 段后距
    s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE  # 行间距
    s._element.rPr.rFonts.set(area, number_font)  # 除中文外其它文字 使用的字体 ，备选项

    # 编辑页眉
    document.settings.odd_and_even_pages_header_footer = True
    section = document.sections[0]
    header = section.header
    # run = header.paragraphs[0]
    # run.style = s
    # p = run.add_run('2019学校12月月考卷试卷副标题答题卡')
    # font = p.font
    # font.name = black_font
    # font.color.rgb = RGBColor(0, 0, 0)
    # font.size = Pt(15)
    # run.alignment = WD_ALIGN_PARAGRAPH.CENTER

    type_test_no, type_sheet = 2, 1
    subject_id = "2111612056"
    main_title = "2019学校12月月考卷试卷副标题答题卡"
    add_prefix_info(header, type_test_no, type_sheet, subject_id, main_title, stu_id_count=8)

    # run = document.paragraphs[0]
    # run.style = s
    # p = run.add_run('2019学校12月月考卷试卷副标题答题卡')
    # font = p.font
    # font.name = black_font
    # font.color.rgb = RGBColor(0, 0, 0)
    # font.size = Pt(15)
    # run.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # p = document.add_paragraph()
    # p.add_run().add_picture(r'D:\校本资源相关\答题卡\scan\scan\prefix_.png')

    p = document.paragraphs[0]
    # p.add_run().add_break()
    run = p.add_run('选择题（请用2B铅笔填涂）')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = black_font
    width5 = (Cm(0.59), Cm(5.34), Cm(5.34), Cm(5.34), Cm(0.59))
    width3 = (Cm(0.59), Cm(16.02), Cm(0.59))
    # choice_group = sheet.get('choice_question',[])
    sub_mark_num_vec = []  # 每页竖长条表示块数量
    sub_mark_num = 0  # 竖长条表示块总数量
    qus_range_choice = []  # 选择题题号范围
    qus_range_fillin = []  # 填空题题号范围
    qus_range_eassy = []  # 简答题题号范围
    sub_mark_choice = []  # 选择题黑块区域
    sub_mark_fillin = []  # 填空题黑块区域
    sub_mark_eassy = []  # 简答题黑块区域
    choice_page = 0  # 选择题所在页数
    fillin_page = 0  # 填空题所在页数
    eassy_page = 0  # 简答题所在页数
    page_num = 0  # 页数
    ques_no = 0  # 题号
    choice_answer_cnt = 4  # 选项列数
    choice_answer_pattern = []  # 每列选择题数目
    choice_col_num = 0  # 选择题每行最多列数
    choice_row_num = 0  # 选择题行数
    # choice_len = len(choice_group)
    # fillin_group = sheet.get('fillin_question',[])
    # fillin_len = len(fillin_group)
    fillin_len = 6
    choice_len = 100
    fill_no = 0
    is_fillin_breakpage = False
    row_n, row_rest = divmod(choice_len, 15)
    choice_row_num = row_n + 1 if row_rest else row_n
    if 0 <= choice_len <= 5:
        choice_col_num = 1
        choice_answer_pattern.append(choice_len)
    elif 6 <= choice_len <= 10:
        choice_col_num = 2
        choice_answer_pattern.append(5)
        choice_answer_pattern.append(choice_len - 5)
    else:
        choice_col_num = 3
        for n in range(divmod(choice_len, 5)[0]):
            choice_answer_pattern.append(5)
        if divmod(choice_len, 5)[1]:
            choice_answer_pattern.append(divmod(choice_len, 5)[1])

    if choice_len > 150:
        ques_no = write_option(ques_no, 60, document)
        document.add_page_break()
        page_num += 1
        # add_barcode(document, page_num + 1)
        else_choice = choice_len - 60
        pre_no, rest = divmod(else_choice, 90)
        for n in range(pre_no):
            page_num += 1
            # add_barcode(document, page_num + 1)
            ques_no = write_option(ques_no, 90, document)
            document.add_page_break()
        ques_no = write_option(ques_no, rest, document)
        writeline(document, black_font)
        choice_page = page_num + 1 if rest else page_num
        ques_no, fill_no, num, page_num, is_fillin_breakpage = choice60_(fillin_len, rest, ques_no, document, page_num)
        page_num += num
        fillin_page = page_num
        first_fillin_page = page_num
    else:
        if choice_len > 60:
            else_choice = choice_len - 60
            ques_no = write_option(ques_no, 60, document)
            document.add_page_break()
            page_num += 1
            # add_barcode(document, page_num + 1)
            ques_no = write_option(ques_no, else_choice, document)
            writeline(document, black_font)
            choice_page = page_num + 1 if else_choice else page_num
            ques_no, fill_no, num, page_num, is_fillin_breakpage = choice60_(fillin_len, else_choice, ques_no, document,
                                                                             page_num)
            page_num += num
            fillin_page = page_num
        else:
            ques_no = write_option(ques_no, choice_len, document)
            writeline(document, black_font)
            choice_page = page_num + 1
            ques_no, fill_no, num, page_num, is_fillin_breakpage = choice0_60(fillin_len, choice_len, ques_no, document,
                                                                              page_num)
            page_num += num
            fillin_page = page_num

    print(ques_no, fill_no, is_fillin_breakpage)
    print(page_num, choice_page, fillin_page)
    print('page_num{}'.format(page_num))
    subjec_len = 0
    rest_subjec_len = subjec_len
    nn = 0
    if is_fillin_breakpage and fillin_len:
        if fill_no == 0:
            nn = 2 if subjec_len >= 2 else subjec_len
            rest_subjec_len = subjec_len - nn
            for n in range(nn):
                ques_no, page_num = sub_subjective(n, ques_no, page_num, document)
            eassy_page = page_num
            if rest_subjec_len > 0:
                document.add_page_break()
                page_num += 1
            # add_barcode(document, page_num)
        elif 7 >= fill_no >= 1:
            nn = 1 if subjec_len >= 1 else subjec_len
            rest_subjec_len = subjec_len - nn
            for n in range(nn):
                ques_no, page_num = sub_subjective(n, ques_no, page_num, document)
            eassy_page = page_num
            if rest_subjec_len > 0:
                document.add_page_break()
                page_num += 1
            # add_barcode(document, page_num)
        else:
            nn = 0
            if rest_subjec_len > 0:
                document.add_page_break()
                page_num += 1
            # add_barcode(document, page_num)
            eassy_page = page_num
    else:
        if subjec_len:
            document.add_page_break()
            page_num += 1
        # add_barcode(document, page_num)
        eassy_page = page_num
    div1, div2 = divmod(rest_subjec_len, 2)
    for n in range(rest_subjec_len):
        ques_no, page_num = sub_subjective(n, ques_no, page_num, document)

    write_eng(ques_no, document)
    # write_chn(ques_no, document)
    write_forbidden(document)
    # sub_page_num = 1 if div2 else 0
    # page_num += sub_page_num
    #  return sub_mark_num_vec except eassy
    if not fillin_len:
        fillin_page = 0
    if not subjec_len:
        eassy_page = 0
    # first_choice_page = 1 if choice_len else 0
    if choice_len:
        print('fillin_page', fillin_page)
        print('choice_page', choice_page)
        if fillin_page >= choice_page:
            if choice_len > 60:
                sub_mark_num_vec.append(4)
                rest_ = choice_len - 60
                a, b = divmod(rest_, 90)
                if fillin_page > choice_page:
                    sub_mark_num_vec.append(8)
                    a = a + 1
                else:
                    sub_mark_num_vec.append(8)
                for i in range(a):
                    sub_mark_num_vec.append(4)
            else:
                if fillin_page == choice_page:
                    sub_mark_num_vec.append(8)
                else:
                    sub_mark_num_vec.append(8)
                    for i in range(fillin_page - 2):
                        sub_mark_num_vec.append(4)
                    if eassy_page == fillin_page:
                        sub_mark_num_vec.append(4 + nn * 4)
                    else:
                        sub_mark_num_vec.append(4)

        else:
            if eassy_page == choice_page:
                sub_mark_num_vec.append(4)
                sub_mark_num_vec.append(8)
            else:
                for i in range(choice_page):
                    sub_mark_num_vec.append(4)
    else:
        if fillin_len:
            if fillin_page > 1:
                for i in range(fillin_page - 1):
                    sub_mark_num_vec.append(4)
                if eassy_page == fillin_page:
                    sub_mark_num_vec.append(4 + nn * 4)
                else:
                    sub_mark_num_vec.append(4)
            else:
                sub_mark_num_vec.append(4)
        else:
            sub_mark_num_vec.append(2)
    if sub_mark_num_vec:
        sub_mark_num_vec[0] = sub_mark_num_vec[0] + 2
    #  return sub_mark_num_vec with eassy
    for n in range(div1):
        sub_mark_num_vec.append(8)
    if div2 != 0:
        sub_mark_num_vec.append(4)
    #  return sub_mark for choice, fillin, eassy
    block_num = 0
    start = 2
    step = 3
    if choice_len and choice_page:
        for n in range(choice_page):
            block_num += 1
            sub_mark_choice.append(start + step * (block_num - 1) + block_num - 1)
            sub_mark_choice.append(start + step * block_num + block_num - 1)
    if fillin_page >= choice_page and divmod(choice_len, 90)[1] <= 45:
        for n in range(fillin_page):
            block_num += 1
            sub_mark_fillin.append(start + step * (block_num - 1) + block_num - 1)
            sub_mark_fillin.append(start + step * block_num + block_num - 1)
    else:
        for n in range(fillin_page - choice_page):
            block_num += 1
            sub_mark_fillin.append(start + step * (block_num - 1) + block_num - 1)
            sub_mark_fillin.append(start + step * block_num + block_num - 1)
    for n in range(subjec_len):
        block_num += 1
        sub_mark_eassy.append(start + step * (block_num - 1) + block_num - 1)
        sub_mark_eassy.append(start + step * block_num + block_num - 1)

    #  return sub_mark_num
    for num in sub_mark_num_vec:
        sub_mark_num += num

    print(sub_mark_choice)
    print(sub_mark_fillin)
    print(sub_mark_eassy)
    print(sub_mark_num)
    print('sub_mark_num_vec', sub_mark_num_vec)
    print('page_num{}'.format(page_num))
    print(choice_answer_cnt, choice_answer_pattern, choice_col_num)
    #
    # add_postfix_info(document, 3)
    # print(len(document.sections))

    document.save(r'C:\Users\j20687\Desktop\answersheet1.docx')
