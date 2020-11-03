import os
import re
import threading
import time
from io import BytesIO

import docx
import lxml.etree as ET
from pydocx import PyDocX

'''
OMML2MML_XSL是OMML转MathML的XSL文件，
一般在'C:\Program Files (x86)\Microsoft Office\Office15\OMML2MML.XSL'可找到，可复制到项目目录里
'''
OMML2MML_XSL = r"C:\Program Files (x86)\Microsoft Office\Office15\OMML2MML.XSL"


def read4word(file_path):
    '''
    :param file: 为传入的.docx文件对象,以二进制格式打开,
    若传入的是doc文件，则先转为docx，但是转换文档后的公式会变成图片格式。所以传入文件最好为docx文档
    :return:
    '''
    if file_path.split('.')[-1] == 'doc':
        file_name = file_path.split("/")[-1]
        file_name = ".".join(file_name.split(".")[:-1])
        docx_path = r'%s.docx' % file_name
        os.system('soffice --headless --invisible --convert-to docx {} --outdir ./'.format(file_path))
        with open(docx_path, 'rb') as f:
            Exam = read_word(f)
        os.remove(docx_path)
    else:
        with open(file_path, "rb") as f:
            Exam = read_word(f)
    return Exam


def read_word(file):
    proxy = []
    doc = docx.Document(file)
    for para in doc.paragraphs:
        proxy.append(para._element.xml)  # 返回docx文档的xml文件

    threads = []
    # q = Queue
    q = {}
    t1 = threading.Thread(target=getMathml, args=(proxy, 'mmls_in_para', q))
    threads.append(t1)
    html = PyDocX.to_html(file)
    t2 = threading.Thread(target=get_img, args=(html, 'images_in_para', 'images_in_table', q))
    threads.append(t2)
    t1.start()
    t2.start()
    t1.join()
    t2.join()
    # for t in threads:
    # 	t.setDaemon(True)
    # 	t.start()
    # t.join()
    mmls_in_para = q.get('mmls_in_para')

    images_in_para, images_in_table = q.get('images_in_para'), q.get('images_in_table')

    table_html, table_para = get_table(doc, images_in_table)  # 返回表格数据

    paragraphs = get_para_html(proxy, mmls_in_para, images_in_para, table_html, table_para)  # 得到文档的段落信息，返回的是HTML标签

    return paragraphs


def getMathml(proxy, string1, q, flag=True):
    '''
    遍历xml，找到m:oMath
    :param element: xml文件根节点
    :return:
    '''

    mmls = []
    ommls = []
    re_math = re.compile(r'<m:oMath>(.*?)</m:oMath>', re.S)
    i = 0
    for xml in proxy:
        omml = re_math.findall(xml)
        if omml != []:
            for ss in omml:
                string = (str(ss))
                str1 = '<xml-fragment xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml" xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk" xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml" xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:w10="urn:schemas-microsoft-com:office:word" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas">'
                string = str1 + string + '</xml-fragment>'
                i += 1
                ommls.append(bytes(string, encoding='utf-8'))
    if ommls:
        xslt = ET.parse(OMML2MML_XSL)
        transform = ET.XSLT(xslt)
        omml2mml(ommls, transform, mmls)
    if flag:
        q[string1] = mmls
    else:
        return mmls


def omml2mml(ommls, transform, mmls):
    for omml in ommls:
        dom = ET.parse(BytesIO(omml))
        newdom = transform(dom)
        string = str(ET.tostring(newdom, pretty_print=False), encoding='utf-8')
        string = string.replace("xmlns:m=\"http://schemas.openxmlformats.org/officeDocument/2006/math\"", "")
        string = string.replace("xmlns:mml", "xmlns")
        string = string.replace("mml:", "").replace("\n", "")
        mmls.append(string)


def get_table(doc, images_in_table):
    """

    :param doc:原文档
    :param images_in_table: 表格中的图片
    :return: table_html：表格的html标签, table_para：表格所在的段落数
    """
    doc_xml = doc._element.xml
    root = ET.fromstring(doc_xml)
    table_xml = []
    p_tab = []  # 表的单元格数
    table_html = {}
    table_para = []
    table_para_count = []
    i = 0
    j = 0
    for p in root.iter():
        if p.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p':
            i += 1
        if p.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl':
            table_para.append(i)
    for table in doc.tables:
        tab = 0
        xml = table._element.xml
        root1 = ET.fromstring(xml)
        para_count = 0
        for para in root1.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
            para_count += 1
        table_para_count.append(max(para_count, len(table.rows)))
        table_xml.append(xml)
        for row in table.rows:
            for cell in row.cells:
                tab += 1
        p_tab.append(tab)
    for n in range(len(p_tab)):
        if n >= 1:
            n_para = table_para[n]
            for i in range(n):
                n_para -= max(p_tab[i], table_para_count[i])
            n_para += n
            table_para[n] = n_para
    table_para.append(-1)
    mmls_in_table = getMathml(table_xml, "a", {}, flag=False)
    for tbl in root.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl'):
        str_table = ''
        result = []
        result.append('<table border="1">')
        getTableText(tbl, result, mmls_in_table, images_in_table)
        result.append('</tr>')
        result.append('</table>')
        result.remove(result[1])
        for string in result:
            str_table += string
        str_split = str_table.split('</td>')
        str_table = ''
        for s in str_split:
            if s[0:4] in ['</tr', '<td ']:
                str_table += '</td>'
                str_table += s
            else:
                str_table += s
        table_html[table_para[j]] = str_table
        j += 1
    return table_html, table_para


def getTableText(p, result, mmls_in_table, images_in_table):
    if len(list(p)) == 0:
        if p.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t':
            tmp = p.text
            result.append(tmp)
            result.append('</td>')
    else:
        for child in list(p):
            if child.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr':
                result.append('</tr>')
                result.append('<tr>')
            if child.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc':
                result.append('<td ')
                tcw = []
                for sub in child.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcW'):
                    tcw.append(sub.tag)
                if not tcw:
                    result.append('>')
            if child.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcW':
                attrib = child.attrib
                width = attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w']
                result.append('width=\"%spx;\"></td>' % width)
            if child.tag == '{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath':
                if mmls_in_table:
                    result.append(mmls_in_table[0])
                    result.append('</td>')
                    mmls_in_table.pop(0)
            elif child.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing' \
                    or child.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict':
                if images_in_table:
                    result.append(images_in_table[0])
                    result.append('</td>')
                    images_in_table.pop(0)
            elif child.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object':
                if images_in_table:
                    result.append(images_in_table[0])
                    result.append('</td>')
                    images_in_table.pop(0)
            else:
                getTableText(child, result, mmls_in_table, images_in_table)


def get_para_html(proxy, mmls, images_in_para, table_html, table_para):
    """

    :param proxy:docx的xml文档
    :param mmls: 段落中的公式
    :param images_in_para:段落中的图片
    :param table_html:表格html标签
    :param table_para:表格所在段落号
    :return: 返回word段落带p标签的html格式
    """
    results = []
    paragraphs = []  # 返回word段落带p标签的html格式
    i = 1
    para = 0
    for p in proxy:
        result = []
        root = ET.fromstring(p)
        result = getNodeText(root, result, mmls, images_in_para)
        results.append(result)
    for s in results:
        str_para = ''
        for ss in s:
            str_para += str(ss)
        str_para = '<p>' + str_para + '</p>'
        paragraphs.append(str_para)
        if i == table_para[para]:
            paragraphs.append(table_html[table_para[para]])
            para += 1
            i += 1
        i += 1
    return paragraphs


def getNodeText(ele, result, mmls, img_in_docx):
    """

    :param ele:根节点
    :param result: 中间参数
    :param mmls: 公式
    :param img_in_docx:文档中总的图片
    :return: 返回表格数据
    """
    if len(list(ele)) == 0:
        walkdata(ele, result)
    else:
        for child in list(ele):
            if child.tag == '{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath':
                if mmls:
                    result.append(mmls[0])
                    mmls.pop(0)
            elif child.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing' \
                    or child.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict':
                pic = []
                # style = None
                for sub in child.iter('{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}docPr'):
                    attrib = sub.attrib
                    name = attrib['name']
                    if "图片" not in str(name):
                        pic.append(name)
                # for sub in child:
                #     if sub.tag == '{urn:schemas-microsoft-com:vml}shape':
                #         attrib = sub.attrib
                #         style = attrib['style']
                if not pic and img_in_docx:
                    img = img_in_docx[0]
                    # img_list = re.split(re.compile(r'(<.*?src=\".*?\")'), img)
                    # if style:
                    #     img = img_list[1] + " style=\"{}\"".format(style) + img_list[2]
                    result.append(img)
                    img_in_docx.pop(0)
            elif child.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object':  # OLE对象，比如mathtype公式
                if img_in_docx:
                    result.append(img_in_docx[0])
                    img_in_docx.pop(0)
            else:
                getNodeText(child, result, mmls, img_in_docx)
    return result


def walkdata(child, result):
    #  遍历
    if child.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bookmarkStart':
        attrib = child.attrib
        id_num = attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id']
        if id_num != "0":
            result.append('{}.'.format(id_num))
    if child.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t':
        # tmp = str(child.text).replace(' ', '&nbsp;')
        tmp = str(child.text)
        result.append(tmp)


def get_img(html, str1, str2, q):
    """

    :param file:文件对象
    :return:返回文档中的图片
    """
    soup = ET.HTML(html)
    images_in_para = []
    images_in_table = []

    for img in soup.xpath(u"//img"):
        # src = img.get("src")
        # print(src)
        parent = img.getparent().getparent().getparent().getparent()
        if 'table' == parent.tag:
            img.set('src', "$$$$$$$$$$$$$$$$")
            images_in_table.append(str(ET.tostring(img)))
        else:
            img.set('src', "aaaaaaaaaaaa")
            print(ET.tostring(img))
            images_in_para.append(str(ET.tostring(img)))
    #     reg = re.compile('data.*?/(.*?);', re.S)
    #     style_img = reg.findall(img['src'])[0]
    #     strg = img['src'].replace("data:image/wmf;base64,", "").replace("data:image/jpeg;base64,", "")
    #     byte = base64.urlsafe_b64decode(strg)
    #     t0 = int(round(time.time() * 1000))
    #     tmp_path = '/tmp/%d.%s' % (t0, str(style_img))
    #     with open(tmp_path, 'wb') as file:
    #         file.write(byte)
    #     if style_img == 'wmf':
    #         t1 = int(round(time.time() * 1000))
    #         png_path = '/tmp/%d.png' % t1
    #         os.system('convert %s %s' % (tmp_path, png_path))
    #         f = open(png_path, 'rb')
    #         url = put(f)
    #         f.close()
    #         img['src'] = url
    #         if img.find_parents('table') != []:
    #             images_in_table.append(img)
    #         else:
    #             images_in_para.append(img)
    #         os.remove(png_path)
    #         os.remove(tmp_path)
    #     else:
    #         f = open(tmp_path, 'rb')
    #         url = put(f)
    #         f.close()
    #         img['src'] = url
    #         if img.find_parents('table') != []:
    #             images_in_table.append(img)
    #         else:
    #             images_in_para.append(img)
    #         os.remove(tmp_path)
    # return images_in_para, images_in_table
    q[str1] = images_in_para
    q[str2] = images_in_table


def get_option(option, para_str, opt_line=1):
    """

    :param option:中间参数，用于保存选项的数据
    :param para_str:选项字符串数据
    :param opt_line:选项的行数
    :return:匹配答案选项
    """
    re_a = re.compile(r'A\.(.*?)B\.|A\.(.*?)</p>|A、(.*?)B、|A、(.*?)</p>|A\.(.*?)B、|A、(.*?)B\.', re.S | re.M)
    re_b = re.compile(r'B\.(.*?)C\.|B\.(.*?)</p>|B、(.*?)C、|B、(.*?)</p>|B\.(.*?)C、|B、(.*?)C\.', re.S | re.M)
    re_c = re.compile(r'C\.(.*?)D\.|C\.(.*?)</p>|C、(.*?)D、|C、(.*?)</p>|C\.(.*?)D、|C、(.*?)D\.', re.S | re.M)
    re_d = re.compile(r'D\.(.*?)E\.|D\.(.*?)</p>|D、(.*?)E、|D、(.*?)</p>|D\.(.*?)E、|D、(.*?)E\.', re.S | re.M)
    re_e = re.compile(r'E\.(.*?)F\.|E\.(.*?)</p>|E、(.*?)F、|E、(.*?)</p>|E\.(.*?)F、|E、(.*?)F\.', re.S | re.M)
    re_f = re.compile(r'F\.(.*?)G\.|F\.(.*?)</p>|F、(.*?)G、|F、(.*?)</p>|F\.(.*?)G、|F、(.*?)G\.', re.S | re.M)
    re_g = re.compile(r'G\.(.*?)</p>|G、(.*?)</p>', re.S | re.M)
    rest_a = re.compile(r'<p>(.*?)B\.|<p>(.*?)B、', re.S | re.M)
    rest_b = re.compile(r'<p>(.*?)C\.|<p>(.*?)C、', re.S | re.M)
    rest_c = re.compile(r'<p>(.*?)D\.|<p>(.*?)D、', re.S | re.M)
    rest_d = re.compile(r'<p>(.*?)E\.|<p>(.*?)E、', re.S | re.M)
    rest_e = re.compile(r'<p>(.*?)F\.|<p>(.*?)F、', re.S | re.M)
    rest_f = re.compile(r'<p>(.*?)G\.|<p>(.*?)G、', re.S | re.M)
    re_newline = re.compile('.*?A\..*?|.*?B\..*?|.*?C\..*?|.*?D\..*?|.*?E\..*?|.*?F\..*?|.*?G\..*?'
                            '|.*?A、.*?|.*?B、.*?|.*?C、.*?|.*?D、.*?|.*?E、.*?|.*?F、.*?|.*?G、.*?')
    aa = {}
    bb = {}
    cc = {}
    dd = {}
    ee = {}
    ff = {}
    gg = {}
    opt = para_str.replace("【选项】", "")
    if opt != "":
        A = re_a.findall(opt)
        B = re_b.findall(opt)
        C = re_c.findall(opt)
        D = re_d.findall(opt)
        E = re_e.findall(opt)
        F = re_f.findall(opt)
        G = re_g.findall(opt)
        newline = re_newline.findall(opt)
        if A:
            aa['key'] = 'A'
            aa['value'] = str(A[0][0] + A[0][1] + A[0][2] + A[0][3] + A[0][4] + A[0][5]).replace('&nbsp;', ' ')
            option.append(aa)
        if B:
            rest = rest_a.findall(opt)
            if not A and rest:
                rest_str = str(rest[0][0] + rest[0][1]).replace('&nbsp;', ' ')
                if option:
                    remap = option.pop()
                    remap['value'] += rest_str
                    option.append(remap)
            bb['key'] = 'B'
            bb['value'] = str(B[0][0] + B[0][1] + B[0][2] + B[0][3] + B[0][4] + B[0][5]).replace('&nbsp;', ' ')
            option.append(bb)
        if C:
            rest = rest_b.findall(opt)
            if not B and rest:
                rest_str = str(rest[0][0] + rest[0][1]).replace('&nbsp;', ' ')
                if option:
                    remap = option.pop()
                    remap['value'] += rest_str
                    option.append(remap)
            cc['key'] = 'C'
            cc['value'] = str(C[0][0] + C[0][1] + C[0][2] + C[0][3] + C[0][4] + C[0][5]).replace('&nbsp;', ' ')
            option.append(cc)
        if D:
            rest = rest_c.findall(opt)
            if not C and rest:
                rest_str = str(rest[0][0] + rest[0][1]).replace('&nbsp;', ' ')
                if option:
                    remap = option.pop()
                    remap['value'] += rest_str
                    option.append(remap)
            dd['key'] = 'D'
            dd['value'] = str(D[0][0] + D[0][1] + D[0][2] + D[0][3] + D[0][4] + D[0][5]).replace('&nbsp;', ' ')
            option.append(dd)
        if E:
            rest = rest_d.findall(opt)
            if not D and rest:
                rest_str = str(rest[0][0] + rest[0][1]).replace('&nbsp;', ' ')
                if option:
                    remap = option.pop()
                    remap['value'] += rest_str
                    option.append(remap)
            ee['key'] = 'E'
            ee['value'] = str(E[0][0] + E[0][1] + E[0][2] + E[0][3] + E[0][4] + E[0][5]).replace('&nbsp;', ' ')
            option.append(ee)
        if F:
            rest = rest_e.findall(opt)
            if not E and rest:
                rest_str = str(rest[0][0] + rest[0][1]).replace('&nbsp;', ' ')
                if option:
                    remap = option.pop()
                    remap['value'] += rest_str
                    option.append(remap)
            ff['key'] = 'F'
            ff['value'] = str(F[0][0] + F[0][1] + F[0][2] + F[0][3] + F[0][4] + F[0][5]).replace('&nbsp;', ' ')
            option.append(ff)
        if G:
            rest = rest_f.findall(opt)
            if not F and rest:
                rest_str = str(rest[0][0] + rest[0][1]).replace('&nbsp;', ' ')
                if option:
                    remap = option.pop()
                    remap['value'] += rest_str
                    option.append(remap)
            gg['key'] = 'G'
            gg['value'] = str(G[0][0] + G[0][1]).replace('&nbsp;', ' ').replace('<p>', '').replace('</p>', '')
            option.append(gg)
        if newline == [] and opt_line > 4:  # 如果选项换行
            if option:
                remap = option.pop()
                remap['value'] += opt
                option.append(remap)


# def put(f):
#     try:
#         r = file_server.upload_file(f)
#     except Exception as e:
#         logger.error("get sheet image fail: {}".format(e))
#         return ""
#     url = r.url
#     return url


if __name__ == '__main__':
    time1 = time.time()
    file = r'D:\网手阅\测试卡卷\测试卡卷\试卷test-2015年普通高等学校招生全国统一考试（浙江卷）数学（理科）-教师用卷 - 副本.docx'
    # file = r'C:\Users\j20687\Desktop\试题导入模板0422.docx'
    content = read4word(file)
    string = ""
    for i in content:
        string += i
    print(string)
