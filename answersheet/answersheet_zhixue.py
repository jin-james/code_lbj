import json
import logging
import os
import re
import time
import uuid

import pdfplumber
import requests
from MyQR import myqr
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.oxml.xmlchemy import OxmlElement
from docx.shared import Pt, RGBColor, Cm
from pdfminer.pdfparser import PDFSyntaxError

# from nile.utils.constants import CHOICE_QUESTIONS, FILL_IN_QUESTIONS

logger = logging.getLogger(__name__)
imgtmp_base_path = '/tmp/img'
CHOICE_QUESTIONS = ["选择题", "单选题", "单项选择题", "多选题", "多项选择题", "听力选择题", '双选题']  # 选择题
FILL_IN_QUESTIONS = ["填空题", "单词拼写题"]  # 填空题


def ichoice_sub_write(table, n, ques_no, pic_in_sub, isfillin=False, is_score=None):
    if isfillin:
        cell = table.cell(0, 1)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cell.paragraphs[0]
        paragraph_format = p.paragraph_format
        paragraph_format.line_spacing = Pt(40)  # 22磅行间距
        # p.add_run().add_break()
        for num in ques_no:
            p.add_run(str(num) + '题、_____________________________').font.size = Pt(9)
            if is_score:
                p.add_run().add_picture(r'D:\Program Files\code_lbj\answersheet\fillin_scorebar.png', width=Pt(30))
                p.add_run(" ")
            else:
                p.add_run("      ")
            if (num + 1) % 2 == 0:
                p.add_run().add_break()
    else:
        cell = table.cell(1, 1)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cell.paragraphs[0]
        if not pic_in_sub:
            p.add_run().add_break()
        p.add_run(str(ques_no) + '题、').font.size = Pt(9)
    if pic_in_sub:
        add_pic(table, pic_in_sub)
        n -= 3
    for i in range(n):
        p.add_run().add_break()


def set_col_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def hide_frame(table, rows, cols, ischoice=True):
    if ischoice:
        for r in table.rows:
            for c in r.cells:
                set_cell_border(
                    c,
                    top={"sz": 0, "val": "single", "color": "#FFFFFF"},
                    left={"sz": 0, "val": "single", "color": "#FFFFFF"},
                    bottom={"sz": 0, "color": "#FFFFFF", "val": "single"},
                    right={"sz": 0, "val": "single", "color": "#FFFFFF", },
                )
        set_cell_border(
            table.cell(0, 1),
            left={"sz": 5, "val": "single", "color": "#000000"},
            top={"sz": 5, "color": "#000000", "val": "single"},
        )
        set_cell_border(
            table.cell(0, 2),
            top={"sz": 5, "val": "single", "color": "#000000"},
        )
        set_cell_border(
            table.cell(0, 3),
            top={"sz": 5, "color": "#000000", "val": "single"},
            right={"sz": 5, "val": "single", "color": "#000000", },
        )
        set_cell_border(
            table.cell(rows - 1, 1),
            left={"sz": 5, "val": "single", "color": "#000000"},
            bottom={"sz": 5, "color": "#000000", "val": "single"},
        )
        set_cell_border(
            table.cell(rows - 1, 2),
            bottom={"sz": 5, "val": "single", "color": "#000000"},
        )
        set_cell_border(
            table.cell(rows - 1, 3),
            bottom={"sz": 5, "color": "#000000", "val": "single"},
            right={"sz": 5, "val": "single", "color": "#000000", },
        )
        for r in range(1, rows - 1):
            set_cell_border(
                table.cell(r, 1),
                left={"sz": 5, "val": "single", "color": "#000000", },
            )
            set_cell_border(
                table.cell(r, 3),
                right={"sz": 5, "val": "single", "color": "#000000", },
            )
    else:
        for r in range(rows):
            set_cell_border(
                table.cell(r, 0),
                top={"sz": 0, "val": "single", "color": "#FFFFFF"},
                bottom={"sz": 0, "color": "#FFFFFF", "val": "single"},
                left={"sz": 0, "val": "single", "color": "#FFFFFF"},
            )
            set_cell_border(
                table.cell(r, cols - 1),
                top={"sz": 0, "val": "single", "color": "#FFFFFF"},
                bottom={"sz": 0, "color": "#FFFFFF", "val": "single"},
                right={"sz": 0, "val": "single", "color": "#FFFFFF", },
            )


def set_frame(table):
    for r in table.rows:
        for c in r.cells:
            set_cell_border(
                c,
                top={"sz": 5, "val": "single", "color": "#000000"},
                left={"sz": 5, "val": "single", "color": "#000000"},
                bottom={"sz": 5, "color": "#000000", "val": "single"},
                right={"sz": 5, "val": "single", "color": "#000000", },
            )


def hide_frame_single(table):
    for r in table.rows:
        for c in r.cells:
            set_cell_border(
                c,
                top={"sz": 0, "val": "single", "color": "#FFFFFF"},
                left={"sz": 0, "val": "single", "color": "#FFFFFF"},
                bottom={"sz": 0, "color": "#FFFFFF", "val": "single"},
                right={"sz": 0, "val": "single", "color": "#FFFFFF", },
            )


def add_options(p, option_no):
    p.add_run().add_picture(r'D:\Program Files\code_lbj\answersheet\sel-a_.bmp', width=Pt(12.5))
    p.add_run(' ').add_picture(r'D:\Program Files\code_lbj\answersheet\sel-b_.bmp', width=Pt(12.5))
    p.add_run(' ').add_picture(r'D:\Program Files\code_lbj\answersheet\sel-c_.bmp', width=Pt(12.5))
    p.add_run(' ').add_picture(r'D:\Program Files\code_lbj\answersheet\sel-d_.bmp', width=Pt(12.5))
    if int(option_no) == 5:
        p.add_run(' ').add_picture(r'D:\Program Files\code_lbj\answersheet\sel-e_.bmp', width=Pt(12.5))
    if int(option_no) == 6:
        p.add_run(' ').add_picture(r'D:\Program Files\code_lbj\answersheet\sel-e_.bmp', width=Pt(12.5))
        p.add_run(' ').add_picture(r'D:\Program Files\code_lbj\answersheet\sel-f_.bmp', width=Pt(12.5))
    if int(option_no) == 7:
        p.add_run(' ').add_picture(r'D:\Program Files\code_lbj\answersheet\sel-e_.bmp', width=Pt(12.5))
        p.add_run(' ').add_picture(r'D:\Program Files\code_lbj\answersheet\sel-f_.bmp', width=Pt(12.5))
        p.add_run(' ').add_picture(r'D:\Program Files\code_lbj\answersheet\sel-g_.bmp', width=Pt(12.5))


def add_mark(cols, table, rows=None):
    cell1 = table.cell(0, 0)
    cell1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p1 = cell1.paragraphs[0]
    p1.add_run().add_picture(r'D:\Program Files\code_lbj\answersheet\Rectangle-c.png', width=Pt(6))
    p1.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell2 = table.cell(0, cols + 1)
    cell2.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p2 = cell2.paragraphs[0]
    p2.add_run().add_picture(r'D:\Program Files\code_lbj\answersheet\Rectangle-c.png', width=Pt(6))
    p2.alignment = WD_TABLE_ALIGNMENT.RIGHT
    if rows:
        cell1 = table.cell(rows - 1, 0)
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p1 = cell1.paragraphs[0]
        p1.add_run().add_picture(r'D:\Program Files\code_lbj\answersheet\Rectangle-c.png', width=Pt(6))
        p1.alignment = WD_TABLE_ALIGNMENT.LEFT
        cell2 = table.cell(rows - 1, cols + 1)
        cell2.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p2 = cell2.paragraphs[0]
        p2.add_run().add_picture(r'D:\Program Files\code_lbj\answersheet\Rectangle-c.png', width=Pt(6))
        p2.alignment = WD_TABLE_ALIGNMENT.RIGHT


def add_pic(table_sub, pic_in_sub):
    # table_sub.alignment = WD_TABLE_ALIGNMENT.CENTER
    para = table_sub.cell(0, 1).paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    re_image = re.compile("<img.*?>", re.S | re.I)
    re_src = re.compile("src=\"(.*?)\"", re.S | re.I)
    re_height = re.compile("height=\"(.*?)\"", re.S | re.I)
    line = [v for v in pic_in_sub if v]
    line = line[0] if line else ""
    image = re_image.findall(line)
    for img in image:
        string = re_src.search(img)
        string = string[1] if string else ""
        hei_str = re_height.search(img)
        hei_str = hei_str[1] if hei_str else ""
        content = requests.get(string).content
        subfix = string[-32:]
        path = imgtmp_base_path + '/' + subfix + '{}'.format('.png')
        height_v = float(hei_str.split('p')[0])
        if os.path.exists(path):
            para.add_run().add_picture(path, height=Pt(height_v * 0.7))
        with open(path, 'wb')as f:
            f.write(content)
        para.add_run().add_picture(path, height=Pt(height_v * 0.7))
    for i in range(len(pic_in_sub)):
        tmp = pic_in_sub[i]
        if not tmp:
            pic_in_sub[i] = []  # 剔除已经添加的图片数据
            break


def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def write_option(choice_num, choice_opt, choice_answer_pattern, rows, document):
    width5 = (Cm(0.59), Cm(5.73), Cm(5.73), Cm(5.73), Cm(0.59))
    choice_len = len(choice_num)
    if choice_len:
        cols = 3
        table_choice = document.add_table(rows=rows, cols=cols + 2, style='Table Grid')
        set_col_widths(table_choice, width5)
        hide_frame(table_choice, rows, cols + 2)
        c = 1
        r = 0
        p = table_choice.cell(r, c).paragraphs[0]
        p.add_run().add_break()
        start = 0
        for count in choice_answer_pattern:
            ques_no = 0
            num = choice_num[start: start + count]
            option = choice_opt[start: start + count]
            choice_length = len(num)
            for i in range(choice_length):
                ques_no += 1
                choice_no = num[i]
                option_no = option[i]
                if choice_no >= 100:
                    p.add_run(str(choice_no) + " ").font.size = Pt(9)
                elif choice_no >= 10:
                    p.add_run(' ' + str(choice_no) + ' ').font.size = Pt(9)
                else:
                    p.add_run("  " + str(choice_no) + " ").font.size = Pt(9)
                add_options(p, option_no)
                if ques_no != choice_length:
                    p.add_run().add_break()
            if c < 3:
                c += 1
                p.add_run().add_break()  # 另起一列的时候在前一列末加一空行
                p = table_choice.cell(r, c).paragraphs[0]
                p.add_run().add_break()  # 另起一列的时候在前加一空行
            else:
                if r < rows - 1:
                    r += 1
                    c = 1
                    p.add_run().add_break()  # 另起一行的时候在前一单元格末加一空行
                    p = table_choice.cell(r, c).paragraphs[0]
                    p.add_run().add_break()  # 另起一行的时候在前加一空行
            start += count
        add_mark(cols, table_choice)


def sub_subjective(number, document, pic_in_sub, score, simi_score):
    score = 28
    simi_score = 1
    width3 = (Cm(0.59), Cm(17.19), Cm(0.59))
    rows = 2
    cols = 1
    table_sub = document.add_table(rows=rows, cols=cols + 2, style='Table Grid')
    set_col_widths(table_sub, width3)
    if score is not None:
        add_mark(cols, table_sub)
        hide_frame(table_sub, rows, cols + 2, ischoice=False)
        cell0 = table_sub.cell(0, 1)
        cell0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        sub_tbl = cell0.add_table(rows=1, cols=22)
        # for cell in sub_tbl.column_cells(0):
        #     cell.width = Pt(28)
        if simi_score:
            par = sub_tbl.cell(0, 21).paragraphs[0]
            pf = par.add_run("{}".format(0.5))
            font = pf.font
            font.size = Pt(7.5)
        if 0 < score <= 20:
            index = 20
            for i in range(score + 1):
                par = sub_tbl.cell(0, index).paragraphs[0]
                pf = par.add_run("{}".format(i))
                font = pf.font
                font.size = Pt(9)
                index -= 1
        else:
            t, s = divmod(score, 10)
            index = 0
            for i in range(20, 10, -1):
                par = sub_tbl.cell(0, i).paragraphs[0]
                pf = par.add_run("{}".format(index))
                font = pf.font
                font.size = Pt(9)
                index += 1
            index = 9
            for j in range(int(t + 1)):
                par = sub_tbl.cell(0, index).paragraphs[0]
                pf = par.add_run("{}".format(j))
                font = pf.font
                font.size = Pt(9)
                index -= 1
        for r in sub_tbl.rows:
            for c in r.cells:
                set_cell_border(
                    c,
                    top={"sz": 5, "val": "single", "color": "#000000"},
                    left={"sz": 5, "val": "single", "color": "#000000"},
                    bottom={"sz": 5, "color": "#000000", "val": "single"},
                    right={"sz": 5, "val": "single", "color": "#000000", },
                )
        set_cell_border(
            table_sub.cell(0, 1),
            left={"sz": 0, "val": "single", "color": "#FFFFFF"},
            right={"sz": 0, "val": "single", "color": "#FFFFFF", },
            top={"sz": 0, "val": "single", "color": "#FFFFFF", },
            bottom={"sz": 0, "val": "single", "color": "#FFFFFF", },
        )
        set_cell_border(
            table_sub.cell(0, 0),
            right={"sz": 0, "val": "single", "color": "#FFFFFF", },
        )
        set_cell_border(
            table_sub.cell(0, 2),
            left={"sz": 0, "val": "single", "color": "#FFFFFF"},
        )
        table_sub.cell(1, 0).merge(table_sub.cell(0, 0))
        table_sub.cell(1, 2).merge(table_sub.cell(0, 2))

    else:
        table_sub.cell(1, 1).merge(table_sub.cell(0, 1))
        add_mark(cols, table_sub)
        hide_frame(table_sub, rows, cols + 2, ischoice=False)
    ichoice_sub_write(table_sub, 18, number, pic_in_sub, is_score=score)


def sub_fillin(fillin_len, fs_num_list, document, scan_type):
    width3 = (Cm(0.59), Cm(17.19), Cm(0.59))
    if fillin_len:
        rows = 1
        cols = 1
        table_fillin = document.add_table(rows=rows, cols=cols + 2, style='Table Grid')
        set_col_widths(table_fillin, width3)
        hide_frame(table_fillin, rows, cols + 2, ischoice=False)
        ichoice_sub_write(table_fillin, 0, fs_num_list, [], isfillin=True, is_score=scan_type)
        add_mark(cols, table_fillin)


def writeline(document, black_font):
    p = document.add_paragraph()
    # p.add_run().add_break()
    run = p.add_run('非选择题（请在各试题的答题区内作答）')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = black_font


def get_ques_range(choice_num, qus_range_choice):
    sort_index_ = {}
    qus_range_choice_ = []
    for item in qus_range_choice:
        if item:
            start = item[0]
            sort_index_[start] = item
    for i in choice_num:
        for key, value in sort_index_.items():
            if i == key:
                qus_range_choice_.append(value)
    return qus_range_choice_


def add_prefix_info(header, type_test_no, type_sheet, subject_id, main_title, stu_id_count=8):
    '''
    :param header:
    :param type_test_no: 0:条形码，1：准考证，2：学籍号，3：短考号
    :param main_title: 标题
    :param type_sheet: 0:网阅，1：手阅
    :param stu_id_count: 学号长度，默认为8位
    :param subject_id: 学科识别号
    :return:
    '''
    if type_sheet == 0:
        if type_test_no == 0 or type_test_no == 1:
            table_pre = header.add_table(rows=2, cols=2, width=Cm(22))
            table_pre.alignment = WD_TABLE_ALIGNMENT.CENTER
            # for i in range(2):
            #     for cell in table_pre.column_cells(i):
            #         cell.width = Cm(0.86)
            #     for cell in table_pre.row_cells(i):
            #         cell.width = Cm(0.86)
            prefix_title(table_pre, main_title)
            # p.add_run().add_break()
            table_pre.cell(0, 0).merge(table_pre.cell(0, 1))
            cell10 = table_pre.cell(1, 0)
            p = cell10.paragraphs[0]
            run = p.add_run("姓名：____________  班级：____________")
            run.font.bold = True
            run.font.name = u'黑体'
            p.add_run().add_break()
            run = p.add_run("考场/座位号：_______________")
            run.font.bold = True
            run.font.name = u'黑体'
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run.add_break()
            if type_test_no == 0:
                prefix_notice_1(cell10)
                cell11 = table_pre.cell(1, 1)
                prefix_codebar(cell11)
                prefix_absent(cell11)
                p = cell10.paragraphs[-1]
                p.add_run().add_break()
            else:
                prefix_notice_2(cell10)
                prefix_absent(cell10)
                cell11 = table_pre.cell(1, 1)
                prefix_admission_card(cell11, stu_id_count)
                p = cell11.paragraphs[-1]
                p.add_run().add_break()
        else:
            table_pre = header.add_table(rows=3, cols=2, width=Cm(22))
            table_pre.alignment = WD_TABLE_ALIGNMENT.CENTER
            prefix_title(table_pre, main_title)
            table_pre.cell(0, 0).merge(table_pre.cell(0, 1))
            cell10 = table_pre.cell(1, 0)
            p = cell10.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("姓名：____________  班级：____________  考场/座位号：_______________")
            run.font.bold = True
            run.font.name = u'黑体'
            table_pre.cell(1, 0).merge(table_pre.cell(1, 1))
            cell20 = table_pre.cell(2, 0)
            prefix_notice_3(cell20)
            prefix_absent(cell20)
            cell21 = table_pre.cell(2, 1)
            prefix_admission_card(cell21, stu_id_count)
            p = cell21.paragraphs[-1]
            p.add_run().add_break()

    else:
        path = get_2dcode(subject_id)
        if type_test_no == 0 or type_test_no == 3:
            table_pre = header.add_table(rows=2, cols=2, width=Cm(22))
            table_pre.alignment = WD_TABLE_ALIGNMENT.CENTER
            prefix_title(table_pre, main_title)
            table_pre.cell(0, 0).merge(table_pre.cell(0, 1))
            cell10 = table_pre.cell(1, 0)
            prefix_2dcode(cell10, path)
            cell11 = table_pre.cell(1, 1)
            if type_test_no == 0:
                prefix_codebar(cell11)
            else:
                prefix_short_card(cell11)
        if type_test_no == 1:
            table_pre = header.add_table(rows=2, cols=2, width=Cm(22))
            table_pre.alignment = WD_TABLE_ALIGNMENT.CENTER
            prefix_title(table_pre, main_title)
            cell10 = table_pre.cell(1, 0)
            prefix_2dcode(cell10, path)
            table_pre.cell(0, 1).merge(table_pre.cell(1, 1))
            cell01 = table_pre.cell(0, 1)
            prefix_admission_card(cell01, stu_id_count)
            p = cell01.paragraphs[-1]
            p.add_run().add_break()
        if type_test_no == 2:
            table_pre = header.add_table(rows=3, cols=2, width=Cm(22))
            table_pre.alignment = WD_TABLE_ALIGNMENT.CENTER
            prefix_title(table_pre, main_title)
            table_pre.cell(0, 0).merge(table_pre.cell(0, 1))
            cell10 = table_pre.cell(1, 0)
            p = cell10.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("姓名：____________  班级：____________  考场/座位号：_______________")
            run.font.bold = True
            run.font.name = u'黑体'
            table_pre.cell(1, 0).merge(table_pre.cell(1, 1))
            cell20 = table_pre.cell(2, 0)
            prefix_2dcode_(cell20, path)
            cell21 = table_pre.cell(2, 1)
            prefix_admission_card(cell21, stu_id_count)
            p = cell21.paragraphs[-1]
            p.add_run().add_break()


def get_2dcode(stu_id):
    uid = uuid.uuid4()
    path = '{}.png'.format(uid)
    myqr.run(
        words=stu_id,
        # 扫描二维码后，显示的内容，或是跳转的链接
        version=5,  # 设置容错率
        # level='H',  # 控制纠错水平，范围是L、M、Q、H，从左到右依次升高
        # picture='we.png',  # 图片所在目录，可以是动图
        # colorized=False,  # 黑白(False)还是彩色(True)
        # contrast=1.0,  # 用以调节图片的对比度，1.0 表示原始图片。默认为1.0。
        # brightness=1.0,  # 用来调节图片的亮度，用法同上。
        save_name=path,  # 控制输出文件名，格式可以是 .jpg， .png ，.bmp ，.gif
    )
    return path


def prefix_title(table_pre, main_title):
    hide_frame_single(table_pre)
    cell00 = table_pre.cell(0, 0)
    p = cell00.paragraphs[0]
    run = p.add_run(main_title)
    run.font.size = Pt(14)
    run.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def prefix_notice_1(cell):
    tbl_sub = cell.add_table(rows=1, cols=1)
    tbl_sub.cell(0, 0).width = Pt(350)
    set_frame(tbl_sub)
    set_cell_border(
        tbl_sub.cell(0, 0),
        bottom={"sz": 5, "color": "#000000", "val": "single"},
    )
    p = tbl_sub.cell(0, 0).paragraphs[0]
    paragraph_format = p.paragraph_format
    paragraph_format.line_spacing = Pt(12)  # 12磅行间距
    run = p.add_run("注意事项")
    run.font.bold = True
    p.add_run().add_break()
    run = p.add_run("1．答题前，考生先将自己的姓名、班级、考场填写清楚，并认真核对条形码上")
    run.font.size = Pt(9)
    p.add_run().add_break()
    run = p.add_run("的姓名和准考证号。")
    run.font.size = Pt(9)
    p.add_run().add_break()
    run = p.add_run("2．选择题部分请按题号用2B铅笔填涂方框，修改时用橡皮擦干净，不留痕迹。")
    run.font.size = Pt(9)
    p.add_run().add_break()
    run = p.add_run("3．非选择题部分请按题号用0.5毫米黑色墨水签字笔书写，否则作答无效。要")
    run.font.size = Pt(9)
    p.add_run().add_break()
    run = p.add_run("求字体工整、笔迹清晰。作图时，必须用2B铅笔，并描浓。")
    run.font.size = Pt(9)
    p.add_run().add_break()
    run = p.add_run("4．在草稿纸、试题卷上答题无效。")
    run.font.size = Pt(9)
    p.add_run().add_break()
    run = p.add_run("5．请勿折叠答题卡,保持字体工整、笔迹清晰、卡面清洁。")
    run.font.size = Pt(9)


def prefix_notice_2(cell):
    tbl_sub = cell.add_table(rows=1, cols=1)
    tbl_sub.cell(0, 0).width = Pt(350)
    set_frame(tbl_sub)
    set_cell_border(
        tbl_sub.cell(0, 0),
        bottom={"sz": 5, "color": "#000000", "val": "single"},
    )
    p = tbl_sub.cell(0, 0).paragraphs[0]
    paragraph_format = p.paragraph_format
    paragraph_format.line_spacing = Pt(12)  # 12磅行间距
    run = p.add_run("注意事项")
    run.font.bold = True
    p.add_run().add_break()
    run = p.add_run("1．答题前请将姓名、班级、考场、准考证号填写清楚。")
    run.font.size = Pt(9)
    p.add_run().add_break()
    run = p.add_run("2．客观题答题，必须使用2B铅笔填涂，修改时用橡皮擦干净。")
    run.font.size = Pt(9)
    p.add_run().add_break()
    run = p.add_run("3．主观题答题，必须使用黑色签字笔书写。")
    run.font.size = Pt(9)
    p.add_run().add_break()
    run = p.add_run("4．必须在题号对应的答题区域内作答，超出答题区域书写无效。")
    run.font.size = Pt(9)
    p.add_run().add_break()
    run = p.add_run("5．保持答卷清洁、完整。")
    run.font.size = Pt(9)


def prefix_notice_3(cell):
    tbl_sub = cell.add_table(rows=1, cols=1)
    tbl_sub.cell(0, 0).width = Cm(6)
    set_frame(tbl_sub)
    set_cell_border(
        tbl_sub.cell(0, 0),
        bottom={"sz": 5, "color": "#000000", "val": "single"},
    )
    p = tbl_sub.cell(0, 0).paragraphs[0]
    paragraph_format = p.paragraph_format
    paragraph_format.line_spacing = Pt(12)  # 12磅行间距
    run = p.add_run("注意事项")
    run.font.bold = True
    p.add_run().add_break()
    run = p.add_run("1．答题前请将姓名、班级、考场、准考证号")
    run.font.size = Pt(9)
    run.add_break()
    run = p.add_run("填写清楚。")
    run.font.size = Pt(9)
    run.add_break()
    run = p.add_run("2．客观题答题，必须使用2B铅笔填涂，修改")
    run.font.size = Pt(9)
    run.add_break()
    run = p.add_run("用橡皮擦干净。")
    run.font.size = Pt(9)
    run.add_break()
    run = p.add_run("3.必须在题号对应的答题区域内作答，超出")
    run.font.size = Pt(9)
    run.add_break()
    run = p.add_run("答题区域书写无效。")
    run.font.size = Pt(9)
    run.add_break()


def prefix_codebar(cell):
    tbl1 = cell.add_table(rows=1, cols=1)
    tbl1.cell(0, 0).width = Cm(6.32)
    tbl1.cell(0, 0).height = Cm(2.2)
    set_cell_border(
        tbl1.cell(0, 0),
        top={"sz": 5, "val": "dashed", "color": "#000000"},
        left={"sz": 5, "val": "dashed", "color": "#000000"},
        bottom={"sz": 5, "color": "#000000", "val": "dashed"},
        right={"sz": 5, "val": "dashed", "color": "#000000", },
    )
    p = tbl1.cell(0, 0).paragraphs[0]
    run = p.add_run("贴条形码区")
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(200, 200, 200)
    p.add_run().add_break()
    p.add_run().add_break()
    run = p.add_run("(正面朝上，切勿贴出虚线方框)")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(200, 200, 200)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def prefix_absent(cell):
    tbl2 = cell.add_table(rows=1, cols=1)
    tbl2.cell(0, 0).width = Cm(6.57)
    tbl2.cell(0, 0).height = Cm(2)
    set_frame(tbl2)
    p = tbl2.cell(0, 0).paragraphs[0]
    run = p.add_run("正确填涂    ")
    run.add_picture(r'D:\校本资源相关\mark_pic\bmp小尺寸\black.png', width=Pt(14))
    run.font.bold = True
    run.font.size = Pt(9)
    run = p.add_run("    缺考标记    ")
    run.add_picture(r'D:\校本资源相关\mark_pic\bmp小尺寸\white.png', width=Pt(14))
    run.font.bold = True
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def prefix_admission_card(cell, stu_id_count):
    tbl = cell.add_table(rows=3, cols=stu_id_count)
    for i in range(stu_id_count):
        for cell in tbl.column_cells(i):
            cell.width = Cm(0.86)
    set_frame(tbl)
    p = tbl.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("准考证号")
    for i in range(stu_id_count - 1):
        tbl.cell(0, i).merge(tbl.cell(0, i + 1))
    for j in range(stu_id_count):
        p = tbl.cell(2, j).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for k in range(9):
            run = p.add_run("[{}]".format(k))
            run.font.size = Pt(9)
            run.add_break()
        run = p.add_run("[{}]".format(9))
        run.font.size = Pt(9)


def prefix_2dcode(cell, path):
    tbl2 = cell.add_table(rows=1, cols=2)
    # set_frame(tbl2)
    for cell in tbl2.column_cells(0):
        cell.width = Cm(4)
    for cell in tbl2.column_cells(1):
        cell.width = Cm(4)
    p = tbl2.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Pt(90))
    p = tbl2.cell(0, 1).paragraphs[0]
    # p.add_run().add_break()
    run = p.add_run("姓名：_________")
    run.font.bold = True
    run.add_break()
    run.add_break()
    run = p.add_run("班级：_________")
    run.font.bold = True
    run.add_break()
    run.add_break()
    run = p.add_run("考号：_________")
    run.font.bold = True
    run.add_break()
    run.add_break()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def prefix_2dcode_(cell, path):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    tbl2 = cell.add_table(rows=2, cols=1)
    set_frame(tbl2)
    tbl2.cell(0, 0).width = Cm(6)
    tbl2.cell(1, 0).width = Cm(6)
    tbl2.cell(0, 0).height = Cm(6)
    tbl2.cell(1, 0).height = Cm(2)
    p = tbl2.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Pt(90))
    p = tbl2.cell(1, 0).paragraphs[0]
    run = p.add_run("正确填涂    ")
    run.add_picture(r'D:\校本资源相关\mark_pic\bmp小尺寸\black.png', width=Pt(14))
    run.font.bold = True
    run.font.size = Pt(9)


def prefix_short_card(cell):
    cell.vertical = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("准考证号后三位数字")
    run.font.size = Pt(9)
    tbl = cell.add_table(rows=3, cols=3)
    set_frame(tbl)
    for cell in tbl.column_cells(0):
        cell.width = Cm(1)
    for cell in tbl.column_cells(1):
        cell.width = Cm(1)
    for cell in tbl.column_cells(2):
        cell.width = Cm(6.8)
    for cell in tbl.row_cells(0):
        cell.height = Cm(1)
    for cell in tbl.row_cells(1):
        cell.height = Cm(1)
    p = tbl.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("短")
    run.font.size = Pt(9)
    p = tbl.cell(1, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("考")
    run.font.size = Pt(9)
    p = tbl.cell(2, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("号")
    run.font.size = Pt(9)
    for i in range(2):
        tbl.cell(i, 0).merge(tbl.cell(i + 1, 0))
    for j in range(3):
        cell2 = tbl.cell(j, 2)
        cell2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for k in range(9):
            p = cell2.paragraphs[0]
            run = p.add_run("[{}]".format(k) + " ")
            run.font.size = Pt(9)
        run = p.add_run("[{}]".format(9))
        run.font.size = Pt(9)
    # tbl.cell(0, 2).merge(tbl.cell(1, 2)).merge(tbl.cell(2, 2))


def set_cell_background_color(cell):
    # if not isinstance(rgb_color, RGBValue):
    #     print('rgbColor is not RGBValue...', type(rgb_color))
    #     return
    # hr = str(hex(int(rgb_color.r)))[-2:]
    # hg = str(hex(int(rgb_color.g)))[-2:]
    # hb = str(hex(int(rgb_color.b)))[-2:]
    # color_str = hr + hg + hb
    # shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value=color_str))
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="CCCAC9"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading_elm_1)


def add_postfix_info(document, page_count):
    page_binary = get_binary_code(page_count)
    for i in range(page_count):
        section = document.add_section()
        footer = section.footer
        p = footer.paragraphs[0]
        p.add_run('{}'.format(page_binary[i]))


def get_binary_code(page_count):
    sum = []
    for i in range(1, page_count+1):
        binary = []
        m, n = divmod(i, 2)
        recurrent(m, n, binary)
        sum.append(binary)
    blank_count = max(max([len(ls) for ls in sum]), 3)
    for ls in sum:
        gap = blank_count - len(ls)
        for i in range(gap):
            ls.insert(0, 0)
    return sum


def recurrent(m, n, binary):
    if m == 0:
        binary.insert(0, n)
    else:
        binary.insert(0, n)
        m2, n2 = divmod(m, 2)
        recurrent(m2, n2, binary)


def write_forbidden(document):
    table = document.add_table(rows=2, cols=1)
    set_cell_border(
        table.cell(0, 0),
        top={"sz": 5, "val": "single", "color": "#000000"},
        left={"sz": 5, "val": "single", "color": "#000000"},
        # bottom={"sz": 5, "color": "#000000", "val": "single"},
        right={"sz": 5, "val": "single", "color": "#000000", },
    )
    set_cell_border(
        table.cell(1, 0),
        # top={"sz": 5, "val": "single", "color": "#000000"},
        left={"sz": 5, "val": "single", "color": "#000000"},
        bottom={"sz": 5, "color": "#000000", "val": "single"},
        right={"sz": 5, "val": "single", "color": "#000000", },
    )
    cell1 = table.cell(0, 0)
    set_cell_background_color(cell1)
    p = cell1.paragraphs[0]
    p.add_run().add_picture(r'D:\校本资源相关\答题卡\forbidden.png')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_break()
    cell2 = table.cell(1, 0)
    set_cell_background_color(cell2)
    p = cell2.paragraphs[0]
    run = p.add_run("请勿在此区域作答或者做任何标记")
    run.font.size = Pt(26)
    run.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def write4answersheet(answer_sheet_data):
    '''
    answer_sheet_data = {
        "name": paper_data["name"],
        "main_title": "{}-答题卡".format(paper_data["main_title"]),
        "testee": paper_data["testee"],
        "sheet": answer_sheet
    }
    answer_sheet[key].append({
                "q_no": ques_no,
                "pic_in_card": pic_in_card,
                "options": [v["key"] for v in options if v.get("key")],
                "subques_len": len(subquestions) if subquestions else 0
            })
    '''
    sheet = answer_sheet_data.get('sheet', {})
    sheet_name = answer_sheet_data.get('name')
    main_title = answer_sheet_data.get('main_title')
    scan_type = answer_sheet_data.get('scan_type')  # 默认是0：先扫后阅，另一值1：先阅后扫
    simi_score = answer_sheet_data.get('simi_score')  # 默认是0：不支持0.5分，另一值1：支持0.5分
    answer_path = r'C:\Users\j20687\Desktop\answersheet.docx'
    document = Document(answer_path)
    # 设置整个文档的默认字体
    microsoft_font = u'宋体'  # u 表示后面的字符串以 Unicode 格式进行编码
    black_font = u'黑体'
    number_font = 'Times New Roman'
    area = qn('w:eastAsia')
    document.styles['Normal'].font.name = microsoft_font
    document.styles['Normal'].font.size = Pt(10.5)
    document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
    document.styles['Normal']._element.rPr.rFonts.set(area, number_font)
    # 指定段落样式
    styles = document.styles
    # fls = time.time()
    # strr = 'lbj_Style%s' %fls  #自定义的样式的名称
    # strr = strr.replace('.', '')
    # strr = strr + ''.join(random.sample('zyxwvutsrqponmlkjihgfedcbaABCDEFGHIJKLMNOPQRST', 5))
    s = styles.add_style('lbj_Style', WD_STYLE_TYPE.PARAGRAPH)
    s.font.name = microsoft_font
    s.font.size = Pt(10.5)
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.line_spacing = Pt(0)  # 行距值
    s.paragraph_format.space_before = Pt(0)  # 段前距
    s.paragraph_format.space_after = Pt(0)  # 段后距
    s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE  # 行间距
    s._element.rPr.rFonts.set(area, number_font)  # 除中文外其它文字 使用的字体 ，备选项

    # 编辑页眉
    document.settings.odd_and_even_pages_header_footer = True
    section = document.sections[0]
    header = section.header
    # run = header.paragraphs[0]
    # run.style = s
    # p = run.add_run('2019学校12月月考卷试卷副标题答题卡2019学校12月月考卷试卷副标题答题卡2019学校12月月考卷试卷')
    # font = p.font
    # font.name = black_font
    # font.color.rgb = RGBColor(0, 0, 0)
    # font.size = Pt(15)
    # run.alignment = WD_ALIGN_PARAGRAPH.CENTER

    type_test_no, type_sheet = 1, 0
    subject_id = "2111612056"
    main_title = "2019学校12月月考卷试卷副标题答题卡"
    add_prefix_info(header, type_test_no, type_sheet, subject_id, main_title, stu_id_count=8)

    p = document.paragraphs[0]
    # p.add_run().add_break()
    run = p.add_run('选择题（请用2B铅笔填涂）')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = black_font
    # sub_mark_num_vec = []  # 每页竖长条表示块数量
    sub_mark_num = 0  # 竖长条表示块总数量
    sub_mark_choice = []  # 选择题黑块区域
    sub_mark_fillin = []  # 填空题黑块区域
    sub_mark_eassy = []  # 简答题黑块区域
    qus_range_choice = []  # 选择题题号范围
    qus_range_fillin = []  # 填空题题号范围
    qus_range_eassy = []  # 简答题题号范围
    paper_num = 0  # 页数
    ques_no = 0  # 题号
    choice_answer_cnt = []  # 每个区块的选项数
    choice_line_number = []  # 每大行的最多题目数量
    choice_answer_pattern = []  # 每行的每列选择题数目
    choice_col_num = []  # 选择题每行的列数
    choice_row_num = []  # 选择题行数
    option_len = []
    choice_num, mm, combine_ques_index = [], [], []
    count_ques = 0
    for key, value in sheet.items():
        if key == 'choice_question_all':
            for k in range(len(value)):
                q_no = value[k].get("q_no")
                opt_len = len(value[k].get("options")) if len(value[k].get("options")) >= 4 else 4
                mm.append(q_no)
                if not isinstance(q_no, list):
                    option_len.append(opt_len)
                else:
                    for qq in q_no:
                        option_len.append(opt_len)
    for i in range(len(mm)):
        if isinstance(mm[i], list):
            count_ques += len(mm[i])
            combine_ques_index.append(i)
            for j in mm[i]:
                choice_num.append(j)
        else:
            count_ques += 1
            choice_num.append(mm[i])
    # logger.info("option_len========================={}".format(option_len))
    print("option_len========================={}".format(option_len))
    choice_group = sheet.get('choice_question', [])
    choice_len = len(choice_group)
    choice_len_sum = len(choice_num)  # 选择题总数，包括组合题中的小选择题
    fillin_group = sheet.get('fillin_question', [])
    fillin_len = len(fillin_group)
    fs_num = fillin_group[0].get('q_no') if fillin_group else 0  # 填空题开始的题号
    fs_num_list = [ques.get('q_no') for ques in fillin_group]
    subjective_question = sheet.get('subjective_question', [])
    subjec_len = len(subjective_question)
    ss_num = subjective_question[0].get('q_no') if subjective_question else 0  # 主观题开始的题号
    ss_num_list = [ques.get('q_no') for ques in subjective_question]
    ss_score_list = [int(ques.get('score')) for ques in subjective_question]  # 各主观题的分值
    combine_question = sheet.get('combine_question', [])

    pic_in_sub = [[] for x in range(subjec_len)]
    pic_num = 0
    for ques in subjective_question:
        pic_in_card = ques.get('pic_in_card', "")
        pic_in_sub[pic_num] = pic_in_card if pic_in_card else []
        pic_num += 1
    tmp = []
    for i in range(len(mm)):
        if isinstance(mm[i], list):
            if tmp:
                qus_range_choice.append(tmp)
                tmp = []
            qus_range_choice.append(mm[i])
        else:
            tmp.append(mm[i])

    if not combine_question:
        qus_range_choice.append(tmp)
    else:
        if not isinstance(mm[-1], list):
            qus_range_choice.append(tmp)

    if fillin_group:
        qus_range_fillin = fs_num_list
    if subjective_question:
        qus_range_eassy = ss_num_list

    # 选项个数不一样
    group_choice_num_ = []
    group_choice_num = {}  # 按题号顺序各选项数对应的题号集合
    is_multi = False
    if option_len:
        pre = option_len[0]
        group_choice_num = {pre: []}
        group_choice_num[pre].append(choice_num[0])
    for opt_i in range(1, len(option_len)):
        now = option_len[opt_i]
        pre = option_len[opt_i - 1]
        if pre != now:
            is_multi = True
            group_choice_num_.append(group_choice_num)
            group_choice_num = {now: []}
            # group_choice_num[now].append(choice_num[opt_i])
        group_choice_num[now].append(choice_num[opt_i])
    # logger.info("group_choice_num========================={}".format(group_choice_num))
    # if is_multi:
    group_choice_num_.append(group_choice_num)
    # logger.info("group_choice_num_========================={}".format(group_choice_num_))
    cols = 0
    for item in group_choice_num_:
        for key, value in item.items():
            d1, d2 = divmod(len(value), 5)
            for i in range(d1):
                cols += 1
                choice_answer_pattern.append(5)
                choice_answer_cnt.append(key)
            if d2:
                cols += 1
                choice_answer_pattern.append(d2)
                choice_answer_cnt.append(key)
    # logger.info("cols========================={}".format(cols))
    c1, c2 = divmod(cols, 3)
    row = c1 + 1 if c2 else c1
    choice_row_num.append(row)
    for i in range(c1):
        choice_col_num.append(3)
    if c2:
        choice_col_num.append(c2)
    tmp4line = []
    plusn = 0
    for j in range(len(choice_answer_pattern)):
        tmp4line.append(choice_answer_pattern[j])
        plusn += 1
        if plusn == 3:
            choice_line_number.append(max(tmp4line))
            tmp4line = []
            plusn = 0
    if tmp4line:
        choice_line_number.append(max(tmp4line))
    break_sum = 0  # 一页最多存放（5+5+5+5+1）* 3个选项
    break_row = 0
    break_num = 0
    for cl in range(len(choice_line_number)):
        break_sum += choice_line_number[cl]
        if break_sum > 21 and cl > 1:
            break_row = cl
            break_sum -= choice_line_number[cl]
            break
    for br in range(break_row * 3):
        break_num += choice_answer_pattern[br] if len(choice_answer_pattern) >= br else 0

    choice_answer_pattern = [choice_answer_pattern] if not break_num else [choice_answer_pattern[:break_row * 3],
                                                                           choice_answer_pattern[break_row * 3:]]
    choice_answer_cnt = [choice_answer_cnt] if not break_num else [choice_answer_cnt[:break_row * 3],
                                                                   choice_answer_cnt[break_row * 3:]]
    choice_row_num = choice_row_num if not break_num else [break_row, row - break_row]
    choice_col_num = [choice_col_num] if not break_num else [choice_col_num[0:break_row], choice_col_num[break_row:]]
    choice_line_number = [choice_line_number] if not break_num else [choice_line_number[0:break_row],
                                                                     choice_line_number[break_row:]]

    # logger.info("choice_line_number========================={}".format(choice_line_number))
    # logger.info("row========================={}".format(row))
    # logger.info("choice_answer_pattern========================={}".format(choice_answer_pattern))
    if not break_num:
        write_option(choice_num, option_len, choice_answer_pattern[0], row, document)
    else:
        write_option(choice_num[0:break_num], option_len[0:break_num], choice_answer_pattern[0], break_row, document)
        if break_num != 21:
            document.add_page_break()
        write_option(choice_num[break_num:], option_len[break_num:], choice_answer_pattern[1], row - break_row,
                     document)
    # write_option(ques_no, choice_len_sum, option_len, choice_ques_no, document)
    writeline(document, black_font)
    sub_fillin(fillin_len, fs_num_list, document, scan_type)

    if ss_num:
        for ns in range(len(ss_num_list)):
            number = ss_num_list[ns]
            score = ss_score_list[ns] if scan_type else None
            # score = ss_score_list[ns]
            sub_subjective(number, document, pic_in_sub, score, simi_score)

    block_num = 0
    start = 0
    step = 3
    if choice_len_sum:
        for n in range(len(choice_row_num)):
            block_num += 1
            sub_mark_choice.append(start + step * (block_num - 1) - (block_num - 1))
            sub_mark_choice.append(start + step * block_num - (block_num - 1))
    if fillin_len:
        block_num += 1
        sub_mark_fillin.append(start + step * (block_num - 1) - (block_num - 1))
        sub_mark_fillin.append(start + step * block_num - (block_num - 1))
        if not subjec_len:
            index = sub_mark_fillin[-2]
            sub_mark_fillin.pop(-1)
            sub_mark_fillin.append(index + 1)
        else:
            for n in range(subjec_len):
                block_num += 1
                sub_mark_eassy.append(start + step * (block_num - 1) - (block_num - 1))
                sub_mark_eassy.append(start + step * block_num - (block_num - 1))
            index = sub_mark_eassy[-2]
            sub_mark_eassy.pop(-1)
            sub_mark_eassy.append(index + 1)
    else:
        if subjec_len:
            for n in range(subjec_len):
                block_num += 1
                sub_mark_eassy.append(start + step * (block_num - 1) - (block_num - 1))
                sub_mark_eassy.append(start + step * block_num - (block_num - 1))
            index = sub_mark_eassy[-2]
            sub_mark_eassy.pop(-1)
            sub_mark_eassy.append(index + 1)
        else:
            if sub_mark_choice:
                index = sub_mark_choice[-2]
                sub_mark_choice.pop(-1)
                sub_mark_choice.append(index + 1)

    #  return sub_mark_num
    for i in range(len(choice_row_num)):
        sub_mark_num += 2
    # if choice_len_sum:
    #     sub_mark_num += 4
    if fillin_len:
        sub_mark_num += 2
    for num in range(subjec_len):
        sub_mark_num += 2

    t1 = int(round(time.time() * 1000))
    doc_file_name = "{}-答题卡".format(main_title)
    return_path = r'C:\Users\j20687\Desktop\%s.docx' % (doc_file_name)
    document.save(return_path)
    sub_json_data = {}
    tmp_sub_mark_choice = []
    if choice_num:
        for i in range(len(choice_row_num)):
            tmp_sub_mark_choice.append([sub_mark_choice[i * 2], sub_mark_choice[i * 2 + 1]])
    # logger.info("sub_mark_choice========================={}".format(sub_mark_choice))
    sub_json_data['sub_mark_choice'] = tmp_sub_mark_choice
    sub_json_data['sub_mark_fillin'] = sub_mark_fillin
    sub_json_data['sub_mark_eassy'] = sub_mark_eassy
    sub_json_data['sub_mark_num'] = sub_mark_num
    # sub_json_data['sub_mark_num_vec'] = sub_mark_num_vec
    # page = get_pdf_page(return_path)
    # logger.info("page========================={}".format(page))
    sub_json_data['page_num'] = paper_num
    sub_json_data['scan_type'] = scan_type
    sub_json_data['simi_score'] = simi_score
    # sub_json_data['paper_num_count_flag'] = paper_num_count_flag
    sub_json_data['choice_answer_cnt'] = choice_answer_cnt
    sub_json_data['choice_line_number'] = choice_line_number
    sub_json_data['choice_answer_pattern'] = choice_answer_pattern
    sub_json_data['choice_col_num'] = choice_col_num
    sub_json_data['choice_row_num'] = choice_row_num
    sub_json_data['qus_range_choice'] = qus_range_choice
    sub_json_data['qus_range_fillin'] = qus_range_fillin
    sub_json_data['qus_range_eassy'] = qus_range_eassy
    sub_json_data['eassy_score'] = ss_score_list
    # logger.info("sub_json_data========================={}".format(sub_json_data))
    return return_path, sub_json_data


def get4answer_sheet_data(paper_data):
    answer_sheet = {}
    ques_no = 1
    for g_idx, grp in enumerate(paper_data.get("question_group", [])):
        for eq in grp.get("exam_questions", []):
            question = eq.get("question", {})
            score = eq.get("score") or 0
            q_type_name = question.get("q_type", {}).get("name")
            pic_in_card = question.get("pic_in_card", "")
            options = question.get("options") or []
            subquestions = question.get("subs") or []
            if q_type_name in CHOICE_QUESTIONS:
                key = "choice_question"
            elif subquestions:
                key = "combine_question"
            elif q_type_name in FILL_IN_QUESTIONS:
                key = "fillin_question"
            else:
                key = "subjective_question"
            if not answer_sheet.get(key):
                answer_sheet[key] = []
            answer_sheet[key].append({
                "q_no": ques_no,
                "score": score,
                "pic_in_card": pic_in_card,
                "options": [v["key"] for v in options if v.get("key")],
                "subques_len": len(subquestions) if subquestions else 0
            })
            if q_type_name in CHOICE_QUESTIONS or subquestions:
                key = "choice_question_all"
                if not answer_sheet.get(key):
                    answer_sheet[key] = []
                answer_sheet[key].append({
                    "q_no": ques_no,
                    "pic_in_card": pic_in_card,
                    "options": [v["key"] for v in options if v.get("key")],
                    "subques_len": len(subquestions) if subquestions else 0
                })
            if key == "combine_question":
                ques_no += len(subquestions)
            else:
                ques_no += 1
    answer_sheet_data = {
        "name": paper_data.get("name") or paper_data.get("main_title"),
        "main_title": "{}-答题卡".format(paper_data.get("main_title")),
        "testee": paper_data.get("testee"),
        "scan_type": paper_data.get("scan_type") or 0,
        "simi_score": paper_data.get("simi_score") or 0,
        "sheet": answer_sheet
    }
    answer_path, sub_json_data = write4answersheet(answer_sheet_data)
    return answer_path, answer_sheet_data, sub_json_data


def get_pdf_page(docx_path):
    t1 = time.time()
    file_name = docx_path.split("/")[-1]
    file_name = ".".join(file_name.split(".")[:-1])
    pdf_path = '/tmp/{}.pdf'.format(file_name)
    os.system('soffice --headless --invisible --convert-to pdf %s --outdir /tmp' % docx_path)
    try:
        f = pdfplumber.open(pdf_path)
        page = len(f.pages)
    except PDFSyntaxError:
        page = 0
    t2 = time.time()
    logger.info("get_pdf_page costs========================={}s".format(t2 - t1))
    return page


if __name__ == '__main__':
    path = r'C:\Users\j20687\Desktop\paperJson.json'
    with open(path, 'r', encoding='utf-8') as f:
        paper_data = json.load(f)
        paper_data = paper_data.get("data")
        print(paper_data)
        get4answer_sheet_data(paper_data)
