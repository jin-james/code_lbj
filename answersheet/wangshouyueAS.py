import json
import logging
import os
import re
import time
import uuid

import pdfplumber
import requests
from MyQR import myqr
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.oxml.xmlchemy import OxmlElement
from docx.shared import Pt, RGBColor, Cm
from pdfminer.pdfparser import PDFSyntaxError

# from nile.utils.constants import CHOICE_QUESTIONS, FILL_IN_QUESTIONS

logger = logging.getLogger(__name__)
imgtmp_base_path = '/tmp/img'
CHOICE_QUESTIONS = ["选择题", "单选题", "单项选择题", "多选题", "多项选择题", "听力选择题", '双选题']  # 选择题
FILL_IN_QUESTIONS = ["填空题", "单词拼写题"]  # 填空题


def ichoice_sub_write(table, n, ques_no, pic_in_sub, isfillin=False, is_score=0):
    if isfillin:
        cell = table.cell(0, 1)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cell.paragraphs[0]
        paragraph_format = p.paragraph_format
        paragraph_format.line_spacing = Pt(40)  # 22磅行间距
        ques_count = 0
        # p.add_run().add_break()
        for num in ques_no:
            ques_count += 1
            p.add_run(str(num) + '题、_____________________________').font.size = Pt(9)
            if is_score:
                p.add_run().add_picture(r'D:\Program Files\code_lbj\answersheet\fillin_scorebar.png', width=Pt(30))
                p.add_run(" ")
            else:
                p.add_run("      ")
            if ques_count % 2 == 0:
                p.add_run().add_break()
    else:
        cell = table.cell(1, 1)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cell.paragraphs[0]
        if not pic_in_sub:
            p.add_run().add_break()
        p.add_run(str(ques_no) + '题、').font.size = Pt(9)
    if pic_in_sub:
        add_pic(table, pic_in_sub)
        n -= 3
    for i in range(n):
        p.add_run().add_break()


def set_col_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def set_score_column_widths_(table, value):
    for row in table.rows:
        for cell in row.cells:
            cell.width = value


def hide_frame(table, rows, cols, ischoice=True):
    if ischoice:
        for r in table.rows:
            for c in r.cells:
                set_cell_border(
                    c,
                    top={"sz": 0, "val": "single", "color": "#FFFFFF"},
                    left={"sz": 0, "val": "single", "color": "#FFFFFF"},
                    bottom={"sz": 0, "color": "#FFFFFF", "val": "single"},
                    right={"sz": 0, "val": "single", "color": "#FFFFFF", },
                )
        set_cell_border(
            table.cell(0, 1),
            left={"sz": 5, "val": "single", "color": "#000000"},
            top={"sz": 5, "color": "#000000", "val": "single"},
        )
        set_cell_border(
            table.cell(0, 2),
            top={"sz": 5, "val": "single", "color": "#000000"},
        )
        set_cell_border(
            table.cell(0, 3),
            top={"sz": 5, "color": "#000000", "val": "single"},
            right={"sz": 5, "val": "single", "color": "#000000", },
        )
        set_cell_border(
            table.cell(rows - 1, 1),
            left={"sz": 5, "val": "single", "color": "#000000"},
            bottom={"sz": 5, "color": "#000000", "val": "single"},
        )
        set_cell_border(
            table.cell(rows - 1, 2),
            bottom={"sz": 5, "val": "single", "color": "#000000"},
        )
        set_cell_border(
            table.cell(rows - 1, 3),
            bottom={"sz": 5, "color": "#000000", "val": "single"},
            right={"sz": 5, "val": "single", "color": "#000000", },
        )
        for r in range(1, rows - 1):
            set_cell_border(
                table.cell(r, 1),
                left={"sz": 5, "val": "single", "color": "#000000", },
            )
            set_cell_border(
                table.cell(r, 3),
                right={"sz": 5, "val": "single", "color": "#000000", },
            )
    else:
        for r in range(rows):
            set_cell_border(
                table.cell(r, 0),
                top={"sz": 0, "val": "single", "color": "#FFFFFF"},
                bottom={"sz": 0, "color": "#FFFFFF", "val": "single"},
                left={"sz": 0, "val": "single", "color": "#FFFFFF"},
            )
            set_cell_border(
                table.cell(r, cols - 1),
                top={"sz": 0, "val": "single", "color": "#FFFFFF"},
                bottom={"sz": 0, "color": "#FFFFFF", "val": "single"},
                right={"sz": 0, "val": "single", "color": "#FFFFFF", },
            )


def set_frame(table):
    for r in table.rows:
        for c in r.cells:
            set_cell_border(
                c,
                top={"sz": 5, "val": "single", "color": "#000000"},
                left={"sz": 5, "val": "single", "color": "#000000"},
                bottom={"sz": 5, "color": "#000000", "val": "single"},
                right={"sz": 5, "val": "single", "color": "#000000", },
            )


def hide_frame_single(table):
    for r in table.rows:
        for c in r.cells:
            set_cell_border(
                c,
                top={"sz": 0, "val": "single", "color": "#FFFFFF"},
                left={"sz": 0, "val": "single", "color": "#FFFFFF"},
                bottom={"sz": 0, "color": "#FFFFFF", "val": "single"},
                right={"sz": 0, "val": "single", "color": "#FFFFFF", },
            )


def add_options(p, option_no):
    p.add_run().add_picture(r'D:\Program Files\code_lbj\answersheet\sel-a_.bmp', width=Pt(12.5))
    p.add_run(' ').add_picture(r'D:\Program Files\code_lbj\answersheet\sel-b_.bmp', width=Pt(12.5))
    p.add_run(' ').add_picture(r'D:\Program Files\code_lbj\answersheet\sel-c_.bmp', width=Pt(12.5))
    p.add_run(' ').add_picture(r'D:\Program Files\code_lbj\answersheet\sel-d_.bmp', width=Pt(12.5))
    if int(option_no) == 5:
        p.add_run(' ').add_picture(r'D:\Program Files\code_lbj\answersheet\sel-e_.bmp', width=Pt(12.5))
    if int(option_no) == 6:
        p.add_run(' ').add_picture(r'D:\Program Files\code_lbj\answersheet\sel-e_.bmp', width=Pt(12.5))
        p.add_run(' ').add_picture(r'D:\Program Files\code_lbj\answersheet\sel-f_.bmp', width=Pt(12.5))
    if int(option_no) == 7:
        p.add_run(' ').add_picture(r'D:\Program Files\code_lbj\answersheet\sel-e_.bmp', width=Pt(12.5))
        p.add_run(' ').add_picture(r'D:\Program Files\code_lbj\answersheet\sel-f_.bmp', width=Pt(12.5))
        p.add_run(' ').add_picture(r'D:\Program Files\code_lbj\answersheet\sel-g_.bmp', width=Pt(12.5))


def add_mark(cols, table, rows=None):
    cell1 = table.cell(0, 0)
    cell1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p1 = cell1.paragraphs[0]
    p1.add_run().add_picture(r'D:\Program Files\code_lbj\answersheet\Rectangle-c.png', width=Pt(6))
    p1.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell2 = table.cell(0, cols + 1)
    cell2.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p2 = cell2.paragraphs[0]
    p2.add_run().add_picture(r'D:\Program Files\code_lbj\answersheet\Rectangle-c.png', width=Pt(6))
    p2.alignment = WD_TABLE_ALIGNMENT.RIGHT
    if rows:
        cell1 = table.cell(rows - 1, 0)
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p1 = cell1.paragraphs[0]
        p1.add_run().add_picture(r'D:\Program Files\code_lbj\answersheet\Rectangle-c.png', width=Pt(6))
        p1.alignment = WD_TABLE_ALIGNMENT.LEFT
        cell2 = table.cell(rows - 1, cols + 1)
        cell2.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p2 = cell2.paragraphs[0]
        p2.add_run().add_picture(r'D:\Program Files\code_lbj\answersheet\Rectangle-c.png', width=Pt(6))
        p2.alignment = WD_TABLE_ALIGNMENT.RIGHT


def add_pic(table_sub, pic_in_sub):
    # table_sub.alignment = WD_TABLE_ALIGNMENT.CENTER
    para = table_sub.cell(0, 1).paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    re_image = re.compile("<img.*?>", re.S | re.I)
    re_src = re.compile("src=\"(.*?)\"", re.S | re.I)
    re_height = re.compile("height=\"(.*?)\"", re.S | re.I)
    line = [v for v in pic_in_sub if v]
    line = line[0] if line else ""
    image = re_image.findall(line)
    for img in image:
        string = re_src.search(img)
        string = string[1] if string else ""
        hei_str = re_height.search(img)
        hei_str = hei_str[1] if hei_str else ""
        content = requests.get(string).content
        subfix = string[-32:]
        path = imgtmp_base_path + '/' + subfix + '{}'.format('.png')
        height_v = float(hei_str.split('p')[0])
        if os.path.exists(path):
            para.add_run().add_picture(path, height=Pt(height_v * 0.7))
        with open(path, 'wb')as f:
            f.write(content)
        para.add_run().add_picture(path, height=Pt(height_v * 0.7))
    for i in range(len(pic_in_sub)):
        tmp = pic_in_sub[i]
        if not tmp:
            pic_in_sub[i] = []  # 剔除已经添加的图片数据
            break


def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def write_option(choice_num, choice_opt, choice_answer_pattern, rows, document):
    width5 = (Cm(0.59), Cm(6.12), Cm(6.12), Cm(6.12), Cm(0.59))
    choice_len = len(choice_num)
    if choice_len:
        cols = 3
        table_choice = document.add_table(rows=rows, cols=cols + 2, style='Table Grid')
        set_col_widths(table_choice, width5)
        hide_frame(table_choice, rows, cols + 2)
        c = 1
        r = 0
        p = table_choice.cell(r, c).paragraphs[0]
        p.add_run().add_break()
        start = 0
        for count in choice_answer_pattern:
            ques_no = 0
            num = choice_num[start: start + count]
            option = choice_opt[start: start + count]
            choice_length = len(num)
            for i in range(choice_length):
                ques_no += 1
                choice_no = num[i]
                option_no = option[i]
                if choice_no >= 100:
                    p.add_run(str(choice_no) + " ").font.size = Pt(9)
                elif choice_no >= 10:
                    p.add_run(' ' + str(choice_no) + ' ').font.size = Pt(9)
                else:
                    p.add_run("  " + str(choice_no) + " ").font.size = Pt(9)
                add_options(p, option_no)
                if ques_no != choice_length:
                    p.add_run().add_break()
            if c < 3:
                c += 1
                p.add_run().add_break()  # 另起一列的时候在前一列末加一空行
                p = table_choice.cell(r, c).paragraphs[0]
                p.add_run().add_break()  # 另起一列的时候在前加一空行
            else:
                if r < rows - 1:
                    r += 1
                    c = 1
                    p.add_run().add_break()  # 另起一行的时候在前一单元格末加一空行
                    p = table_choice.cell(r, c).paragraphs[0]
                    p.add_run().add_break()  # 另起一行的时候在前加一空行
            start += count
        add_mark(cols, table_choice)


def write_option_vertical(choice_num, choice_opt, document):
    '''
    :param choice_num: 列表，一行中各块的题号[[], [], []]
    :param choice_opt: 列表，一行中各块的每题选项数目,同上
    :param document:
    :return:
    '''
    choice_len = len(choice_num)
    width5 = (Cm(0.59), Cm(6.12), Cm(6.12), Cm(6.12), Cm(0.59))
    if choice_len:
        rows = 1
        cols = len(choice_num)
        table_choice = document.add_table(rows=rows, cols=cols + 2, style='Table Grid')
        set_col_widths(table_choice, width5)
        hide_frame(table_choice, rows, cols + 2)
        c = 1
        r = 0
        p = table_choice.cell(r, c).paragraphs[0]
        p.add_run().add_break()
        for i in range(len(choice_num)):
            for j in range(len(choice_num[i])):
                num = choice_num[i][j]
                opt = choice_opt[i][j]
                if num >= 100:
                    p.add_run(str(num) + " ").font.size = Pt(9)
                elif num >= 10:
                    p.add_run(' ' + str(num) + ' ').font.size = Pt(9)
                else:
                    p.add_run("  " + str(num) + " ").font.size = Pt(9)
                add_options(p, opt)
            c += 1
        add_mark(cols, table_choice)


def write_option_rank(choice_num, choice_opt, document):
    '''
    只支持四个选项的横排，一排五列（并且此时不支持设置每列行数，默认为5行）
    :param choice_num: 列表，一行中各块的题号[[], [], []]
    :param choice_opt: 列表，一行中各块的每题选项数目,同上
    :param document:
    :return:
    '''
    f1 = 18.36
    cols = len(choice_num)
    width = [Cm(0.59)]
    f_per = round(f1/cols, 2)
    for i in range(cols):
        width.append(Cm(f_per))
    width.append(Cm(0.59))
    rows = 1
    table_choice = document.add_table(rows=rows, cols=cols + 2, style='Table Grid')
    set_col_widths(table_choice, width)
    hide_frame(table_choice, rows, cols + 2)
    c = 1
    r = 0
    p = table_choice.cell(r, c).paragraphs[0]
    p.add_run().add_break()
    for i in range(len(choice_num)):
        for j in range(len(choice_num[i])):
            num = choice_num[i][j]
            opt = choice_opt[i][j]
            if num >= 100:
                p.add_run(str(num) + " ").font.size = Pt(9)
            elif num >= 10:
                p.add_run(' ' + str(num) + ' ').font.size = Pt(9)
            else:
                p.add_run("  " + str(num) + " ").font.size = Pt(9)
            add_options(p, opt)
        c += 1
    add_mark(cols, table_choice)


def sub_subjective(number, document, pic_in_sub, score, simi_score, cols_count):
    score = 20
    simi_score = 1
    width3 = (Cm(0.59), Cm(18.36), Cm(0.59))
    if cols_count == 3:
        width3 = (Cm(0.59), Cm(11.82), Cm(0.59))
    rows = 2
    cols = 1
    table_sub = document.add_table(rows=rows, cols=cols + 2, style='Table Grid')
    set_col_widths(table_sub, width3)
    if score is not None:
        add_mark(cols, table_sub)
        hide_frame(table_sub, rows, cols + 2, ischoice=False)
        cell0 = table_sub.cell(0, 1)
        cell0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if cols_count == 3:
            sub_cols = 16
            sub_tbl = cell0.add_table(rows=1, cols=sub_cols)
            set_score_column_widths_(sub_tbl, Cm(0.71))
        else:
            sub_cols = 22
            sub_tbl = cell0.add_table(rows=1, cols=sub_cols)
            set_score_column_widths_(sub_tbl, Cm(0.82))
        if simi_score:
            par = sub_tbl.cell(0, sub_cols - 1).paragraphs[0]
            pf = par.add_run("{}".format(0.5))
            font = pf.font
            font.size = Pt(7.5)
        if cols_count != 3:
            if 0 < score <= 20:
                index = 20
                for i in range(score + 1):
                    par = sub_tbl.cell(0, index).paragraphs[0]
                    pf = par.add_run("{}".format(i))
                    font = pf.font
                    font.size = Pt(9)
                    index -= 1
            else:
                t, s = divmod(score, 10)
                index = 0
                for i in range(sub_cols - 2, 10, -1):
                    par = sub_tbl.cell(0, i).paragraphs[0]
                    pf = par.add_run("{}".format(index))
                    font = pf.font
                    font.size = Pt(9)
                    index += 1
                index = sub_cols - 11 - 1 - 1
                for j in range(int(t + 1)):
                    par = sub_tbl.cell(0, index).paragraphs[0]
                    pf = par.add_run("{}".format(j))
                    font = pf.font
                    font.size = Pt(9)
                    index -= 1
        else:
            t, s = divmod(score, 10)
            index = 0
            for i in range(sub_cols - 2, sub_cols - 12, -1):
                par = sub_tbl.cell(0, i).paragraphs[0]
                pf = par.add_run("{}".format(index))
                font = pf.font
                font.size = Pt(9)
                index += 1
            index = sub_cols - 11 - 1 - 1
            for j in range(int(t + 1)):
                par = sub_tbl.cell(0, index).paragraphs[0]
                pf = par.add_run("{}".format(j))
                font = pf.font
                font.size = Pt(9)
                index -= 1
        for r in sub_tbl.rows:
            for c in r.cells:
                set_cell_border(
                    c,
                    top={"sz": 5, "val": "single", "color": "#000000"},
                    left={"sz": 5, "val": "single", "color": "#000000"},
                    bottom={"sz": 5, "color": "#000000", "val": "single"},
                    right={"sz": 5, "val": "single", "color": "#000000", },
                )
        set_cell_border(
            table_sub.cell(0, 1),
            left={"sz": 0, "val": "single", "color": "#FFFFFF"},
            right={"sz": 0, "val": "single", "color": "#FFFFFF", },
            top={"sz": 0, "val": "single", "color": "#FFFFFF", },
            bottom={"sz": 0, "val": "single", "color": "#FFFFFF", },
        )
        set_cell_border(
            table_sub.cell(0, 0),
            right={"sz": 0, "val": "single", "color": "#FFFFFF", },
        )
        set_cell_border(
            table_sub.cell(0, 2),
            left={"sz": 0, "val": "single", "color": "#FFFFFF"},
        )
        table_sub.cell(1, 0).merge(table_sub.cell(0, 0))
        table_sub.cell(1, 2).merge(table_sub.cell(0, 2))

    else:
        table_sub.cell(1, 1).merge(table_sub.cell(0, 1))
        add_mark(cols, table_sub)
        hide_frame(table_sub, rows, cols + 2, ischoice=False)
    ichoice_sub_write(table_sub, 18, number, pic_in_sub, is_score=score)


def sub_fillin(fillin_len, fs_num_list, document, scan_type, cols_count):
    width3 = (Cm(0.59), Cm(18.36), Cm(0.59))
    if cols_count == 3:
        width3 = (Cm(0.59), Cm(11.82), Cm(0.59))
    if fillin_len:
        rows = 1
        cols = 1
        table_fillin = document.add_table(rows=rows, cols=cols + 2, style='Table Grid')
        set_col_widths(table_fillin, width3)
        hide_frame(table_fillin, rows, cols + 2, ischoice=False)
        ichoice_sub_write(table_fillin, 0, fs_num_list, [], isfillin=True, is_score=scan_type)
        add_mark(cols, table_fillin)


def writeline(document, black_font):
    p = document.add_paragraph()
    # p.add_run().add_break()
    run = p.add_run('非选择题（请在各试题的答题区内作答）')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = black_font


def get_ques_range(choice_num, qus_range_choice):
    sort_index_ = {}
    qus_range_choice_ = []
    for item in qus_range_choice:
        if item:
            start = item[0]
            sort_index_[start] = item
    for i in choice_num:
        for key, value in sort_index_.items():
            if i == key:
                qus_range_choice_.append(value)
    return qus_range_choice_


def add_prefix_info(document, main_title, type_test_no, type_sheet, subject_id, cols_count, absent, paper_kind,
                    stu_id_count=8):
    '''
    :param document:
    :param main_title:标题
    :param type_test_no: 0:条形码，1：准考证，2：学籍号，3：短考号
    :param absent: 0:不添加，1：添加
    :param paper_kind: 0:不添加A/B卷，1：添加A/B卷
    :param main_title: 标题
    :param type_sheet: 0:网阅，1：手阅
    :param stu_id_count: 学号长度，默认为8位
    :param subject_id: 学科识别号
    :return:
    '''
    prefix_title(document, main_title, type_test_no, paper_kind)
    if type_sheet == 0:
        if type_test_no == 0 or type_test_no == 1:
            table_pre = document.add_table(rows=1, cols=2)
            table_pre.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell10 = table_pre.cell(0, 0)
            p = cell10.paragraphs[0]
            run = p.add_run("姓名：____________  班级：____________")
            run.font.bold = True
            run.font.name = u'黑体'
            p.add_run().add_break()
            run = p.add_run("考场/座位号：____________")
            run.font.bold = True
            run.font.name = u'黑体'
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run.add_break()
            if type_test_no == 0:
                prefix_notice_1(cell10)
                cell11 = table_pre.cell(0, 1)
                prefix_codebar(cell11)
                prefix_absent(cell11, paper_kind)
                p = cell10.paragraphs[-1]
                p.add_run().add_break()
            else:
                prefix_notice_2(cell10)
                prefix_absent(cell10, paper_kind)
                cell11 = table_pre.cell(0, 1)
                prefix_admission_card(cell11, type_test_no)
                p = cell11.paragraphs[-1]
                p.add_run().add_break()
        else:
            table_pre = document.add_table(rows=2, cols=2)
            table_pre.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell00 = table_pre.cell(0, 0)
            p = cell00.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if cols_count == 3:
                p.add_run("姓名：____________ 班级：____________ 考场/座位号：____________ ").font.bold = True
            else:
                p.add_run("姓名：______________________  班级：______________________  "
                          "考场/座位号：______________________").font.bold = True
            table_pre.cell(0, 0).merge(table_pre.cell(0, 1))
            cell10 = table_pre.cell(1, 0)
            if cols_count != 3:
                prefix_notice_3(cell10)
                prefix_absent(cell10)
            else:
                prefix_absent_cols3(cell10)
            cell11 = table_pre.cell(1, 1)
            prefix_admission_card(cell11, type_test_no, stu_id_count=stu_id_count)
            p = cell11.paragraphs[-1]
            p.add_run().add_break()
            if cols_count == 3:
                prefix_notice_3_cols3(document)

    else:
        path = get_2dcode(subject_id)
        if type_test_no == 0 or type_test_no == 3:
            table_pre = document.add_table(rows=1, cols=2)
            table_pre.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell10 = table_pre.cell(0, 0)
            prefix_2dcode(cell10, path, type_test_no, cols_count, absent)
            if type_test_no == 0 and cols_count == 3:
                p = document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if cols_count == 3:
                    p.add_run("姓名：____________ 班级：____________ 考场/座位号：____________ ").font.bold = True
                else:
                    p.add_run("姓名：______________________  班级：______________________  "
                              "考场/座位号：______________________").font.bold = True
                p.add_run().add_break()
            cell11 = table_pre.cell(0, 1)
            if type_test_no == 0:
                prefix_codebar(cell11)
            else:
                prefix_short_card(cell11)
        if type_test_no == 1:
            table_pre = document.add_table(rows=1, cols=2)
            table_pre.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell10 = table_pre.cell(0, 0)
            prefix_2dcode(cell10, path, type_test_no, cols_count, absent)
            cell01 = table_pre.cell(0, 1)
            prefix_admission_card(cell01, type_test_no)
            p = cell01.paragraphs[-1]
            p.add_run().add_break()
        if type_test_no == 2:
            table_pre = document.add_table(rows=2, cols=2)
            table_pre.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell10 = table_pre.cell(0, 0)
            p = cell10.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if cols_count == 3:
                p.add_run("姓名：____________ 班级：____________ 考场/座位号：____________ ").font.bold = True
            else:
                p.add_run("姓名：______________________  班级：______________________  "
                          "考场/座位号：______________________").font.bold = True
            table_pre.cell(0, 0).merge(table_pre.cell(0, 1))
            cell20 = table_pre.cell(1, 0)
            if cols_count != 3:
                prefix_2dcode_(cell20, path, absent)
            else:
                prefix_2dcode_cols3(cell20, path, absent)
            cell21 = table_pre.cell(1, 1)
            prefix_admission_card(cell21, type_test_no, stu_id_count=stu_id_count)
            p = cell21.paragraphs[-1]
            p.add_run().add_break()
        if path:
            os.remove(path)


def prefix_title(document, main_title, type_test_no, paper_kind):
    if not (type_test_no == 2 and paper_kind):
        table_pre = document.add_table(rows=1, cols=1)
        cell00 = table_pre.cell(0, 0)
        p = cell00.paragraphs[0]
        run = p.add_run(main_title)
        run.font.size = Pt(14)
        run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        table_pre = document.add_table(rows=1, cols=3)
        widths = [Cm(2), Cm(2), Cm(15.8)]
        set_col_widths(table_pre, widths)
        cell00 = table_pre.cell(0, 0)
        p = cell00.paragraphs[0]
        run = p.add_run("试卷类型")
        run.font.size = Pt(12)
        run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell01 = table_pre.cell(0, 1)
        p = cell01.paragraphs[0]
        run = p.add_run("A ")
        run.font.bold = True
        run.add_picture(r'D:\校本资源相关\mark_pic\bmp小尺寸\white.png', width=Pt(14))
        run.add_break()
        run = p.add_run("B ")
        run.font.bold = True
        run.add_picture(r'D:\校本资源相关\mark_pic\bmp小尺寸\white.png', width=Pt(14))
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell02 = table_pre.cell(0, 2)
        p = cell02.paragraphs[0]
        run = p.add_run(main_title)
        run.font.size = Pt(14)
        run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph()


def get_2dcode(stu_id):
    uid = uuid.uuid4()
    path = '{}.png'.format(uid)
    myqr.run(
        words=stu_id,
        # 扫描二维码后，显示的内容，或是跳转的链接
        version=5,  # 设置容错率
        # level='H',  # 控制纠错水平，范围是L、M、Q、H，从左到右依次升高
        # picture='we.png',  # 图片所在目录，可以是动图
        # colorized=False,  # 黑白(False)还是彩色(True)
        # contrast=1.0,  # 用以调节图片的对比度，1.0 表示原始图片。默认为1.0。
        # brightness=1.0,  # 用来调节图片的亮度，用法同上。
        save_name=path,  # 控制输出文件名，格式可以是 .jpg， .png ，.bmp ，.gif
        save_dir=os.getcwd(),
    )
    return path


def prefix_notice_1(cell):
    cell.width = Cm(10)
    tbl_sub = cell.add_table(rows=1, cols=1)
    tbl_sub.cell(0, 0).width = Cm(9.8)
    set_frame(tbl_sub)
    set_cell_border(
        tbl_sub.cell(0, 0),
        bottom={"sz": 5, "color": "#000000", "val": "single"},
    )
    p = tbl_sub.cell(0, 0).paragraphs[0]
    paragraph_format = p.paragraph_format
    paragraph_format.line_spacing = Pt(12)  # 12磅行间距
    run = p.add_run("注意事项")
    run.font.bold = True
    p.add_run().add_break()
    run = p.add_run("1．答题前，考生先将自己的姓名、班级、考场填写清楚，并认真核对")
    run.font.size = Pt(9)
    p.add_run().add_break()
    run = p.add_run("条形码上的姓名和准考证号。")
    run.font.size = Pt(9)
    p.add_run().add_break()
    run = p.add_run("2．选择题部分请按题号用2B铅笔填涂方框，修改时用橡皮擦干净，不")
    run.font.size = Pt(9)
    p.add_run().add_break()
    run = p.add_run("留痕迹。")
    run.font.size = Pt(9)
    p.add_run().add_break()
    run = p.add_run("3．非选择题部分请按题号用0.5毫米黑色墨水签字笔书写，否则作答")
    run.font.size = Pt(9)
    p.add_run().add_break()
    run = p.add_run("无效。要求字体工整、笔迹清晰。作图时，必须用2B铅笔，并描浓。")
    run.font.size = Pt(9)
    p.add_run().add_break()
    run = p.add_run("4．在草稿纸、试题卷上答题无效。")
    run.font.size = Pt(9)
    p.add_run().add_break()
    run = p.add_run("5．请勿折叠答题卡,保持字体工整、笔迹清晰、卡面清洁。")
    run.font.size = Pt(9)


def prefix_notice_2(cell, cols_count=2):
    tbl_sub = cell.add_table(rows=1, cols=1)
    tbl_sub.cell(0, 0)._width = Cm(5)
    set_frame(tbl_sub)
    set_cell_border(
        tbl_sub.cell(0, 0),
        bottom={"sz": 5, "color": "#000000", "val": "single"},
    )
    p = tbl_sub.cell(0, 0).paragraphs[0]
    paragraph_format = p.paragraph_format
    paragraph_format.line_spacing = Pt(12)  # 12磅行间距
    if cols_count == 3:
        run = p.add_run("注意事项")
        run.font.bold = True
        p.add_run().add_break()
        p.add_run("1．答题前请将姓名、班级、考").font.size = Pt(9)
        p.add_run().add_break()
        p.add_run("场、准考证号填写清楚。").font.size = Pt(9)
        p.add_run().add_break()
        p.add_run("2．客观题答题，必须使用2B铅笔").font.size = Pt(9)
        p.add_run().add_break()
        p.add_run("填涂，修改时用橡皮擦干净。").font.size = Pt(9)
        p.add_run().add_break()
        p.add_run("3．主观题答题，必须使用黑色签").font.size = Pt(9)
        p.add_run().add_break()
        p.add_run("字笔书写。").font.size = Pt(9)
        p.add_run().add_break()
        p.add_run("4．必须在题号对应的答题区域内").font.size = Pt(9)
        p.add_run().add_break()
        p.add_run("作答，超出答题区域书写无效。").font.size = Pt(9)
        p.add_run().add_break()
        p.add_run("5．保持答卷清洁、完整。").font.size = Pt(9)
    else:
        run = p.add_run("注意事项")
        run.font.bold = True
        p.add_run().add_break()
        p.add_run("1．答题前请将姓名、班级、考场、准考证号填写清楚。").font.size = Pt(9)
        p.add_run().add_break()
        p.add_run("2．客观题答题，必须使用2B铅笔填涂，修改时用橡皮擦干净。").font.size = Pt(9)
        p.add_run().add_break()
        p.add_run("3．主观题答题，必须使用黑色签字笔书写。").font.size = Pt(9)
        p.add_run().add_break()
        p.add_run("4．必须在题号对应的答题区域内作答，超出答题区域书写无效。").font.size = Pt(9)
        p.add_run().add_break()
        p.add_run("5．保持答卷清洁、完整。").font.size = Pt(9)


def prefix_notice_3(cell):
    tbl_sub = cell.add_table(rows=1, cols=1)
    tbl_sub.cell(0, 0).width = Cm(6)
    set_frame(tbl_sub)
    set_cell_border(
        tbl_sub.cell(0, 0),
        bottom={"sz": 5, "color": "#000000", "val": "single"},
    )
    p = tbl_sub.cell(0, 0).paragraphs[0]
    paragraph_format = p.paragraph_format
    paragraph_format.line_spacing = Pt(12)  # 12磅行间距
    run = p.add_run("注意事项")
    run.font.bold = True
    p.add_run().add_break()
    run = p.add_run("1．答题前请将姓名、班级、考场、准考证号")
    run.font.size = Pt(9)
    run.add_break()
    run = p.add_run("填写清楚。")
    run.font.size = Pt(9)
    run.add_break()
    run = p.add_run("2．客观题答题，必须使用2B铅笔填涂，修改")
    run.font.size = Pt(9)
    run.add_break()
    run = p.add_run("用橡皮擦干净。")
    run.font.size = Pt(9)
    run.add_break()
    run = p.add_run("3.必须在题号对应的答题区域内作答，超出")
    run.font.size = Pt(9)
    run.add_break()
    run = p.add_run("答题区域书写无效。")
    run.font.size = Pt(9)
    run.add_break()


def prefix_notice_3_cols3(document):
    tbl_sub = document.add_table(rows=1, cols=1)
    set_frame(tbl_sub)
    set_cell_border(
        tbl_sub.cell(0, 0),
        bottom={"sz": 5, "color": "#000000", "val": "single"},
    )
    p = tbl_sub.cell(0, 0).paragraphs[0]
    paragraph_format = p.paragraph_format
    paragraph_format.line_spacing = Pt(12)  # 12磅行间距
    run = p.add_run("注意事项")
    run.font.bold = True
    p.add_run().add_break()
    run = p.add_run("1．答题前请将姓名、班级、考场、准考证号填写清楚。")
    run.font.size = Pt(9)
    run.add_break()
    run = p.add_run("2．客观题答题，必须使用2B铅笔填涂，修改用橡皮擦干净。")
    run.font.size = Pt(9)
    run.add_break()
    run = p.add_run("3.必须在题号对应的答题区域内作答，超出答题区域书写无效。")
    run.font.size = Pt(9)
    run.add_break()


def prefix_codebar(cell):
    tbl1 = cell.add_table(rows=1, cols=1)
    tbl1.cell(0, 0).width = Cm(7.1)
    tbl1.cell(0, 0).height = Cm(4)
    set_cell_border(
        tbl1.cell(0, 0),
        top={"sz": 5, "val": "dashed", "color": "#000000"},
        left={"sz": 5, "val": "dashed", "color": "#000000"},
        bottom={"sz": 5, "color": "#000000", "val": "dashed"},
        right={"sz": 5, "val": "dashed", "color": "#000000", },
    )
    p = tbl1.cell(0, 0).paragraphs[0]
    run = p.add_run("贴条形码区")
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(200, 200, 200)
    p.add_run().add_break()
    p.add_run().add_break()
    p.add_run().add_break()
    run = p.add_run("(正面朝上，切勿贴出虚线方框)")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(200, 200, 200)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def prefix_absent(cell, paper_kind=0):
    tbl2 = cell.add_table(rows=1, cols=1)
    tbl2.cell(0, 0).width = Cm(7.1)
    tbl2.cell(0, 0).height = Cm(3)
    set_frame(tbl2)
    p = tbl2.cell(0, 0).paragraphs[0]
    run = p.add_run("正确填涂      ")
    run.add_picture(r'D:\校本资源相关\mark_pic\bmp小尺寸\black.png', width=Pt(14))
    run.font.bold = True
    run.font.size = Pt(9)
    run = p.add_run("   缺考标记      ")
    run.add_picture(r'D:\校本资源相关\mark_pic\bmp小尺寸\white.png', width=Pt(14))
    run.font.bold = True
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if paper_kind:
        run = p.add_run("试卷类型    A ")
        run.add_picture(r'D:\校本资源相关\mark_pic\bmp小尺寸\white.png', width=Pt(14))
        run.font.bold = True
        run.font.size = Pt(9)
        run = p.add_run("               B ")
        run.add_picture(r'D:\校本资源相关\mark_pic\bmp小尺寸\white.png', width=Pt(14))
        run.font.bold = True
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def prefix_absent_cols3(cell):
    tbl2 = cell.add_table(rows=1, cols=1)
    tbl2.cell(0, 0).width = Cm(1.6)
    tbl2.cell(0, 0).height = Cm(2)
    # set_frame(tbl2)
    p = tbl2.cell(0, 0).paragraphs[0]
    p.add_run("正确填涂").font.size = Pt(9)
    p.add_run().add_break()
    p.add_run().add_picture(r'D:\校本资源相关\mark_pic\bmp小尺寸\black.png', width=Pt(14))
    p.add_run().add_break()
    p.add_run().add_break()
    p.add_run().add_break()
    p.add_run("缺考标记").font.size = Pt(9)
    p.add_run().add_break()
    p.add_run().add_picture(r'D:\校本资源相关\mark_pic\bmp小尺寸\white.png', width=Pt(14))
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def prefix_admission_card(cell, type_test_no, stu_id_count=8):
    tbl = cell.add_table(rows=3, cols=stu_id_count)
    for i in range(stu_id_count):
        for cell in tbl.column_cells(i):
            cell.width = Cm(0.8)
    set_frame(tbl)
    p = tbl.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if type_test_no == 1:
        p.add_run("准考证号")
    else:
        p.add_run("考号")
    for i in range(stu_id_count - 1):
        tbl.cell(0, i).merge(tbl.cell(0, i + 1))
    for j in range(stu_id_count):
        p = tbl.cell(2, j).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for k in range(9):
            run = p.add_run("[{}]".format(k))
            run.font.size = Pt(9)
            run.add_break()
        run = p.add_run("[{}]".format(9))
        run.font.size = Pt(9)


def prefix_signature(p):
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run().add_break()
    run = p.add_run("姓名：_________")
    run.font.bold = True
    run.add_break()
    run.add_break()
    run = p.add_run("班级：_________")
    run.font.bold = True
    run.add_break()
    run.add_break()
    run = p.add_run("考号：_________")
    run.font.bold = True
    run.add_break()
    run.add_break()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def prefix_absent_mark(p):
    p.add_run().add_break()
    run = p.add_run("缺考")
    run.font.bold = True
    run.font.size = Pt(9)
    run.add_break()
    run = p.add_run("标记")
    run.add_break()
    run.add_picture(r'D:\校本资源相关\mark_pic\bmp小尺寸\white.png', width=Pt(14))
    run.font.bold = True
    run.font.size = Pt(9)


def prefix_2dcode(cell, path, type_test_no, cols_count=1, absent=0):
    if cols_count == 2 or cols_count == 1:
        if absent == 1:
            tbl2 = cell.add_table(rows=1, cols=3)
            p = tbl2.cell(0, 0).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run().add_picture(path, width=Pt(90))

            p = tbl2.cell(0, 1).paragraphs[0]
            prefix_signature(p)

            p = tbl2.cell(0, 2).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run().add_break()
            prefix_absent_mark(p)
        else:
            tbl2 = cell.add_table(rows=1, cols=2)
            # set_frame(tbl2)
            for cell in tbl2.column_cells(0):
                cell.width = Cm(4)
            for cell in tbl2.column_cells(1):
                cell.width = Cm(4)
            p = tbl2.cell(0, 0).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run().add_picture(path, width=Pt(90))
            p = tbl2.cell(0, 1).paragraphs[0]
            prefix_signature(p)
    else:
        tbl2 = cell.add_table(rows=2, cols=2)
        p = tbl2.cell(0, 0).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Pt(90))
        if absent:
            p = tbl2.cell(0, 1).paragraphs[0]
            prefix_absent_mark(p)
        if type_test_no != 0:
            p = tbl2.cell(1, 0).paragraphs[0]
            tbl2.cell(1, 0).merge(tbl2.cell(1, 1))
            prefix_signature(p)


def prefix_2dcode_(cell, path, absent):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    tbl2 = cell.add_table(rows=2, cols=1)
    set_frame(tbl2)
    tbl2.cell(0, 0).width = Cm(6)
    tbl2.cell(1, 0).width = Cm(6)
    tbl2.cell(0, 0).height = Cm(6)
    tbl2.cell(1, 0).height = Cm(2)
    p = tbl2.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Pt(90))
    p = tbl2.cell(1, 0).paragraphs[0]
    run = p.add_run("正确填涂  ")
    run.add_picture(r'D:\校本资源相关\mark_pic\bmp小尺寸\black.png', width=Pt(14))
    run.font.bold = True
    run.font.size = Pt(9)
    if absent:
        run = p.add_run("    缺考标记  ")
        run.add_picture(r'D:\校本资源相关\mark_pic\bmp小尺寸\white.png', width=Pt(14))
        run.font.bold = True
        run.font.size = Pt(9)


def prefix_2dcode_cols3(cell, path, absent):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run().add_break()
    p.add_run().add_picture(path, width=Cm(2))
    p.add_run().add_break()
    run = p.add_run("正确填涂  ")
    run.add_picture(r'D:\校本资源相关\mark_pic\bmp小尺寸\black.png', width=Pt(14))
    run.font.bold = True
    run.font.size = Pt(9)
    if absent:
        p.add_run().add_break()
        run = p.add_run("缺考标记  ")
        run.add_picture(r'D:\校本资源相关\mark_pic\bmp小尺寸\white.png', width=Pt(14))
        run.font.bold = True
        run.font.size = Pt(9)


def prefix_short_card(cell):
    cell.vertical = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("准考证号后三位数字")
    run.font.size = Pt(9)
    tbl = cell.add_table(rows=3, cols=3)
    set_frame(tbl)
    for cell in tbl.column_cells(0):
        cell.width = Cm(1)
    for cell in tbl.column_cells(1):
        cell.width = Cm(1)
    for cell in tbl.column_cells(2):
        cell.width = Cm(6.8)
    for cell in tbl.row_cells(0):
        cell.height = Cm(1)
    for cell in tbl.row_cells(1):
        cell.height = Cm(1)
    p = tbl.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("短")
    run.font.size = Pt(9)
    p = tbl.cell(1, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("考")
    run.font.size = Pt(9)
    p = tbl.cell(2, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("号")
    run.font.size = Pt(9)
    for i in range(2):
        tbl.cell(i, 0).merge(tbl.cell(i + 1, 0))
    for j in range(3):
        cell2 = tbl.cell(j, 2)
        cell2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for k in range(9):
            p = cell2.paragraphs[0]
            run = p.add_run("[{}]".format(k) + " ")
            run.font.size = Pt(9)
        run = p.add_run("[{}]".format(9))
        run.font.size = Pt(9)
    # tbl.cell(0, 2).merge(tbl.cell(1, 2)).merge(tbl.cell(2, 2))


def set_cell_background_color(cell):
    # if not isinstance(rgb_color, RGBValue):
    #     print('rgbColor is not RGBValue...', type(rgb_color))
    #     return
    # hr = str(hex(int(rgb_color.r)))[-2:]
    # hg = str(hex(int(rgb_color.g)))[-2:]
    # hb = str(hex(int(rgb_color.b)))[-2:]
    # color_str = hr + hg + hb
    # shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value=color_str))
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="CCCAC9"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading_elm_1)


def add_postfix_info(document, page_count):
    page_binary = get_binary_code(page_count)
    for i in range(page_count):
        section = document.add_section()
        footer = section.footer
        p = footer.paragraphs[0]
        p.add_run('{}'.format(page_binary[i]))


def get_binary_code(page_count):
    sum = []
    for i in range(1, page_count + 1):
        binary = []
        m, n = divmod(i, 2)
        recurrent(m, n, binary)
        sum.append(binary)
    blank_count = max(max([len(ls) for ls in sum]), 3)
    for ls in sum:
        gap = blank_count - len(ls)
        for i in range(gap):
            ls.insert(0, 0)
    return sum


def recurrent(m, n, binary):
    if m == 0:
        binary.insert(0, n)
    else:
        binary.insert(0, n)
        m2, n2 = divmod(m, 2)
        recurrent(m2, n2, binary)


def write_forbidden(document):
    table = document.add_table(rows=2, cols=1)
    set_cell_border(
        table.cell(0, 0),
        top={"sz": 5, "val": "single", "color": "#000000"},
        left={"sz": 5, "val": "single", "color": "#000000"},
        # bottom={"sz": 5, "color": "#000000", "val": "single"},
        right={"sz": 5, "val": "single", "color": "#000000", },
    )
    set_cell_border(
        table.cell(1, 0),
        # top={"sz": 5, "val": "single", "color": "#000000"},
        left={"sz": 5, "val": "single", "color": "#000000"},
        bottom={"sz": 5, "color": "#000000", "val": "single"},
        right={"sz": 5, "val": "single", "color": "#000000", },
    )
    cell1 = table.cell(0, 0)
    set_cell_background_color(cell1)
    p = cell1.paragraphs[0]
    p.add_run().add_picture(r'D:\校本资源相关\答题卡\forbidden.png')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_break()
    cell2 = table.cell(1, 0)
    set_cell_background_color(cell2)
    p = cell2.paragraphs[0]
    run = p.add_run("请勿在此区域作答或者做任何标记")
    run.font.size = Pt(26)
    run.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def write4answersheet(answer_sheet_data):
    '''
    answer_sheet_data = {
        "name": paper_data["name"],
        "main_title": "{}-答题卡".format(paper_data["main_title"]),
        "testee": paper_data["testee"],
        "sheet": answer_sheet
    }
    answer_sheet[key].append({
                "q_no": ques_no,
                "pic_in_card": pic_in_card,
                "options": [v["key"] for v in options if v.get("key")],
                "subques_len": len(subquestions) if subquestions else 0
            })
    '''
    link_url = "http://www.baidu.com"  # APP下载链接地址
    sheet = answer_sheet_data.get('sheet', {})
    sheet_name = answer_sheet_data.get('name')
    main_title = answer_sheet_data.get('main_title')
    subject_id = answer_sheet_data.get('subject_id')
    scan_type = answer_sheet_data.get('scan_type')  # 默认是0：网阅卡，另一值1：手阅
    type_exam_no = answer_sheet_data.get('type_exam_no')  # 0:条形码，1：准考证，2：学籍号，3：短考号
    absent = answer_sheet_data.get('absent')  # 默认是0：不加缺考标记，另一值1：添加缺考标记
    paper_kind = answer_sheet_data.get('paper_kind')  # 默认是0：不加A/B卷，另一值1：添加A/B卷
    simi_score = answer_sheet_data.get('simi_score')  # 默认是0：不支持0.5分，另一值1：支持0.5分
    cols_count = answer_sheet_data.get('cols_count')  # 默认是1：一栏，A4；2：两栏，A3；3：三栏，A3
    answer_path = r'C:\Users\j20687\Desktop\answersheet_3.docx'
    document = Document(answer_path)
    # 设置整个文档的默认字体
    microsoft_font = u'宋体'  # u 表示后面的字符串以 Unicode 格式进行编码
    black_font = u'黑体'
    number_font = 'Times New Roman'
    area = qn('w:eastAsia')
    document.styles['Normal'].font.name = microsoft_font
    document.styles['Normal'].font.size = Pt(10.5)
    document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
    document.styles['Normal']._element.rPr.rFonts.set(area, number_font)
    # 指定段落样式
    styles = document.styles
    # fls = time.time()
    # strr = 'lbj_Style%s' %fls  #自定义的样式的名称
    # strr = strr.replace('.', '')
    # strr = strr + ''.join(random.sample('zyxwvutsrqponmlkjihgfedcbaABCDEFGHIJKLMNOPQRST', 5))
    s = styles.add_style('lbj_Style', WD_STYLE_TYPE.PARAGRAPH)
    s.font.name = microsoft_font
    s.font.size = Pt(10.5)
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.line_spacing = Pt(0)  # 行距值
    s.paragraph_format.space_before = Pt(0)  # 段前距
    s.paragraph_format.space_after = Pt(0)  # 段后距
    s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE  # 行间距
    s._element.rPr.rFonts.set(area, number_font)  # 除中文外其它文字 使用的字体 ，备选项

    # 编辑页眉
    document.settings.odd_and_even_pages_header_footer = True
    section = document.sections[0]
    header = section.header
    header_tbl = header.add_table(1, 2, width=Cm(13))
    set_col_widths(header_tbl, [Cm(8.75), Cm(4.25)])
    header_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    p = header_tbl.cell(0, 0).paragraphs[0]
    run = p.add_run('扫码下载蘑信学生端查分！   ')
    run.bold = True
    run.font.size = Pt(9)
    p.add_run().add_break()
    p.add_run('（用户名和初始密码均为准考证号）').font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cell = header_tbl.cell(0, 1)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p = cell.paragraphs[0]
    path = get_2dcode(link_url)
    p.add_run().add_picture(path, width=Cm(1.3))
    os.remove(path)

    # 标题
    # p = document.paragraphs[0]
    # run = p.add_run(main_title)
    # run.font.size = Pt(14)
    # run.font.bold = True
    # p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 前置信息
    type_exam_no, scan_type, absent = 2, 0, 1  # 考号版本；答题卡模式
    cols_count = 3
    paper_kind = 1
    add_prefix_info(document, main_title, type_exam_no, scan_type, subject_id, cols_count, absent, paper_kind,
                    stu_id_count=12)

    p = document.add_paragraph()
    # p.add_run().add_break()
    run = p.add_run('选择题（请用2B铅笔填涂）')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = black_font
    # sub_mark_num_vec = []  # 每页竖长条表示块数量
    sub_mark_num = 0  # 竖长条表示块总数量
    sub_mark_choice = []  # 选择题黑块区域
    sub_mark_fillin = []  # 填空题黑块区域
    sub_mark_eassy = []  # 简答题黑块区域
    qus_range_choice = []  # 选择题题号范围
    qus_range_fillin = []  # 填空题题号范围
    qus_range_eassy = []  # 简答题题号范围
    paper_num = 0  # 页数
    ques_no = 0  # 题号
    choice_answer_cnt = []  # 每个区块的选项数
    choice_line_number = []  # 每大行的最多题目数量
    choice_answer_pattern = []  # 每行的每列选择题数目
    choice_col_num = []  # 选择题每行的列数
    choice_row_num = []  # 选择题行数
    option_len = []
    choice_num, mm, combine_ques_index = [], [], []
    count_ques = 0
    for key, value in sheet.items():
        if key == 'choice_question_all':
            for k in range(len(value)):
                q_no = value[k].get("q_no")
                opt_len = len(value[k].get("options")) if len(value[k].get("options")) >= 4 else 4
                mm.append(q_no)
                if not isinstance(q_no, list):
                    option_len.append(opt_len)
                else:
                    for qq in q_no:
                        option_len.append(opt_len)
    for i in range(len(mm)):
        if isinstance(mm[i], list):
            count_ques += len(mm[i])
            combine_ques_index.append(i)
            for j in mm[i]:
                choice_num.append(j)
        else:
            count_ques += 1
            choice_num.append(mm[i])
    # logger.info("option_len========================={}".format(option_len))
    print("option_len========================={}".format(option_len))
    choice_group = sheet.get('choice_question', [])
    choice_len = len(choice_group)
    choice_len_sum = len(choice_num)  # 选择题总数，包括组合题中的小选择题
    fillin_group = sheet.get('fillin_question', [])
    fillin_len = len(fillin_group)
    fs_num = fillin_group[0].get('q_no') if fillin_group else 0  # 填空题开始的题号
    fs_num_list = [ques.get('q_no') for ques in fillin_group]
    subjective_question = sheet.get('subjective_question', [])
    subjec_len = len(subjective_question)
    ss_num = subjective_question[0].get('q_no') if subjective_question else 0  # 主观题开始的题号
    ss_num_list = [ques.get('q_no') for ques in subjective_question]
    ss_score_list = [int(ques.get('score')) for ques in subjective_question]  # 各主观题的分值
    combine_question = sheet.get('combine_question', [])

    pic_in_sub = [[] for x in range(subjec_len)]
    pic_num = 0
    for ques in subjective_question:
        pic_in_card = ques.get('pic_in_card', "")
        pic_in_sub[pic_num] = pic_in_card if pic_in_card else []
        pic_num += 1
    tmp = []
    for i in range(len(mm)):
        if isinstance(mm[i], list):
            if tmp:
                qus_range_choice.append(tmp)
                tmp = []
            qus_range_choice.append(mm[i])
        else:
            tmp.append(mm[i])

    if not combine_question:
        qus_range_choice.append(tmp)
    else:
        if not isinstance(mm[-1], list):
            qus_range_choice.append(tmp)

    if fillin_group:
        qus_range_fillin = fs_num_list
    if subjective_question:
        qus_range_eassy = ss_num_list

    # 选项个数不一样
    group_choice_num_ = []
    group_choice_num = {}  # 按题号顺序各选项数对应的题号集合
    is_multi = False
    if option_len:
        pre = option_len[0]
        group_choice_num = {pre: []}
        group_choice_num[pre].append(choice_num[0])
    for opt_i in range(1, len(option_len)):
        now = option_len[opt_i]
        pre = option_len[opt_i - 1]
        if pre != now:
            is_multi = True
            group_choice_num_.append(group_choice_num)
            group_choice_num = {now: []}
            # group_choice_num[now].append(choice_num[opt_i])
        group_choice_num[now].append(choice_num[opt_i])
    # logger.info("group_choice_num========================={}".format(group_choice_num))
    # if is_multi:
    group_choice_num_.append(group_choice_num)
    # logger.info("group_choice_num_========================={}".format(group_choice_num_))
    cols = 0
    for item in group_choice_num_:
        for key, value in item.items():
            d1, d2 = divmod(len(value), 5)
            for i in range(d1):
                cols += 1
                choice_answer_pattern.append(5)
                choice_answer_cnt.append(key)
            if d2:
                cols += 1
                choice_answer_pattern.append(d2)
                choice_answer_cnt.append(key)
    # logger.info("cols========================={}".format(cols))
    c1, c2 = divmod(cols, 3)
    row = c1 + 1 if c2 else c1
    choice_row_num.append(row)
    for i in range(c1):
        choice_col_num.append(3)
    if c2:
        choice_col_num.append(c2)
    tmp4line = []
    plusn = 0
    for j in range(len(choice_answer_pattern)):
        tmp4line.append(choice_answer_pattern[j])
        plusn += 1
        if plusn == 3:
            choice_line_number.append(max(tmp4line))
            tmp4line = []
            plusn = 0
    if tmp4line:
        choice_line_number.append(max(tmp4line))
    break_sum = 0  # 一页最多存放（5+5+5+5+1）* 3个选项
    break_row = 0
    break_num = 0
    for cl in range(len(choice_line_number)):
        break_sum += choice_line_number[cl]
        if break_sum > 21 and cl > 1:
            break_row = cl
            break_sum -= choice_line_number[cl]
            break
    for br in range(break_row * 3):
        break_num += choice_answer_pattern[br] if len(choice_answer_pattern) >= br else 0

    choice_answer_pattern = [choice_answer_pattern] if not break_num else [choice_answer_pattern[:break_row * 3],
                                                                           choice_answer_pattern[break_row * 3:]]
    choice_answer_cnt = [choice_answer_cnt] if not break_num else [choice_answer_cnt[:break_row * 3],
                                                                   choice_answer_cnt[break_row * 3:]]
    choice_row_num = choice_row_num if not break_num else [break_row, row - break_row]
    choice_col_num = [choice_col_num] if not break_num else [choice_col_num[0:break_row], choice_col_num[break_row:]]
    choice_line_number = [choice_line_number] if not break_num else [choice_line_number[0:break_row],
                                                                     choice_line_number[break_row:]]

    # logger.info("choice_line_number========================={}".format(choice_line_number))
    # logger.info("row========================={}".format(row))
    # logger.info("choice_answer_pattern========================={}".format(choice_answer_pattern))
    if not break_num:
        write_option(choice_num, option_len, choice_answer_pattern[0], row, document)
    else:
        write_option(choice_num[0:break_num], option_len[0:break_num], choice_answer_pattern[0], break_row, document)
        if break_num != 21:
            document.add_page_break()
        write_option(choice_num[break_num:], option_len[break_num:], choice_answer_pattern[1], row - break_row,
                     document)
    # write_option(ques_no, choice_len_sum, option_len, choice_ques_no, document)
    writeline(document, black_font)
    sub_fillin(fillin_len, fs_num_list, document, scan_type, cols_count)

    if ss_num:
        for ns in range(len(ss_num_list)):
            number = ss_num_list[ns]
            score = ss_score_list[ns] if scan_type else None
            # score = ss_score_list[ns]
            sub_subjective(number, document, pic_in_sub, score, simi_score, cols_count)

    block_num = 0
    start = 0
    step = 3
    if choice_len_sum:
        for n in range(len(choice_row_num)):
            block_num += 1
            sub_mark_choice.append(start + step * (block_num - 1) - (block_num - 1))
            sub_mark_choice.append(start + step * block_num - (block_num - 1))
    if fillin_len:
        block_num += 1
        sub_mark_fillin.append(start + step * (block_num - 1) - (block_num - 1))
        sub_mark_fillin.append(start + step * block_num - (block_num - 1))
        if not subjec_len:
            index = sub_mark_fillin[-2]
            sub_mark_fillin.pop(-1)
            sub_mark_fillin.append(index + 1)
        else:
            for n in range(subjec_len):
                block_num += 1
                sub_mark_eassy.append(start + step * (block_num - 1) - (block_num - 1))
                sub_mark_eassy.append(start + step * block_num - (block_num - 1))
            index = sub_mark_eassy[-2]
            sub_mark_eassy.pop(-1)
            sub_mark_eassy.append(index + 1)
    else:
        if subjec_len:
            for n in range(subjec_len):
                block_num += 1
                sub_mark_eassy.append(start + step * (block_num - 1) - (block_num - 1))
                sub_mark_eassy.append(start + step * block_num - (block_num - 1))
            index = sub_mark_eassy[-2]
            sub_mark_eassy.pop(-1)
            sub_mark_eassy.append(index + 1)
        else:
            if sub_mark_choice:
                index = sub_mark_choice[-2]
                sub_mark_choice.pop(-1)
                sub_mark_choice.append(index + 1)

    #  return sub_mark_num
    for i in range(len(choice_row_num)):
        sub_mark_num += 2
    # if choice_len_sum:
    #     sub_mark_num += 4
    if fillin_len:
        sub_mark_num += 2
    for num in range(subjec_len):
        sub_mark_num += 2

    t1 = int(round(time.time() * 1000))
    doc_file_name = "test-答题卡"
    return_path = r'C:\Users\j20687\Desktop\%s.docx' % (doc_file_name)
    document.save(return_path)
    sub_json_data = {}
    tmp_sub_mark_choice = []
    if choice_num:
        for i in range(len(choice_row_num)):
            tmp_sub_mark_choice.append([sub_mark_choice[i * 2], sub_mark_choice[i * 2 + 1]])
    # logger.info("sub_mark_choice========================={}".format(sub_mark_choice))
    sub_json_data['sub_mark_choice'] = tmp_sub_mark_choice
    sub_json_data['sub_mark_fillin'] = sub_mark_fillin
    sub_json_data['sub_mark_eassy'] = sub_mark_eassy
    sub_json_data['sub_mark_num'] = sub_mark_num
    # sub_json_data['sub_mark_num_vec'] = sub_mark_num_vec
    # page = get_pdf_page(return_path)
    # logger.info("page========================={}".format(page))
    sub_json_data['page_num'] = paper_num
    sub_json_data['scan_type'] = scan_type
    sub_json_data['simi_score'] = simi_score
    # sub_json_data['paper_num_count_flag'] = paper_num_count_flag
    sub_json_data['choice_answer_cnt'] = choice_answer_cnt
    sub_json_data['choice_line_number'] = choice_line_number
    sub_json_data['choice_answer_pattern'] = choice_answer_pattern
    sub_json_data['choice_col_num'] = choice_col_num
    sub_json_data['choice_row_num'] = choice_row_num
    sub_json_data['qus_range_choice'] = qus_range_choice
    sub_json_data['qus_range_fillin'] = qus_range_fillin
    sub_json_data['qus_range_eassy'] = qus_range_eassy
    sub_json_data['eassy_score'] = ss_score_list
    # logger.info("sub_json_data========================={}".format(sub_json_data))
    return return_path, sub_json_data


def get4answer_sheet_data(paper_data):
    answer_sheet = {}
    ques_no = 1
    for g_idx, grp in enumerate(paper_data.get("question_group", [])):
        for eq in grp.get("exam_questions", []):
            question = eq.get("question", {})
            score = eq.get("score") or 0
            q_type_name = question.get("q_type", {}).get("name")
            pic_in_card = question.get("pic_in_card", "")
            options = question.get("options") or []
            subquestions = question.get("subs") or []
            if q_type_name in CHOICE_QUESTIONS:
                key = "choice_question"
            elif subquestions:
                key = "combine_question"
            elif q_type_name in FILL_IN_QUESTIONS:
                key = "fillin_question"
            else:
                key = "subjective_question"
            if not answer_sheet.get(key):
                answer_sheet[key] = []
            answer_sheet[key].append({
                "q_no": ques_no,
                "score": score,
                "pic_in_card": pic_in_card,
                "options": [v["key"] for v in options if v.get("key")],
                "subques_len": len(subquestions) if subquestions else 0
            })
            if q_type_name in CHOICE_QUESTIONS or subquestions:
                key = "choice_question_all"
                if not answer_sheet.get(key):
                    answer_sheet[key] = []
                answer_sheet[key].append({
                    "q_no": ques_no,
                    "pic_in_card": pic_in_card,
                    "options": [v["key"] for v in options if v.get("key")],
                    "subques_len": len(subquestions) if subquestions else 0
                })
            if key == "combine_question":
                ques_no += len(subquestions)
            else:
                ques_no += 1
    answer_sheet_data = {
        "name": paper_data.get("name") or paper_data.get("main_title"),
        "subject_id": paper_data.get("subject").get("uid"),
        "main_title": "{}-答题卡".format(paper_data.get("main_title")),
        "testee": paper_data.get("testee"),
        "scan_type": paper_data.get("scan_type") or 0,
        "type_exam_no": paper_data.get("type_exam_no") or 0,
        "absent": paper_data.get("absent") or 0,
        "paper_kind": paper_data.get("paper_kind") or 0,
        "simi_score": paper_data.get("simi_score") or 0,
        "cols_count": paper_data.get("cols_count") or 1,
        "sheet": answer_sheet
    }
    answer_path, sub_json_data = write4answersheet(answer_sheet_data)
    sheet_json = get_json(sub_json_data)
    return answer_path, answer_sheet_data, sheet_json


def get_pdf_page(docx_path):
    t1 = time.time()
    file_name = docx_path.split("/")[-1]
    file_name = ".".join(file_name.split(".")[:-1])
    pdf_path = '/tmp/{}.pdf'.format(file_name)
    os.system('soffice --headless --invisible --convert-to pdf %s --outdir /tmp' % docx_path)
    try:
        f = pdfplumber.open(pdf_path)
        page = len(f.pages)
    except PDFSyntaxError:
        page = 0
    t2 = time.time()
    logger.info("get_pdf_page costs========================={}s".format(t2 - t1))
    return page


def pre_info(sub_json_data):
    pre_info_json = {}
    pre_info_json['height'] = 1122
    pre_info_json['width'] = 793
    pre_info_json['scan_type'] = sub_json_data['scan_type']
    pre_info_json['simi_score'] = sub_json_data['simi_score']
    pre_info_json[
        'mark_base64_str'] = "/9j/4AAQSkZJRgABAQEAeAB4AAD/4QBORXhpZgAATU0AKgAAAAgABAMBAAUAAAABAAAAPlEQAAEAAAABAQAAAFERAAQAAAABAAASdFESAAQAAAABAAASdAAAAAAAAYagAACxj//bAEMAAgEBAgEBAgICAgICAgIDBQMDAwMDBgQEAwUHBgcHBwYHBwgJCwkICAoIBwcKDQoKCwwMDAwHCQ4PDQwOCwwMDP/bAEMBAgICAwMDBgMDBgwIBwgMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDP/AABEIABsAGwMBIgACEQEDEQH/xAAfAAABBQEBAQEBAQAAAAAAAAAAAQIDBAUGBwgJCgv/xAC1EAACAQMDAgQDBQUEBAAAAX0BAgMABBEFEiExQQYTUWEHInEUMoGRoQgjQrHBFVLR8CQzYnKCCQoWFxgZGiUmJygpKjQ1Njc4OTpDREVGR0hJSlNUVVZXWFlaY2RlZmdoaWpzdHV2d3h5eoOEhYaHiImKkpOUlZaXmJmaoqOkpaanqKmqsrO0tba3uLm6wsPExcbHyMnK0tPU1dbX2Nna4eLj5OXm5+jp6vHy8/T19vf4+fr/xAAfAQADAQEBAQEBAQEBAAAAAAAAAQIDBAUGBwgJCgv/xAC1EQACAQIEBAMEBwUEBAABAncAAQIDEQQFITEGEkFRB2FxEyIygQgUQpGhscEJIzNS8BVictEKFiQ04SXxFxgZGiYnKCkqNTY3ODk6Q0RFRkdISUpTVFVWV1hZWmNkZWZnaGlqc3R1dnd4eXqCg4SFhoeIiYqSk5SVlpeYmZqio6Slpqeoqaqys7S1tre4ubrCw8TFxsfIycrS09TV1tfY2dri4+Tl5ufo6ery8/T19vf4+fr/2gAMAwEAAhEDEQA/APyo+OP7XPxP0D4n+JPs/jvxpbgaxeeX5WsXMa4M8gIwHAwAAPwrhb39tH4s74kj+JPj0GIkEJ4jvMEe4Elc18VbsXPjDUpH88MbqRtpORneSf8AP1rl5Bh/kygbDBgPvZNaRhHsc7nK57GP26vi/bNGB8TviLBKfmOPEl6RyMdPN9Ko6x+2X8W9Q1KWZvit8RiXI6+Kr0HgAf8APQ+leXz3qw3oMo8zBAyeD2qDV7KS71GWSOFmViOc5ycDP60nBLoJTkaPjvV49W8T3l0jBFuJDIBjBw2SB+Rrn5JmeRSWYBPlyB0r+nzxv/wb5fsg6XoH2iD4QokxAO7/AISfWT6jobvHauQu/wDggt+ydE9sF+E6ASx7m/4qTWOT/wCBVbxhdEuXLLY/m6tpAXAkf7xyWYflWnc2K3ExcPCAcfeOD0r9zPjZ/wAEcv2b/B2uyRad8OFt41UkA67qcmPlP965NcHd/wDBKP4AzTBm8AqWZFJP9taj/dH/AE3pThYUZ+R//9k="
    pre_info_json[
        'target_base64_str'] = "/9j/4AAQSkZJRgABAQEAeAB4AAD/4QBORXhpZgAATU0AKgAAAAgABAMBAAUAAAABAAAAPlEQAAEAAAABAQAAAFERAAQAAAABAAASdFESAAQAAAABAAASdAAAAAAAAYagAACxj//bAEMAAgEBAgEBAgICAgICAgIDBQMDAwMDBgQEAwUHBgcHBwYHBwgJCwkICAoIBwcKDQoKCwwMDAwHCQ4PDQwOCwwMDP/bAEMBAgICAwMDBgMDBgwIBwgMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDP/AABEIABkACAMBIgACEQEDEQH/xAAfAAABBQEBAQEBAQAAAAAAAAAAAQIDBAUGBwgJCgv/xAC1EAACAQMDAgQDBQUEBAAAAX0BAgMABBEFEiExQQYTUWEHInEUMoGRoQgjQrHBFVLR8CQzYnKCCQoWFxgZGiUmJygpKjQ1Njc4OTpDREVGR0hJSlNUVVZXWFlaY2RlZmdoaWpzdHV2d3h5eoOEhYaHiImKkpOUlZaXmJmaoqOkpaanqKmqsrO0tba3uLm6wsPExcbHyMnK0tPU1dbX2Nna4eLj5OXm5+jp6vHy8/T19vf4+fr/xAAfAQADAQEBAQEBAQEBAAAAAAAAAQIDBAUGBwgJCgv/xAC1EQACAQIEBAMEBwUEBAABAncAAQIDEQQFITEGEkFRB2FxEyIygQgUQpGhscEJIzNS8BVictEKFiQ04SXxFxgZGiYnKCkqNTY3ODk6Q0RFRkdISUpTVFVWV1hZWmNkZWZnaGlqc3R1dnd4eXqCg4SFhoeIiYqSk5SVlpeYmZqio6Slpqeoqaqys7S1tre4ubrCw8TFxsfIycrS09TV1tfY2dri4+Tl5ufo6ery8/T19vf4+fr/2gAMAwEAAhEDEQA/APx6/Zm+Juv6r8afh9HLrmqSSah4osftDfa5C7l79CxY5zznrRWL+yBCZv2g/hkoTcp8V6XkY6/6fFiitlBW2OZT1ZP+xXKP+GkvhWq/f/4S7SVORkAHUIf8/jRT/wBi/wD5Ob+F/wD2OOj/APpxioojsQ1q/U//2Q=="
    pre_info_json["page_template_base64_str"] = [
        '/9j/4AAQSkZJRgABAQEAeAB4AAD/2wBDAAIBAQIBAQICAgICAgICAwUDAwMDAwYEBAMFBwYHBwcGBwcICQsJCAgKCAcHCg0KCgsMDAwMBwkODw0MDgsMDAz/2wBDAQICAgMDAwYDAwYMCAcIDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAz/wAARCAAYAA8DASIAAhEBAxEB/8QAHwAAAQUBAQEBAQEAAAAAAAAAAAECAwQFBgcICQoL/8QAtRAAAgEDAwIEAwUFBAQAAAF9AQIDAAQRBRIhMUEGE1FhByJxFDKBkaEII0KxwRVS0fAkM2JyggkKFhcYGRolJicoKSo0NTY3ODk6Q0RFRkdISUpTVFVWV1hZWmNkZWZnaGlqc3R1dnd4eXqDhIWGh4iJipKTlJWWl5iZmqKjpKWmp6ipqrKztLW2t7i5usLDxMXGx8jJytLT1NXW19jZ2uHi4+Tl5ufo6erx8vP09fb3+Pn6/8QAHwEAAwEBAQEBAQEBAQAAAAAAAAECAwQFBgcICQoL/8QAtREAAgECBAQDBAcFBAQAAQJ3AAECAxEEBSExBhJBUQdhcRMiMoEIFEKRobHBCSMzUvAVYnLRChYkNOEl8RcYGRomJygpKjU2Nzg5OkNERUZHSElKU1RVVldYWVpjZGVmZ2hpanN0dXZ3eHl6goOEhYaHiImKkpOUlZaXmJmaoqOkpaanqKmqsrO0tba3uLm6wsPExcbHyMnK0tPU1dbX2Nna4uPk5ebn6Onq8vP09fb3+Pn6/9oADAMBAAIRAxEAPwD6c1LQfiL/AMF2v2+PjT4at/i38SfhD+zj+z1rf/CHwf8ACB6mNK1XxZ4giDC8eS4KsRFAwKhWjZGV4mXJLEdL+xx8U/iT/wAEtv8AgplpP7KHxQ+I3iL4qfCz4paDd+IPhl4v8V3P2jW7G6tVeW70u6uP+WqrEkkgd9oX90qY37Fsf8Gwc8t18C/2npbwAanJ+0P4oa6DIEkDG304ncuBj5t/GODmuE/4OH7yaw/4KN/sOzWEskGoInjvDwMVmVTptiByvzY+/wD+Pe9ADPhj+0dof/Bvd+3T8efD3xuGqaB8A/j54rn+Ifgbxla6Vc6la2upXIzf6Zci2R5EdcR+UojJ8uEsWbd8tn9nvWLj/gud/wAFdPCv7QOg6TrFr+zL8ANC1HSPCusajZyWLeNdYvont7qSGORQxt41YgnClXt4weXZEKKAP//Z',
        '/9j/4AAQSkZJRgABAQEAeAB4AAD/2wBDAAIBAQIBAQICAgICAgICAwUDAwMDAwYEBAMFBwYHBwcGBwcICQsJCAgKCAcHCg0KCgsMDAwMBwkODw0MDgsMDAz/2wBDAQICAgMDAwYDAwYMCAcIDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAz/wAARCAAXAAgDASIAAhEBAxEB/8QAHwAAAQUBAQEBAQEAAAAAAAAAAAECAwQFBgcICQoL/8QAtRAAAgEDAwIEAwUFBAQAAAF9AQIDAAQRBRIhMUEGE1FhByJxFDKBkaEII0KxwRVS0fAkM2JyggkKFhcYGRolJicoKSo0NTY3ODk6Q0RFRkdISUpTVFVWV1hZWmNkZWZnaGlqc3R1dnd4eXqDhIWGh4iJipKTlJWWl5iZmqKjpKWmp6ipqrKztLW2t7i5usLDxMXGx8jJytLT1NXW19jZ2uHi4+Tl5ufo6erx8vP09fb3+Pn6/8QAHwEAAwEBAQEBAQEBAQAAAAAAAAECAwQFBgcICQoL/8QAtREAAgECBAQDBAcFBAQAAQJ3AAECAxEEBSExBhJBUQdhcRMiMoEIFEKRobHBCSMzUvAVYnLRChYkNOEl8RcYGRomJygpKjU2Nzg5OkNERUZHSElKU1RVVldYWVpjZGVmZ2hpanN0dXZ3eHl6goOEhYaHiImKkpOUlZaXmJmaoqOkpaanqKmqsrO0tba3uLm6wsPExcbHyMnK0tPU1dbX2Nna4uPk5ebn6Onq8vP09fb3+Pn6/9oADAMBAAIRAxEAPwDsv+Cxvxp/aJ+Pnxi+AHizxRodz8JvgPF8d/Dnhzw54Vu5lfWPGMwnknOs3ojJENuvkBYLdssfMeRgMRklfTf/AAcsf8ih+yN/2cV4Y/8AQLuigA/4OWP+RQ/ZG/7OK8Mf+gXdFFFAH//Z',
        '/9j/4AAQSkZJRgABAQEAeAB4AAD/2wBDAAIBAQIBAQICAgICAgICAwUDAwMDAwYEBAMFBwYHBwcGBwcICQsJCAgKCAcHCg0KCgsMDAwMBwkODw0MDgsMDAz/2wBDAQICAgMDAwYDAwYMCAcIDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAz/wAARCAAYAA8DASIAAhEBAxEB/8QAHwAAAQUBAQEBAQEAAAAAAAAAAAECAwQFBgcICQoL/8QAtRAAAgEDAwIEAwUFBAQAAAF9AQIDAAQRBRIhMUEGE1FhByJxFDKBkaEII0KxwRVS0fAkM2JyggkKFhcYGRolJicoKSo0NTY3ODk6Q0RFRkdISUpTVFVWV1hZWmNkZWZnaGlqc3R1dnd4eXqDhIWGh4iJipKTlJWWl5iZmqKjpKWmp6ipqrKztLW2t7i5usLDxMXGx8jJytLT1NXW19jZ2uHi4+Tl5ufo6erx8vP09fb3+Pn6/8QAHwEAAwEBAQEBAQEBAQAAAAAAAAECAwQFBgcICQoL/8QAtREAAgECBAQDBAcFBAQAAQJ3AAECAxEEBSExBhJBUQdhcRMiMoEIFEKRobHBCSMzUvAVYnLRChYkNOEl8RcYGRomJygpKjU2Nzg5OkNERUZHSElKU1RVVldYWVpjZGVmZ2hpanN0dXZ3eHl6goOEhYaHiImKkpOUlZaXmJmaoqOkpaanqKmqsrO0tba3uLm6wsPExcbHyMnK0tPU1dbX2Nna4uPk5ebn6Onq8vP09fb3+Pn6/9oADAMBAAIRAxEAPwD618Q/siftJ/8ABXr9qT4kal8RfiR8Xv2bP2fPBesvongvwx4Wm/sTXvFHkgB9VuJypPkSFt0aurg5wFQxF5cP9lXU/in/AMEk/wDgsb4N/Zm174r+OvjX8G/jx4fv9Z8L3Hi66fU9c8L39jbzTTRGbOWhZLZsqiqhaZCEQrI8n6k/Ej4j6D8H/AGs+KvFGrWOheHPD1nLqGpajeSiKCzt41LPI7HoAoJr83f+CUngXXf+Ckf7fvjP9u3xho11oHhVtLk8B/BzS7qMLNNocc0hk1aQMNytM7zBBxgXE4+dRHIwB8jf8FU/+Cxfwb/bn/b+1P8AZ3+LPjzUvhr+zF8KdZkTxibTTb661D4j6pY3Cr/Z4+zQu9tax3Eb5kyC4i3KQxjZP0M/Zf8A+Dgf9jf44eO/DXww+Gfj+SfU7i3a00bSYPCep2UEcNtbvJ5aGS2SONUhhbAyBhQBzgUUUAf/2Q==',
        '/9j/4AAQSkZJRgABAQEAeAB4AAD/2wBDAAIBAQIBAQICAgICAgICAwUDAwMDAwYEBAMFBwYHBwcGBwcICQsJCAgKCAcHCg0KCgsMDAwMBwkODw0MDgsMDAz/2wBDAQICAgMDAwYDAwYMCAcIDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAz/wAARCAAZAA4DASIAAhEBAxEB/8QAHwAAAQUBAQEBAQEAAAAAAAAAAAECAwQFBgcICQoL/8QAtRAAAgEDAwIEAwUFBAQAAAF9AQIDAAQRBRIhMUEGE1FhByJxFDKBkaEII0KxwRVS0fAkM2JyggkKFhcYGRolJicoKSo0NTY3ODk6Q0RFRkdISUpTVFVWV1hZWmNkZWZnaGlqc3R1dnd4eXqDhIWGh4iJipKTlJWWl5iZmqKjpKWmp6ipqrKztLW2t7i5usLDxMXGx8jJytLT1NXW19jZ2uHi4+Tl5ufo6erx8vP09fb3+Pn6/8QAHwEAAwEBAQEBAQEBAQAAAAAAAAECAwQFBgcICQoL/8QAtREAAgECBAQDBAcFBAQAAQJ3AAECAxEEBSExBhJBUQdhcRMiMoEIFEKRobHBCSMzUvAVYnLRChYkNOEl8RcYGRomJygpKjU2Nzg5OkNERUZHSElKU1RVVldYWVpjZGVmZ2hpanN0dXZ3eHl6goOEhYaHiImKkpOUlZaXmJmaoqOkpaanqKmqsrO0tba3uLm6wsPExcbHyMnK0tPU1dbX2Nna4uPk5ebn6Onq8vP09fb3+Pn6/9oADAMBAAIRAxEAPwD7E8K/sSfHH/grj408afEH4y/Fn9oH9njwLZ+ILzR/AXw/8E6mfC99FYW0zQ/2hqTyQMZ5Lho/MjBUqiHKOUkAE/8AwS9+NHxS/ZZ/4Kg/Fn9jr4i/EPxF8X/Dfhnwha+OvBfiXXM3Wt21i0ttby2l5MAGlbzLhdpIbiNjuG8Iv2b+3PF8fJvgiB+zlN8LofiAL+My/wDCeR3j6c1l5cgkERtSGW48zySpcMm0OCMkEfB3/BBbxPZeFP2tf2g9A+N9lr+lftp6rNb6v47utaurZ7TV9JjCQ2Umk+QFjFiiPCCApYGWNWdgqBABP2b/APgvwP2I5PE3wo/b0n1vwD8UvCur3g03xSnhe5m0bx1pzzySQXFmLGBlURxtHFgqRtERaQymVUx/2BNA1b/grV/wVj+KP7Vln4c8TeDfgja+CLf4eeDdR1Szazu/F+27iunvoo2ORACjgHGCJYxkOkqJ99/tyf8AIoaH/wBfbf8AoBr2Twr/AMixpv8A16xf+gCgD//Z',
        '/9j/4AAQSkZJRgABAQEAeAB4AAD/2wBDAAIBAQIBAQICAgICAgICAwUDAwMDAwYEBAMFBwYHBwcGBwcICQsJCAgKCAcHCg0KCgsMDAwMBwkODw0MDgsMDAz/2wBDAQICAgMDAwYDAwYMCAcIDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAz/wAARCAAXABEDASIAAhEBAxEB/8QAHwAAAQUBAQEBAQEAAAAAAAAAAAECAwQFBgcICQoL/8QAtRAAAgEDAwIEAwUFBAQAAAF9AQIDAAQRBRIhMUEGE1FhByJxFDKBkaEII0KxwRVS0fAkM2JyggkKFhcYGRolJicoKSo0NTY3ODk6Q0RFRkdISUpTVFVWV1hZWmNkZWZnaGlqc3R1dnd4eXqDhIWGh4iJipKTlJWWl5iZmqKjpKWmp6ipqrKztLW2t7i5usLDxMXGx8jJytLT1NXW19jZ2uHi4+Tl5ufo6erx8vP09fb3+Pn6/8QAHwEAAwEBAQEBAQEBAQAAAAAAAAECAwQFBgcICQoL/8QAtREAAgECBAQDBAcFBAQAAQJ3AAECAxEEBSExBhJBUQdhcRMiMoEIFEKRobHBCSMzUvAVYnLRChYkNOEl8RcYGRomJygpKjU2Nzg5OkNERUZHSElKU1RVVldYWVpjZGVmZ2hpanN0dXZ3eHl6goOEhYaHiImKkpOUlZaXmJmaoqOkpaanqKmqsrO0tba3uLm6wsPExcbHyMnK0tPU1dbX2Nna4uPk5ebn6Onq8vP09fb3+Pn6/9oADAMBAAIRAxEAPwD9zP2gPjx4X/Zg+Cfij4h+NdSTSPCvg/TpdU1O7ZS5jijXJCqOXdjhVRQSzMqgEkV+M/8AwS3+Nnxj/aX/AODim1+JfxcsptAtPiV8HNQ8TeBvDsrkSaB4bk1Vbeyilj5CTyC1aZ8E7jMG+XPlp9GfH7ULX/gtp/wUDHwpS/02b9l79nPWUu/HsjXB8rx74mjDeVo2AwD21qSHmzuVnypGdjLf8Wavp1v/AMHVvg94rqyS0i/ZseJWWRRGhGv3eFGDgcdqAP0toqp/wkFh/wA/tp/39X/GigD4j+IX/Btd+xT8VPH2ueKNe+C/2/XfEmoT6pqNz/wl2vRfaLmeRpZZNiXqou52Y4UBRnAAHFfB/if/AIIVfsraf/wcHeHfgdD8Ldnwuv8A4IP4vn0X/hJdXPmaoNYuLYXH2g3Xnj9zGi7BJs4ztySaKKAP1k/4dsfBX/oS/wDyr3//AMfooooA/9k=',
        '/9j/4AAQSkZJRgABAQEAeAB4AAD/2wBDAAIBAQIBAQICAgICAgICAwUDAwMDAwYEBAMFBwYHBwcGBwcICQsJCAgKCAcHCg0KCgsMDAwMBwkODw0MDgsMDAz/2wBDAQICAgMDAwYDAwYMCAcIDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAz/wAARCAAYAA8DASIAAhEBAxEB/8QAHwAAAQUBAQEBAQEAAAAAAAAAAAECAwQFBgcICQoL/8QAtRAAAgEDAwIEAwUFBAQAAAF9AQIDAAQRBRIhMUEGE1FhByJxFDKBkaEII0KxwRVS0fAkM2JyggkKFhcYGRolJicoKSo0NTY3ODk6Q0RFRkdISUpTVFVWV1hZWmNkZWZnaGlqc3R1dnd4eXqDhIWGh4iJipKTlJWWl5iZmqKjpKWmp6ipqrKztLW2t7i5usLDxMXGx8jJytLT1NXW19jZ2uHi4+Tl5ufo6erx8vP09fb3+Pn6/8QAHwEAAwEBAQEBAQEBAQAAAAAAAAECAwQFBgcICQoL/8QAtREAAgECBAQDBAcFBAQAAQJ3AAECAxEEBSExBhJBUQdhcRMiMoEIFEKRobHBCSMzUvAVYnLRChYkNOEl8RcYGRomJygpKjU2Nzg5OkNERUZHSElKU1RVVldYWVpjZGVmZ2hpanN0dXZ3eHl6goOEhYaHiImKkpOUlZaXmJmaoqOkpaanqKmqsrO0tba3uLm6wsPExcbHyMnK0tPU1dbX2Nna4uPk5ebn6Onq8vP09fb3+Pn6/9oADAMBAAIRAxEAPwD3Pxx+wt8JP+Cw3/Baf9qG++N+iT3vgb4GaT4Z8GaFdLrFxpkL3csM93cDzY3QM6SSSKUydoZDxuGfsD9h/wD4IUfszfsD/Gi1+Jfwm8LappXiRLGaxivJPEN3fwvBMAHGyWRkOcDkDPHWvHfAf/Br38B9f+IHxK8UfGq88QfGTU/iJ4z1LxgLZ9Qv9A03SpLyZ5SiQWl2DJKN5UzO+WAGFUcV5d4E+GR/4IQf8Fh/hf8ADP4ea94nuf2dP2iNB1l4/Beo6lLfW/hLUtNga6aezaYu4D5UEbgW+0OXZ/LiCgG/8Dv+Cg/x4/4JG/8ACUfCn9oz4SftC/HqytfEWoX3g34jeC9HXxI+u6XcTNMiXpaePyZ42dgI/lCJtRU2Rq72P2XfhJ8VP+CuH/BUXwt+1D8TPhx4p+EHwa+D+iXuj+AvCfim2W21nXrq8ilgub25gK74oykuQN2CYYChYeYWKKAP/9k=',
        '/9j/4AAQSkZJRgABAQEAeAB4AAD/2wBDAAIBAQIBAQICAgICAgICAwUDAwMDAwYEBAMFBwYHBwcGBwcICQsJCAgKCAcHCg0KCgsMDAwMBwkODw0MDgsMDAz/2wBDAQICAgMDAwYDAwYMCAcIDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAz/wAARCAAYABADASIAAhEBAxEB/8QAHwAAAQUBAQEBAQEAAAAAAAAAAAECAwQFBgcICQoL/8QAtRAAAgEDAwIEAwUFBAQAAAF9AQIDAAQRBRIhMUEGE1FhByJxFDKBkaEII0KxwRVS0fAkM2JyggkKFhcYGRolJicoKSo0NTY3ODk6Q0RFRkdISUpTVFVWV1hZWmNkZWZnaGlqc3R1dnd4eXqDhIWGh4iJipKTlJWWl5iZmqKjpKWmp6ipqrKztLW2t7i5usLDxMXGx8jJytLT1NXW19jZ2uHi4+Tl5ufo6erx8vP09fb3+Pn6/8QAHwEAAwEBAQEBAQEBAQAAAAAAAAECAwQFBgcICQoL/8QAtREAAgECBAQDBAcFBAQAAQJ3AAECAxEEBSExBhJBUQdhcRMiMoEIFEKRobHBCSMzUvAVYnLRChYkNOEl8RcYGRomJygpKjU2Nzg5OkNERUZHSElKU1RVVldYWVpjZGVmZ2hpanN0dXZ3eHl6goOEhYaHiImKkpOUlZaXmJmaoqOkpaanqKmqsrO0tba3uLm6wsPExcbHyMnK0tPU1dbX2Nna4uPk5ebn6Onq8vP09fb3+Pn6/9oADAMBAAIRAxEAPwD9J/8Agth+394h/ZK+B+g/D74VxS6p+0B8ddQ/4RTwFZWxRpbGaTCzam6tnEVsrg7irKHeMsNgcr6p/wAEx/2AfD3/AATY/ZC8OfDXRXh1DVYVOoeJdbEZWbxFq8oU3N7IWJYlmAVdxJWOONcnbX5l/sm/san/AILvf8FQfjz8f/iZ4t8ZaP4H+DPi65+HPgTSfDGvvpsvl2sckVwXuIVWWOORJo5j5UgZ2uZELGNAG9I/aI+HHin/AIIMftdfA/xX8N/iP8RPGHwX+Nvjqz+H/ifwJ4x1+bWbfR5bv/UXunzS7pImQLM7KxYuVClirDywDO+G3jL4qf8ABAX9sD446fqfwR+JXxh/Z4+M/iy58deH9d+H2mDVtU0LULrHn2dzbbkAUBVUMzIMRIQZDIyx6XxJ1n4kf8F5/wBuH4JwaZ8IviR8Kv2cPgT4qTxrrPiDxxpn9i6n4m1W1P8AottZ2rl2MJyMucfJLLuMToiSFFAH/9k=',
        '/9j/4AAQSkZJRgABAQEAeAB4AAD/2wBDAAIBAQIBAQICAgICAgICAwUDAwMDAwYEBAMFBwYHBwcGBwcICQsJCAgKCAcHCg0KCgsMDAwMBwkODw0MDgsMDAz/2wBDAQICAgMDAwYDAwYMCAcIDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAz/wAARCAAXAA8DASIAAhEBAxEB/8QAHwAAAQUBAQEBAQEAAAAAAAAAAAECAwQFBgcICQoL/8QAtRAAAgEDAwIEAwUFBAQAAAF9AQIDAAQRBRIhMUEGE1FhByJxFDKBkaEII0KxwRVS0fAkM2JyggkKFhcYGRolJicoKSo0NTY3ODk6Q0RFRkdISUpTVFVWV1hZWmNkZWZnaGlqc3R1dnd4eXqDhIWGh4iJipKTlJWWl5iZmqKjpKWmp6ipqrKztLW2t7i5usLDxMXGx8jJytLT1NXW19jZ2uHi4+Tl5ufo6erx8vP09fb3+Pn6/8QAHwEAAwEBAQEBAQEBAQAAAAAAAAECAwQFBgcICQoL/8QAtREAAgECBAQDBAcFBAQAAQJ3AAECAxEEBSExBhJBUQdhcRMiMoEIFEKRobHBCSMzUvAVYnLRChYkNOEl8RcYGRomJygpKjU2Nzg5OkNERUZHSElKU1RVVldYWVpjZGVmZ2hpanN0dXZ3eHl6goOEhYaHiImKkpOUlZaXmJmaoqOkpaanqKmqsrO0tba3uLm6wsPExcbHyMnK0tPU1dbX2Nna4uPk5ebn6Onq8vP09fb3+Pn6/9oADAMBAAIRAxEAPwD72+JH/Br9+x18WfiJr/irXfAOv3Wt+JtRuNV1CdfFOoRLNcTytLKwVZQqgu7HAAAzgCvgrx//AMG/f7Nn7Rv/AAVhi+A/wf8ADmreHvBvwl8PPrXxU10a5dX8ovb2MjTNLh82YhJMBp3IXlQVyCuD+0/7b37VWifsQfsk/ED4seIWiOmeBtGm1HyXlWL7bOBst7ZWYgb5p2iiUZ5aRQOTXgH/AAQa/Zq1f4L/ALCOneOPGjT3XxS+PN7J8SPGN5cRyRzvdahiWGFkckoIbcxJs4Ctv4BJoA8X/wCDkLWZvjjqH7L/AOy9atcND8f/AImWY8QQRsIhc6Hp8kUl4hlJGza08Mo2EOTb4XOdrfptZWcWnWcVvBGkUECCOONRhUUDAAHoAKKKAP/Z',
        '/9j/4AAQSkZJRgABAQEAeAB4AAD/2wBDAAIBAQIBAQICAgICAgICAwUDAwMDAwYEBAMFBwYHBwcGBwcICQsJCAgKCAcHCg0KCgsMDAwMBwkODw0MDgsMDAz/2wBDAQICAgMDAwYDAwYMCAcIDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAz/wAARCAAYAA8DASIAAhEBAxEB/8QAHwAAAQUBAQEBAQEAAAAAAAAAAAECAwQFBgcICQoL/8QAtRAAAgEDAwIEAwUFBAQAAAF9AQIDAAQRBRIhMUEGE1FhByJxFDKBkaEII0KxwRVS0fAkM2JyggkKFhcYGRolJicoKSo0NTY3ODk6Q0RFRkdISUpTVFVWV1hZWmNkZWZnaGlqc3R1dnd4eXqDhIWGh4iJipKTlJWWl5iZmqKjpKWmp6ipqrKztLW2t7i5usLDxMXGx8jJytLT1NXW19jZ2uHi4+Tl5ufo6erx8vP09fb3+Pn6/8QAHwEAAwEBAQEBAQEBAQAAAAAAAAECAwQFBgcICQoL/8QAtREAAgECBAQDBAcFBAQAAQJ3AAECAxEEBSExBhJBUQdhcRMiMoEIFEKRobHBCSMzUvAVYnLRChYkNOEl8RcYGRomJygpKjU2Nzg5OkNERUZHSElKU1RVVldYWVpjZGVmZ2hpanN0dXZ3eHl6goOEhYaHiImKkpOUlZaXmJmaoqOkpaanqKmqsrO0tba3uLm6wsPExcbHyMnK0tPU1dbX2Nna4uPk5ebn6Onq8vP09fb3+Pn6/9oADAMBAAIRAxEAPwD6A/Zo/Y38C/8ABdz9rD9pvxf+0i2veMdN+EXxNvvh14T8Dx6/d6dpnh6z0/5PtXl20kTtNcMzlnY5yjDkABdT9mf4Nx/8Emv+C7/gj9n34Ma94jvvgn8XfAt9r+seBtQ1mTULbwHcW3ntFqFuZi8qR3EsXlFdwLvcSMzOEiVPHPg78GfjR/wVi/4KJfGL9oD9lDxxp/7MPwzTUH8Iah4rtRJq0vxOvbMqDfmxLfZJIwj/ACTAgjd1dnk8vvv2J28Vf8Em/wDgq4fC37T1vpPxI8b/AB80i+bwz8chfzvd3ltYItxLpEllKWSyRB87fZykbE2wIlI3IAUP2Gvjr8TP+Deu38e/An4ifs+fGv4k/C7/AIS6/wBZ+Hfir4daNHryS6dOVItrhN8RjkXaHPmuZC8soAKIjN0Xh/Q/iJ/wXj/4KSfC34ka18HviD8F/gN8BNL1mO3l8a2LadrPifUdStkgdIYM/LGgWFt4Lp+5cFtzhVKKAP/Z',
        '/9j/4AAQSkZJRgABAQEAeAB4AAD/2wBDAAIBAQIBAQICAgICAgICAwUDAwMDAwYEBAMFBwYHBwcGBwcICQsJCAgKCAcHCg0KCgsMDAwMBwkODw0MDgsMDAz/2wBDAQICAgMDAwYDAwYMCAcIDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAz/wAARCAAYABADASIAAhEBAxEB/8QAHwAAAQUBAQEBAQEAAAAAAAAAAAECAwQFBgcICQoL/8QAtRAAAgEDAwIEAwUFBAQAAAF9AQIDAAQRBRIhMUEGE1FhByJxFDKBkaEII0KxwRVS0fAkM2JyggkKFhcYGRolJicoKSo0NTY3ODk6Q0RFRkdISUpTVFVWV1hZWmNkZWZnaGlqc3R1dnd4eXqDhIWGh4iJipKTlJWWl5iZmqKjpKWmp6ipqrKztLW2t7i5usLDxMXGx8jJytLT1NXW19jZ2uHi4+Tl5ufo6erx8vP09fb3+Pn6/8QAHwEAAwEBAQEBAQEBAQAAAAAAAAECAwQFBgcICQoL/8QAtREAAgECBAQDBAcFBAQAAQJ3AAECAxEEBSExBhJBUQdhcRMiMoEIFEKRobHBCSMzUvAVYnLRChYkNOEl8RcYGRomJygpKjU2Nzg5OkNERUZHSElKU1RVVldYWVpjZGVmZ2hpanN0dXZ3eHl6goOEhYaHiImKkpOUlZaXmJmaoqOkpaanqKmqsrO0tba3uLm6wsPExcbHyMnK0tPU1dbX2Nna4uPk5ebn6Onq8vP09fb3+Pn6/9oADAMBAAIRAxEAPwD3jwB8Az/wcI/8FCv2gZ/jJqniaT9nb4AeJn8B+GvA9jqc2m2ur6nbsy3d5dGBwzkNGHQ7g22dFG3YwbnP2jP2X4f+DeX9vL9nTxR+z5q+taP8GfjZ4zt/BXjbwDqGrzX1nLPOUiivIPO3yB1jd3Lb9we3iXJSRlHr2kfszftHf8Ekv+CgXxc8d/B34aR/Hj4D/H3WT4l1jw5Ya1b6br3hbWJC7zTxfanEcsTu8nAPKtGP3flZk8L/AGevjD8Tf+C9n/BbPw5d+Nfh+fhz8Jf2OdQu7+70OTUo9V8/xAJvLt0uJoiITcrJEkm1BIsItJk3ky7iAfpp/wAFT/2pfFX7I/7FfivxH8P/AAvrfjL4jajGujeFNK0zT5b2R9Ruf3cU8iRj5YYcmV2YquI9u4FhWF/wR3/4J6p/wTc/Yj8P+DNSlXUvHusu/iDxvqxm+0SaprNyA1wxmIDSrHxEjNyVjBPLEkooA//Z']
    pre_info_json['page_num'] = sub_json_data['page_num']
    # pre_info_json['paper_num_count_flag'] = sub_json_data['paper_num_count_flag']
    pre_info_json['submark'] = [0, 1]  # 缺考标记
    # pre_info_json['sub_mark_num_vec'] = sub_json_data['sub_mark_num_vec']
    pre_info_json['sub_mark_num'] = sub_json_data['sub_mark_num']
    pre_info_json['question_type'] = 0
    pre_info_json['barcode_topleft_width'] = 409
    pre_info_json['barcode_topleft_height'] = 72
    pre_info_json['barcode_width'] = 268
    pre_info_json['barcode_height'] = 86
    pre_info_json['paper_mark_width'] = 22
    pre_info_json['paper_mark_height'] = 22
    pre_info_json['paper_mark_to_mark'] = 698
    pre_info_json['question_mark_width'] = 7
    pre_info_json['question_mark_height'] = 23
    pre_info_json['absent_topleft_width'] = 43
    pre_info_json['absent_topleft_height'] = 135
    pre_info_json['absent_rightdown_width'] = 63
    pre_info_json['absent_rightdown_height'] = 124
    return pre_info_json


def choice_info(sub_json_data):
    choice_info_json = {}
    choice_answer_inf = {}
    choice_answer_inf["choice_answer_cnt"] = sub_json_data['choice_answer_cnt']
    choice_answer_inf["choice_answer_pattern"] = sub_json_data['choice_answer_pattern']
    choice_info_json['choice_answer_inf'] = choice_answer_inf
    choice_col_num = sub_json_data['choice_col_num']
    choice_info_json['choice_col_num'] = choice_col_num
    choice_row_num = sub_json_data['choice_row_num']
    choice_line_number = sub_json_data['choice_line_number']
    choice_info_json["choice_row_num"] = choice_row_num
    choice_info_json["choice_line_number"] = choice_line_number
    # choice_info_json['col_to_col_gap'] = 420
    choice_info_json['horizon_gap'] = 9
    choice_info_json['choice_row_vertical_interval'] = 50
    left_to_mark = [36, 238, 440]
    # left_to_mark = left_to_mark[0:choice_col_num[0]]
    choice_info_json['left_to_mark'] = left_to_mark
    choice_info_json['horizon_choice_mark_to_mark'] = 621
    choice_info_json['question_type'] = 1
    choice_info_json["qus_range"] = sub_json_data['qus_range_choice']
    choice_info_json['rect_height'] = 13
    choice_info_json['rect_thick'] = 2
    choice_info_json['rect_width'] = 16
    # choice_info_json['scale'] = 0.54474708171206221
    choice_info_json['submark'] = sub_json_data['sub_mark_choice']
    top_to_mark = [13]
    # top_to_mark_60_ = [13, 159, 305, 451, 597, 743]
    # if choice_row_num[0] <= 4:
    #     top_to_mark = top_to_mark_60[0:choice_row_num[0]]
    # else:
    #     top_to_mark = []
    #     round, rest = divmod(choice_row_num[0] - 4, 6)
    #     for n in range(round):
    #         top_to_mark = top_to_mark_60.append(top_to_mark_60_)
    #     top_to_mark.append(top_to_mark_60_[0:rest])
    choice_info_json['top_to_mark'] = top_to_mark
    choice_info_json['vertical_gap'] = 8
    return choice_info_json


def fillin_info(sub_json_data):
    fillin_info = {}
    blankfillin_answer_inf = {
        "blankfillin_answer_cnt": 4,
        "blankfillin_answer_pattern": [4]
    }
    fillin_info['blankfillin_answer_inf'] = blankfillin_answer_inf
    fillin_info['question_type'] = 2
    fillin_info["qus_range"] = sub_json_data['qus_range_fillin']
    fillin_info['submark'] = sub_json_data['sub_mark_fillin']
    return fillin_info


def eassy_info(sub_json_data):
    eassy_info = {}
    eassy_answer_inf = {
        "eassy_answer_cnt": None,
        "eassy_answer_pattern": [2]
    }
    eassy_info['eassy_answer_inf'] = eassy_answer_inf
    eassy_info['question_type'] = 3
    eassy_info["qus_range"] = sub_json_data['qus_range_eassy']
    eassy_info["eassy_score"] = sub_json_data['eassy_score']
    eassy_info["score_distance_h"] = 15
    eassy_info["score_distance_v"] = 12
    eassy_info["score_width"] = [29, 24, 25, 28, 31]
    eassy_info["score_height"] = 22

    eassy_info['submark'] = sub_json_data['sub_mark_eassy']
    return eassy_info


def get_json(sub_json_data):
    json_data = []
    # 总体数据
    pre_info_json = pre_info(sub_json_data)
    # 选择题数据
    choice_info_json = choice_info(sub_json_data)
    # 填空题数据
    fillin_info_json = fillin_info(sub_json_data)
    # 简答题数据
    eassy_info_json = eassy_info(sub_json_data)
    json_data.append(pre_info_json)
    json_data.append(choice_info_json)
    json_data.append(fillin_info_json)
    json_data.append(eassy_info_json)
    return json_data


if __name__ == '__main__':
    path = r'C:\Users\j20687\Desktop\paperJson.json'
    with open(path, 'r', encoding='utf-8') as f:
        paper_data = json.load(f)
        paper_data = paper_data.get("data")
        print(paper_data)
        get4answer_sheet_data(paper_data)
