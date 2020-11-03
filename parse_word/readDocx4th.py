import docx
import re
from bs4 import BeautifulSoup
from pydocx import PyDocX
import base64
import os
import time
import xml.etree.cElementTree as ET

def read4word(file):
	'''
	:param file: 为传入的.docx文件对象,以二进制格式打开
	:return:
	'''
	Exam = {}
	proxy = []
	doc = docx.Document(file)
	for para in doc.paragraphs:
		proxy.append(para._element.xml)
	ommls = getMathml(proxy)
	# end_para_no, QueStyle_para_no, questions, exam_name, subject, style = read_doc4para_no(doc)
	# Questions = word2html(file, end_para_no, QueStyle_para_no, questions, style)
	# Exam['exam_name'] = exam_name
	# Exam['subject'] = subject
	# Exam['data'] = Questions
	return Exam


def getMathml(proxy):
	'''
	遍历xml，找到m:oMath
	:param element: xml文件根节点
	:return:
	'''
	i = 1
	ommls = []
	re_math = re.compile(r'<m:oMath>.*?</m:oMath>', re.S)
	for p in proxy:
		string = str(p)
		omml = re_math.findall(string)
		if omml != []:
			inner_str = ""
			for ss in omml:
				inner_str += str(ss).replace('\n', '')
				# .replace("±", "&#x00B1;").replace("∑", "&sum;")
			ommls.append(inner_str)
		i = i+1
	return ommls


def read_doc4para_no(doc):
	style = {}
	i = 1
	questions = {}
	question = []
	option = []
	answer = []
	analysis = []
	end_para_no = []
	QueStyle_para_no = []
	exam_name = ""
	subject = ""
	re_style = re.compile('.*?、(.*?题)', re.S)
	style_no = 0

	para = doc.paragraphs
	for p in range(len(para)):
		if "【试卷】" in para[p].text:
			exam_name = para[p].text.replace("【试卷】", "")
		if "【科目】" in para[p].text:
			subject = para[p].text.replace("【科目】", "")
		if re_style.findall(para[p].text):
			QueStyle_para_no.append(i)
			style[str(style_no)] = re_style.findall(para[p].text)[0]
			style_no += 1
		if "【题文】" in para[p].text:
			question.append(i)
		if "【选项】" in para[p].text:
			option.append(i)
		if "【答案】" in para[p].text:
			answer.append(i)
		if "【解析】" in para[p].text:
			analysis.append(i)
		if para[p].text == "【结束】":
			end_para_no.append(i)
		i = i+1
	questions['question'] = question
	questions['option'] = option
	questions['answer'] = answer
	questions['analysis'] = analysis
	QueStyle_para_no.append(-1)
	return end_para_no, QueStyle_para_no, questions, exam_name, subject, style


def word2html(file, end_para_no, QueStyle_para_no, questions, style):
	Questions = []
	html = PyDocX.to_html(file)
	soup = BeautifulSoup(html, 'lxml')
	title_no = 0  # 题目序号

	# for img in soup.findAll('img'):
	# 	reg = re.compile('data.*?/(.*?);', re.S)
	# 	style_img = reg.findall(img['src'])[0]
	# 	strg = img['src'].replace("data:image/wmf;base64,", "").replace("data:image/jpeg;base64,", "")
	# 	byte = base64.urlsafe_b64decode(strg)
	# 	t0 = int(round(time.time() * 1000))
	# 	tmp_path = '/tmp/%d.%s' % (t0, str(style_img))
	# 	with open(tmp_path, 'wb') as file:
	# 		file.write(byte)
	# 	if style_img == 'wmf':
	# 		t1 = int(round(time.time() * 1000))
	# 		png_path = '/tmp/%d.png' % t1
	# 		os.system('convert %s %s' % (tmp_path, png_path))
	# 		f = open(png_path, 'rb')
	# 		url = put(f)
	# 		f.close()
	# 		img['src'] = url
	# 		os.remove(png_path)
	# 		os.remove(tmp_path)
	# 	else:
	# 		f = open(tmp_path, 'rb')
	# 		url = put(f)
	# 		f.close()
	# 		img['src'] = url
	# 		os.remove(tmp_path)

	patt = re.compile('>(.*)(<img.*?>)(.*?)<|>(.*?)<')
	re_a = re.compile(r'A\.(.*?)B\.|A\.(.*?)</p>', re.S | re.M)
	re_b = re.compile(r'B\.(.*?)C\.|B\.(.*?)</p>', re.S | re.M)
	re_c = re.compile(r'C\.(.*?)D\.|C\.(.*?)</p>', re.S | re.M)
	re_d = re.compile(r'D\.(.*?)E\.|D\.(.*?)</p>', re.S | re.M)
	re_e = re.compile(r'E\.(.*?)F\.|E\.(.*?)</p>', re.S | re.M)
	re_f = re.compile(r'F\.(.*?)</p>', re.S | re.M)
	re_newline = re.compile('.*?A\..*?|.*?B\..*?|.*?C\..*?|.*?D\..*?|.*?E\..*?|.*?F\..*?')

	for p in range(0, (len(QueStyle_para_no)-1)):
		i = 1
		question = ""
		opt_line = 0
		option = []
		answer = ""
		analysis = ""
		stop = QueStyle_para_no[p + 1] if QueStyle_para_no[p + 1] != -1 else end_para_no[-1]+1
		for item in soup.find_all('p'):
			para_str = ""
			if QueStyle_para_no[p] <= i < stop:
				group = patt.findall(str(item))
				for g in group:
					string = g[0]+g[1]+g[2].replace(' ', '&nbsp;')+g[3].replace(' ', '&nbsp;')
					span_str = str(string).replace('</span>', '').replace('<span class="pydocx-left">', '')
					if span_str != "":
						span_str = '<span>' + span_str + '</span>'
						para_str += span_str
				para_str = "<p>" + para_str + "</p>"
				para_str = para_str.replace("\n", "")
				if style['%s' % str(p)] == "单选题" or style['%s' % str(p)] == "多选题" or style['%s' % str(p)] == "判断题":
					if questions['option'][title_no] <= i < questions['answer'][title_no]:
						opt_line = questions['answer'][title_no] - questions['option'][title_no]
						aa = {}
						bb = {}
						cc = {}
						dd = {}
						ee = {}
						ff = {}
						opt = para_str.replace("【选项】", "")
						A = re_a.findall(opt)
						B = re_b.findall(opt)
						C = re_c.findall(opt)
						D = re_d.findall(opt)
						E = re_e.findall(opt)
						F = re_f.findall(opt)
						newline = re_newline.findall(opt)
						if A != []:
							aa['key'] = 'A'
							aa['value'] = str(A[0][0] + A[0][1]).replace('<span class="pydocx-tab">', '')\
								.replace('</span>', '').replace('<span>', '').replace('&nbsp;', ' ')
							option.append(aa)
						if B != []:
							bb['key'] = 'B'
							bb['value'] = str(B[0][0] + B[0][1]).replace('<span class="pydocx-tab">', '')\
								.replace('</span>', '').replace('<span>', '').replace('&nbsp;', ' ')
							option.append(bb)
						if C != []:
							cc['key'] = 'C'
							cc['value'] = str(C[0][0] + C[0][1]).replace('<span class="pydocx-tab">', '')\
								.replace('</span>', '').replace('<span>', '').replace('&nbsp;', ' ')
							option.append(cc)
						if D != []:
							dd['key'] = 'D'
							dd['value'] = str(D[0][0] + D[0][1]).replace('<span class="pydocx-tab">', '')\
								.replace('</span>', '').replace('<span>', '').replace('&nbsp;', ' ')
							option.append(dd)
						if E != []:
							ee['key'] = 'E'
							ee['value'] = str(E[0][0] + E[0][1]).replace('<span class="pydocx-tab">', '')\
								.replace('</span>', '').replace('<span>', '').replace('&nbsp;', ' ')
							option.append(ee)
						if F != []:
							ff['key'] = 'F'
							ff['value'] = str(F[0][0] + F[0][1]).replace('<span class="pydocx-tab">', '')\
								.replace('</span>', '').replace('<span>', '').replace('&nbsp;', ' ')
							option.append(ff)
						if newline == [] and opt_line > 4:  # 如果选项换行
							remap = option.pop()
							remap['value'] += opt
							option.append(remap)
					end = questions['option'][title_no]
				else:
					end = questions['answer'][title_no]
				if questions['question'][title_no] <= i < end:
					question += para_str.replace("【题文】", "")
				if questions['answer'][title_no] <= i < questions['analysis'][title_no]:
					answer += para_str.replace("【答案】", "")
				if questions['analysis'][title_no] <= i < end_para_no[title_no]:
					analysis += para_str.replace("【解析】", "")
				if "【结束】" in para_str:
					mm = {}
					mm['type'] = style['%s' % str(p)]
					mm['question'] = question
					mm['opt_line'] = opt_line
					mm['option'] = option
					mm['answer'] = answer
					mm['analysis'] = analysis
					Questions.append(mm)
					question = ""
					opt_line = 0
					option = []
					answer = ""
					analysis = ""
					title_no += 1
			i += 1
	return Questions


def put(byte):
	url = "--------------------"
	return url


if __name__ == '__main__':
	path = 'C:\\Users\\j20687\\Desktop\\demo1.docx'
	Exam = read4word(path)
	# print(Exam['exam_name'])
	# print(Exam['subject'])
	# for q in Exam['data']:
	# 	print(q)






