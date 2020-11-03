import docx
import re
from bs4 import BeautifulSoup
from pydocx import PyDocX
import base64
import os
import time
import lxml.etree as ET
from io import BytesIO


'''
OMML2MML_XSL是OMML转MathML的XSL文件，
一般在'C:\Program Files (x86)\Microsoft Office\Office15\OMML2MML.XSL'可找到，可复制到项目目录里
'''
OMML2MML_XSL = r"C:\Program Files (x86)\Microsoft Office\Office15\OMML2MML.XSL"

def read4word(file):
	'''
	:param file: 为传入的.docx文件对象,以二进制格式打开
	:return:
	'''
	Exam = {}
	proxy = []
	doc = docx.Document(file)
	for para in doc.paragraphs:
		proxy.append(para._element.xml)
	mmls_in_para = getMathml(proxy)
	images_in_para, images_in_table = get_img(file)
	table_html, table_para = get_table(doc, images_in_table)
	paragraphs = get_para_html(proxy, mmls_in_para, images_in_para, table_html, table_para)
	end_para_no, QueStyle_para_no, questions, exam_name, subject, style = read_doc4para_no(doc, table_para)
	Questions = word2html(end_para_no, QueStyle_para_no, questions, style, paragraphs)
	Exam['exam_name'] = exam_name
	Exam['subject'] = subject
	Exam['data'] = Questions
	return Exam


def getMathml(proxy):
	'''
	遍历xml，找到m:oMath
	:param element: xml文件根节点
	:return:
	'''

	mmls = []
	ommls = []
	re_math = re.compile(r'<m:oMath>(.*?)</m:oMath>', re.S)
	# <w:t xml:space=\"preserve\"> </w:t>
	for xml in proxy:
		omml = re_math.findall(xml)
		if omml != []:
			for ss in omml:
				string = (str(ss))
				# string = string.replace('<m:oMath>\n', '').replace('</m:oMath>', '')
				# string = string.replace('<m:r>\n', '', 1)
				str1 = '<xml-fragment xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml" xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk" xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml" xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:w10="urn:schemas-microsoft-com:office:word" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas">'
				string = str1 + string + '</xml-fragment>'
				ommls.append(bytes(string, encoding='utf-8'))

	for omml in ommls:
		dom = ET.parse(BytesIO(omml))
		xslt = ET.parse(OMML2MML_XSL)
		transform = ET.XSLT(xslt)
		newdom = transform(dom)
		string = str(ET.tostring(newdom, pretty_print=True), encoding='utf-8')
		string = string.replace("xmlns:m=\"http://schemas.openxmlformats.org/officeDocument/2006/math\"", "")
		string = string.replace("xmlns:mml", "xmlns")
		string = string.replace("mml:", "").replace("\n", "")
		mmls.append(string)
	return mmls


def get_table(doc, images_in_table):
	doc_xml = doc._element.xml
	root = ET.fromstring(doc_xml)
	table_xml = []
	p_tab = []  # 表的单元格数
	table_html = {}
	table_para = []
	i = 0
	j = 0
	for p in root.iter():
		if p.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p':
			i += 1
		if p.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl':
			table_para.append(i)
	for table in doc.tables:
		tab = 0
		xml = table._element.xml
		table_xml.append(xml)
		for row in table.rows:
			for cell in row.cells:
				tab += 1
		p_tab.append(tab)
	for n in range(len(p_tab)):
		if n >= 1:
			table_para[n] = table_para[n] - p_tab[n-1] + n
	table_para.append(-1)
	mmls_in_table = getMathml(table_xml)
	for tbl in root.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl'):
		str_table = ''
		result = []
		result.append('<table border="1">')
		getTableText(tbl, result, mmls_in_table, images_in_table)
		result.append('</tr>')
		result.append('</table>')
		result.remove(result[1])
		for string in result:
			str_table += string
		table_html[table_para[j]] = str_table
		j += 1
	return table_html, table_para


def getTableText(p, result, mmls_in_table, images_in_table):
	if len(list(p)) == 0:
		if p.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t':
			tmp = p.text
			result.append(tmp)
			result.append('</td>')
	else:
		for child in list(p):
			if child.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr':
				result.append('</tr>')
				result.append('<tr>')
			if child.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc':
				result.append('<td>')
			if child.tag == '{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath':
				result.append(mmls_in_table[0])
				result.append('</td>')
				mmls_in_table.pop(0)
			elif child.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing':
				result.append(images_in_table[0])
				result.append('</td>')
				images_in_table.pop(0)
			elif child.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object':
				result.append(images_in_table[0])
				result.append('</td>')
				images_in_table.pop(0)
			else:
				getTableText(child, result, mmls_in_table, images_in_table)


def get_para_html(proxy, mmls, images_in_para, table_html, table_para):
	'''
	返回word段落带p标签的html格式
	:param proxy:
	:param paragraphs:
	:return:
	'''
	results = []
	paragraphs = []  # 返回word段落带p标签的html格式
	i = 1
	para = 0
	for p in proxy:
		result = []
		root = ET.fromstring(p)
		getNodeText(root, result, mmls, images_in_para)
		results.append(result)
	for s in results:
		str_para = ''
		for ss in s:
			# string = str(ss).replace(' ', '&nbsp;')
			str_para += str(ss)
		str_para = '<p>' + str_para + '</p>'
		paragraphs.append(str_para)
		if i == table_para[para]:
			paragraphs.append(table_html[table_para[para]])
			para += 1
			i += 1
		i += 1
	return paragraphs


def getNodeText(ele, result, mmls, img_in_docx):
	if len(list(ele)) == 0:
		walkdata(ele, result)
	else:
		for child in list(ele):
			if child.tag == '{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath':
				result.append(mmls[0])
				mmls.pop(0)
			elif child.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing':
				result.append(img_in_docx[0])
				img_in_docx.pop(0)
			elif child.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object':  # OLE对象，比如mathtype公式
				result.append(img_in_docx[0])
				img_in_docx.pop(0)
			else:
				getNodeText(child, result, mmls, img_in_docx)


def walkdata(child, result):
	if child.tag =='{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t':
		tmp = str(child.text).replace(' ', '&nbsp;')
		result.append(tmp)


def read_doc4para_no(doc, table_para):
	style = {}
	i = 1
	no = 0
	questions = {}  # 所有题的题文、选项、答案、解析的段落号
	question = []
	combine_ques = []  # 组合题的大题文
	option = []
	answer = []
	analysis = []
	end_para_no = []
	QueStyle_para_no = []
	exam_name = ""
	subject = ""
	re_style = re.compile('^\S{1,2}、(.*?题)')
	style_no = 0
	string = ["主观题", "英语题", "填空题"]


	para = doc.paragraphs
	for p in range(len(para)):
		if "【试卷】" in para[p].text:
			exam_name = para[p].text.replace("【试卷】", "")
		if "【科目】" in para[p].text:
			subject = para[p].text.replace("【科目】", "")
		if re_style.findall(para[p].text):
			QueStyle_para_no.append(i)
			# style[str(style_no)] = re_style.findall(para[p].text)[0]
			style[str(style_no)] = para[p].text
			style_no += 1
		# if re_style.findall(para[p].text) and str(re_style.findall(para[p].text)[0]) in string:
		# 	option.append(0)
		if "【大题】" in para[p].text:
			combine_ques.append(i)
		if "【题文】" in para[p].text:
			question.append(i)
		if "【选项】" in para[p].text:
			option.append(i)
		if "【答案】" in para[p].text:
			answer.append(i)
		if "【解析】" in para[p].text:
			analysis.append(i)
		if para[p].text == "【结束】":
			end_para_no.append(i)
		if i == table_para[no]:
			i += 1
			no += 1
		i = i+1
	add_zero = 0
	for ll in range(len(style.keys())):
		if '填空题' in style[str(ll)] or '主观题' in style[str(ll)] or '英语题' in style[str(ll)]:
			j = 0
			for number in end_para_no:
				if QueStyle_para_no[ll] < number < QueStyle_para_no[ll+1]:
					add_zero += 1
			for i in range(len(answer)):
				if answer[i] > QueStyle_para_no[ll]:
					j = i
					break
			for i in range(0, add_zero):
				option.insert(j, 0)
	combine_ques.append(-1)
	questions['combine_ques'] = combine_ques
	questions['question'] = question
	questions['option'] = option
	questions['answer'] = answer
	questions['analysis'] = analysis
	QueStyle_para_no.append(-1)
	print(questions)
	return end_para_no, QueStyle_para_no, questions, exam_name, subject, style


def get_img(file):
	html = PyDocX.to_html(file)
	soup = BeautifulSoup(html, 'lxml')
	images_in_para = []
	images_in_table = []

	for img in soup.find_all('img'):
		if img.find_parents('table') != []:
			img['src'] = "$$$$$$$$$$$$$$$$"
			images_in_table.append(str(img))
		else:
			img['src'] = "aaaaaaaaaaaa"
			images_in_para.append(str(img))
		# reg = re.compile('data.*?/(.*?);', re.S)
		# style_img = reg.findall(img['src'])[0]
		# strg = img['src'].replace("data:image/wmf;base64,", "").replace("data:image/jpeg;base64,", "")
		# byte = base64.urlsafe_b64decode(strg)
		# t0 = int(round(time.time() * 1000))
		# tmp_path = '/tmp/%d.%s' % (t0, str(style_img))
		# with open(tmp_path, 'wb') as file:
		# 	file.write(byte)
		# if style_img == 'wmf':
		# 	t1 = int(round(time.time() * 1000))
		# 	png_path = '/tmp/%d.png' % t1
		# 	os.system('convert %s %s' % (tmp_path, png_path))
		# 	f = open(png_path, 'rb')
		# 	url = put(f)
		# 	f.close()
		# 	img['src'] = url
		# 	if img.find_parents('table') != []:
		# 		images_in_table.append(img)
		# 	else:
		# 		images_in_para.append(img)
		# 	os.remove(png_path)
		# 	os.remove(tmp_path)
		# else:
		# 	f = open(tmp_path, 'rb')
		# 	url = put(f)
		# 	f.close()
		# 	img['src'] = url
		# 	if img.find_parents('table') != []:
		# 		images_in_table.append(img)
		# 	else:
		# 		images_in_para.append(img)
		# 	os.remove(tmp_path)
	return images_in_para, images_in_table


def word2html(end_para_no, QueStyle_para_no, questions, style, paragraphs):
	Questions = []
	title_no = 0  # 题目序号
	string_style = ["单选题", "多选题", "判断题"]

	for p in range(0, (len(QueStyle_para_no)-1)):
		i = 1
		question = ""
		opt_line = 0
		option = []
		answer = ""
		analysis = ""
		stop = QueStyle_para_no[p + 1] if QueStyle_para_no[p + 1] != -1 else end_para_no[-1]+1
		style_p = style[str(p)]
		if '解答题' in style_p or '组合题' in style_p:
			combine_question(questions, Questions, title_no, paragraphs, end_para_no, stop)
		else:
			for item in paragraphs:
				para_str = ""
				if QueStyle_para_no[p] <= i < stop:
					para_str += str(item)
					para_str = para_str.replace("\n", "")
					if style['%s' % str(p)] in string_style:
						if questions['option'][title_no] <= i < questions['answer'][title_no]:
							opt_line = questions['answer'][title_no] - questions['option'][title_no]
							get_option(option, title_no, questions, para_str, i)
						end = questions['option'][title_no]
					else:
						end = questions['answer'][title_no]
					if questions['question'][title_no] <= i < end:
						question += para_str.replace("【题文】", "")
					if questions['answer'][title_no] <= i < questions['analysis'][title_no]:
						answer += para_str.replace("【答案】", "")
					if questions['analysis'][title_no] <= i < end_para_no[title_no]:
						analysis += para_str.replace("【解析】", "")
					if "【结束】" in para_str:
						mm = {}
						mm['type'] = style['%s' % str(p)]
						mm['question'] = question
						mm['opt_line'] = opt_line
						mm['option'] = option
						mm['answer'] = answer
						mm['analysis'] = analysis
						Questions.append(mm)
						question = ""
						opt_line = 0
						option = []
						answer = ""
						analysis = ""
						title_no += 1
				i += 1
	return Questions


def get_option(option, title_no, questions, para_str, i):
	re_a = re.compile(r'A\.(.*?)B\.|A\.(.*?)</p>|A、(.*?)B、|A、(.*?)</p>|A\.(.*?)B、|A、(.*?)B\.', re.S | re.M)
	re_b = re.compile(r'B\.(.*?)C\.|B\.(.*?)</p>|B、(.*?)C、|B、(.*?)</p>|B\.(.*?)C、|B、(.*?)C\.', re.S | re.M)
	re_c = re.compile(r'C\.(.*?)D\.|C\.(.*?)</p>|C、(.*?)D、|C、(.*?)</p>|C\.(.*?)D、|C、(.*?)D\.', re.S | re.M)
	re_d = re.compile(r'D\.(.*?)E\.|D\.(.*?)</p>|D、(.*?)E、|D、(.*?)</p>|D\.(.*?)E、|D、(.*?)E\.', re.S | re.M)
	re_e = re.compile(r'E\.(.*?)F\.|E\.(.*?)</p>|E、(.*?)F、|E、(.*?)</p>|E\.(.*?)F、|E、(.*?)F\.', re.S | re.M)
	re_f = re.compile(r'F\.(.*?)</p>|F、(.*?)</p>', re.S | re.M)
	re_newline = re.compile('.*?A\..*?|.*?B\..*?|.*?C\..*?|.*?D\..*?|.*?E\..*?|.*?F\..*?|.*?A、.*?|.*?B、.*?|.*?C、.*?|.*?D、.*?|.*?E、.*?|.*?F、.*?')
	if questions['option'][title_no] <= i < questions['answer'][title_no]:
		opt_line = questions['answer'][title_no] - questions['option'][title_no]
		aa = {}
		bb = {}
		cc = {}
		dd = {}
		ee = {}
		ff = {}
		opt = para_str.replace("【选项】", "")
		if opt != "":
			A = re_a.findall(opt)
			B = re_b.findall(opt)
			C = re_c.findall(opt)
			D = re_d.findall(opt)
			E = re_e.findall(opt)
			F = re_f.findall(opt)
			newline = re_newline.findall(opt)
			if A != []:
				aa['key'] = 'A'
				aa['value'] = str(A[0][0] + A[0][1] + A[0][2] + A[0][3] + A[0][4] + A[0][5]).replace('&nbsp;', '')
				option.append(aa)
			if B != []:
				bb['key'] = 'B'
				bb['value'] = str(B[0][0] + B[0][1] + B[0][2] + B[0][3] + B[0][4] + B[0][5]).replace('&nbsp;', '')
				option.append(bb)
			if C != []:
				cc['key'] = 'C'
				cc['value'] = str(C[0][0] + C[0][1] + C[0][2] + C[0][3] + C[0][4] + C[0][5]).replace('&nbsp;', '')
				option.append(cc)
			if D != []:
				dd['key'] = 'D'
				dd['value'] = str(D[0][0] + D[0][1] + D[0][2] + D[0][3] + D[0][4] + D[0][5]).replace('&nbsp;', '')
				option.append(dd)
			if E != []:
				ee['key'] = 'E'
				ee['value'] = str(E[0][0] + E[0][1] + E[0][2] + E[0][3] + E[0][4] + E[0][5]).replace('&nbsp;', '')
				option.append(ee)
			if F != []:
				ff['key'] = 'F'
				ff['value'] = str(F[0][0] + F[0][1]).replace('&nbsp;', '')
				option.append(ff)
			if newline == [] and opt_line > 4:  # 如果选项换行
				if option != []:
					remap = option.pop()
					remap['value'] += opt
					option.append(remap)


def combine_question(questions, Questions, title_no, paragraphs, end_para_no, stop):
	for com in range(0, (len(questions['combine_ques'])-1)):
		i = 1
		question = ""
		opt_line = 0
		option = []
		answer = ""
		analysis = ""

		combine_ques = ""
		combine = {}
		sub_questions = []
		first_ques = title_no
		everyCombine_stop = questions['combine_ques'][com + 1] if questions['combine_ques'][com + 1] != -1 else stop
		for item in paragraphs:
			para_str = ""
			if questions['combine_ques'][com] <= i < everyCombine_stop:
				para_str += str(item)
				para_str = para_str.replace("\n", "")
				if questions['combine_ques'][com] <= i < questions['question'][first_ques]:
					combine_ques += para_str.replace("【大题】", "")
				if questions['question'][title_no] <= i < questions['option'][title_no]:
					question += para_str.replace("【题文】", "")
				if questions['option'][title_no] <= i < questions['answer'][title_no]:
					opt_line = questions['answer'][title_no] - questions['option'][title_no]
					get_option(option, title_no, questions, para_str, i)
				if questions['answer'][title_no] <= i < questions['analysis'][title_no]:
					answer += para_str.replace("【答案】", "")
				if questions['analysis'][title_no] <= i < end_para_no[title_no]:
					analysis += para_str.replace("【解析】", "")
				if "【结束】" in para_str:
					mm = {}
					mm['question'] = question
					mm['opt_line'] = opt_line
					mm['option'] = option
					mm['answer'] = answer
					mm['analysis'] = analysis
					sub_questions.append(mm)
					question = ""
					opt_line = 0
					option = []
					answer = ""
					analysis = ""
					title_no += 1
			i += 1
		combine['type'] = '组合题'
		combine['question'] = combine_ques
		combine['opt_line'] = 0
		combine['option'] = []
		combine['answer'] = ''
		combine['analysis'] = ''
		combine['subquestions'] = sub_questions
		Questions.append(combine)


def put(byte):
	url = "--------------------"
	return url


if __name__ == '__main__':
	path = r'C:\Users\j20687\Desktop\demo1.docx'
	Exam = read4word(path)
	print(Exam['exam_name'])
	print(Exam['subject'])
	for q in Exam['data']:
		print(q)






