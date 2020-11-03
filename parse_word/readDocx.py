from io import StringIO
import docx
import zipfile
import os
import shutil
from PIL import Image
from PIL import WmfImagePlugin
from win32com import client
from tika import parser
from docx.enum.style import WD_STYLE_TYPE
import xml.etree.cElementTree as ET
import sys
import re
from bs4 import BeautifulSoup
import  xml.dom.minidom
dom= xml.dom.minidom.parse



def read4word(path, zip_path, tmp_path, pic_path, math_path, txt_path, main_txt_path):
	del_path = [tmp_path, math_path, pic_path]
	for pa in del_path:
		for i in os.listdir(pa):
			if os.path.isdir(os.path.join(pa, i)):
				shutil.rmtree(os.path.join(pa, i))

	for i in os.listdir(main_txt_path):
		if os.path.isdir(os.path.join(main_txt_path, i)):
			for file in os.listdir(os.path.join(main_txt_path, i)):
				os.remove(os.path.join(main_txt_path, i, file))




	proxy = []
	os.rename(path, zip_path)
	f = zipfile.ZipFile(zip_path, 'r')  # 解压
	for file in f.namelist():
		f.extract(file, tmp_path)  # 提取图片并保存
	f.close()
	os.rename(zip_path, path)
	rawdir = os.path.join(tmp_path, 'word/media')
	allpic = os.listdir(rawdir)
	math_pic = [f for f in allpic if re.search('wmf$', f)]
	jpg_pic = [f for f in allpic if re.search('jpeg$', f) or re.search('jpg$', f)]
	# print(math_pic)
	for i in range(len(allpic)):
		if allpic[i] in math_pic:
			Image.open(os.path.join(rawdir + "/" + allpic[i]))\
				.save(os.path.join(math_path + '/image%s.jpg' %str(i)))
		if allpic[i] in jpg_pic:
			shutil.copy(os.path.join(rawdir + "/" + allpic[i]), os.path.join(pic_path + '/image%s.jpg' %str(i)))
	docXmlPath = os.path.join(tmp_path, 'word/document.xml')
	try:
		tree = ET.parse(docXmlPath)
		root = tree.getroot()
	except Exception as e:
		print("parse xml文档 fail")
		sys.exit()
	doc = docx.Document(path)
	for para in doc.paragraphs:
		proxy.append(para._element.xml)
	image_no, x_off, y_off, x_ext, y_ext, para_no, x_coor, y_coor = getPic_rId(proxy)
	end_para_no = read_doc4para_no(doc)
	word2html(path, txt_path, html_path, end_para_no, main_txt_path)
	# print(image_no, x_off, y_off, x_ext, y_ext, para_no, x_coor, y_coor)
	# add_pic(doc, para_no, x_coor, y_coor)
	# imgs = getImage(rIds, doc)
	# print(imgs)
	# saveImg(imgs, pic_path)

	# for i in os.listdir(tmp_path):
	# 	if os.path.isdir(os.path.join(tmp_path, i)):
	# 		shutil.rmtree(os.path.join(tmp_path, i))




def getPic_rId(proxy):
	'''
	遍历xml，找到math图片及JPG图片所在的段落并获得坐标
	:param element: xml文件根节点
	:return:
	'''
	i = 1
	x_coor = []
	y_coor = []
	x_off = []
	y_off = []
	x_ext = []
	y_ext = []
	para_no = []
	image_no = []
	for p in proxy:
		root = ET.fromstring(p)  # 直接获取string对象中的根节点,此root是一个Element
		# for elem in root.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
		# 	# print(elem.tag, elem.attrib, elem.text)
		# 	print(elem.text)
		# print("-------------------------------")
		pictr_str = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"
		pictrs = root.findall(pictr_str)
		math_str = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object"
		image_str = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
		for pic in pictrs:
			math_pict = pic.findall(math_str)
			image_pict = pic.findall(image_str)
			if len(image_pict) > 0:
				image_no.append(i)
			if len(math_pict) > 0:
				para_no.append(i)
				x_coor.append(math_pict[0].attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}dxaOrig'])
				y_coor.append(math_pict[0].attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}dyaOrig'])
		for xfrm in root.iter('{http://schemas.openxmlformats.org/drawingml/2006/main}xfrm'):
			for off in xfrm.iter('{http://schemas.openxmlformats.org/drawingml/2006/main}off'):
				x_off.append(off.attrib['x'])
				y_off.append(off.attrib['y'])
			for ext in xfrm.iter('{http://schemas.openxmlformats.org/drawingml/2006/main}ext'):
				x_ext.append(ext.attrib['cx'])
				y_ext.append(ext.attrib['cy'])
		i = i+1
	# print(x_off, y_off, x_ext, y_ext)
	return image_no, x_off, y_off, x_ext, y_ext, para_no, x_coor, y_coor


def read_doc4para_no(doc):
	i = 1
	end_para_no = []
	demo = docx.Document()
	para = doc.paragraphs
	for p in range(len(para)):
		# para_xml = para[p]._element.xml
		# root = ET.fromstring(para_xml)
		# for child in root.iter():
		# 	print(child.tag)

		demo.add_paragraph(para[p].text)
		if para[p].text == "【结束】":
			demo.save('C:\\Users\\j20687\\Desktop\\demo\\demo%s.docx' % str(i))
			demo = docx.Document()
			end_para_no.append(i)
		i = i+1
	return end_para_no
	# run = doc.paragraphs[0].add_run()
	# run.add_picture(r'C:\Users\j20687\Desktop\math\image2.jpg')
	#
	# doc.save(r'C:\Users\j20687\Desktop\111.docx')


def getImage(rIds, doc):
	imgs = []
	for rId in rIds:
		imgs.append(doc.part.related_parts[rId])
	return imgs


def saveImg(imgs, pic_path):
	i = 1
	for img in imgs:
		f = open(os.path.join(pic_path + "/img%d.jpg" % i))
		f.write(img.blob)
		f.close()
		i = i + 1


def word2html(path, txt_path, html_path, end_para_no, main_txt_path):
	word = client.Dispatch('Word.Application')
	word.Visible = 0
	word.DisplayAlerts = 0
	# doc_new = word.Documents.Add()
	doc = word.Documents.Open(path)
	# doc_str = word.ActiveDocument
	# for s in range(len(doc_str.paragraphs)):
	# 	para = doc_str.paragraphs[s].Range
	# 	doc_new = word.Documents.Add()
	# 	rang = doc_new.Range()
	# 	rang.InsertAfter(para)
	# 	print(para)
	# 	doc_new.SaveAs(pre_path + str(s) + '.html', 10)
	# 	doc_new.Close()
	doc.SaveAs(html_path + '.html', 10)  # 选用 wdFormatFilteredHTML的话公式图片将存储为gif格式
	doc.Close()
	word.Quit()
	file_path = html_path+'.html'
	htmlfile = open(file_path, 'r', encoding='gb2312')
	htmlhandle = htmlfile.read()
	soup = BeautifulSoup(htmlhandle, 'lxml')
	para_txt = open(main_txt_path + '/main/main0.txt', 'w', encoding='utf-8')
	i = 1
	for item in soup.find_all('p'):
		if i <= end_para_no[0]:
			para_txt.write(str(item))
			para_txt.write('\n')
		i += 1
	para_txt.close()
	para_no = len(end_para_no)
	for p in range(1, para_no):
		para_txt = open(main_txt_path + '/main/main%d.txt' %p, 'w', encoding='utf-8')
		i = 1
		for item in soup.find_all('p'):
			if i > end_para_no[p-1] and i <= end_para_no[p]:
				para_txt.write(str(item))
				para_txt.write('\n')
			i += 1
		para_txt.close()
	prefix, postfix = store2txtFile(txt_path, main_txt_path, str(soup))
	everyQue2word(prefix, postfix, main_txt_path)
	# get_word(main_txt_path)


def store2txtFile(txt_path, main_txt_path, data):
	file = open(txt_path, 'w', encoding='utf-8')
	file.write(data)
	file.close()
	prefix = main_txt_path + '/prefix.txt'
	postfix = main_txt_path + '/postfix.txt'
	html_prefix = open(prefix, 'w+', encoding='utf-8')
	html_postfix = open(postfix, 'w+', encoding='utf-8')
	with open(txt_path, 'r', encoding='utf-8') as html:
		lines = html.readlines()
		post_line = lines[-3:]
		for ll in post_line:
			html_postfix.write(ll)
		for line in lines:
			html_prefix.write(line)
			if re.compile("<div class").match(line):
				break
		html_postfix.close()
		html_prefix.close()
	return prefix, postfix


def everyQue2word(prefix, postfix, main_txt_path):

	rawdir = os.path.join(main_txt_path, 'main')
	allque = os.listdir(rawdir)
	ques_no = len(os.listdir(rawdir))
	for i in range(ques_no):
		html_prefix = open(prefix, 'r', encoding='utf-8')
		html_postfix = open(postfix, 'r', encoding='utf-8')
		main = open(os.path.join(main_txt_path + '\\main\\' + allque[i]), 'r', encoding='utf-8')
		question = open(main_txt_path + '/word/ques%d' %i + '.txt', 'w', encoding='utf-8')
		for line in html_prefix:
			question.write(line)
		for line in main:
			question.write(line)
		for line in html_postfix:
			question.write(line)
		main.close()
		question.close()
		html_postfix.close()
		html_prefix.close()


def get_word(main_txt_path):
	'''
	将文件先转成html，再转换为Word文档
	:param ques_list:
	:param main_txt_path:
	:return:
	'''
	path1 = main_txt_path + "\\word"
	path2 = main_txt_path + "\\word" + '\\'

	ques_list = os.listdir(path1)
	for filename in ques_list:
		portion = os.path.splitext(filename)
		if portion[1] == ".txt":
			newname = portion[0] + ".html"
			filenamedir = path2 + filename
			newnamedir = path2 + newname
			os.rename(filenamedir, newnamedir)
	ques_list = os.listdir(path1)
	for filename in ques_list:
		portion = os.path.splitext(filename)
		filenamedir = path2 + filename
		newname = portion[0] + ".docx"
		newnamedir = main_txt_path + newname
		word = client.Dispatch('Word.Application')
		doc = word.Documents.Add(filenamedir)
		doc.SaveAs(newnamedir, FileFormat=12)
		doc.Close()
		word.Quit()











if __name__ == '__main__':
	path = 'C:\\Users\\j20687\\Desktop\\demo.docx'
	txt_path = 'C:\\Users\\j20687\\Desktop\\HTML.txt'
	main_txt_path = 'C:\\Users\\j20687\\Desktop\\Main'
	math_path = 'C:\\Users\\j20687\\Desktop\\math'
	zip_path = 'C:\\Users\\j20687\\Desktop\\demo.zip'
	tmp_path = 'C:\\Users\\j20687\\Desktop\\tmp'
	pic_path = 'C:\\Users\\j20687\\Desktop\\pictures'
	html_path = 'C:\\Users\\j20687\\Desktop\\HTML'
	read4word(path, zip_path, tmp_path, pic_path, math_path, txt_path, main_txt_path)




