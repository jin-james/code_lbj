import logging
import sys
import os
import pygame
from docx.shared import Pt
from pygame.locals import *
# from hubarcode.code128 import Code128Encoder
from PIL import Image, ImageDraw, ImageFont
from docx import Document
import barcode
import docx
from barcode.writer import ImageWriter
from docx.enum.text import WD_ALIGN_PARAGRAPH

# logging.getLogger("code128").setLevel(logging.DEBUG)
# logging.getLogger("code128").addHandler(logging.StreamHandler(sys.stdout))
# 12-80  10-80 10-60 12-60 10-70 12-70 10-70
''' 
使用huBarcode,pygame和PIL生成条形码 
'''
def get_barcode_file(stu_list, per_count):
    # 1 生成条形码
    text = "2111602056".upper()
    #   encoder = Code128Encoder(text,options={"ttf_font":r"C:\Users\j20687\Desktop","ttf_fontsize":12,
    # "bottom_border":15,"height":70,"label_border":2})
    # encoder.save("test.png",bar_width=1)
    code = barcode.get_barcode_class('code128')
    encoder = code('{}'.format(text), writer=ImageWriter())
    info_path = r'C:\Users\j20687\Desktop\{}'.format(text)
    doc_path = r"C:\Users\j20687\Desktop\barcode.docx"
    encoder.save(info_path, options={
        "format": 'BMP',
        "text_distance": 1,  # 文本和条码之间的距离，单位为毫米
        "quiet_zone": 1.2,  # 两端空白宽度，单位为毫米
        "font_size": 8,
        "module_height": 4,  # 条形码高度
        'module_width': 0.25,  # 每个条码宽度（？），单位为毫米
        "text": u'金林波',
        # 'center_text'：默认值True，是否居中显示文本
    })
    pic_path = info_path + '.bmp'

    # # 2 生成条码描述
    # pygame.init()
    # content = u"金林波"
    # font = pygame.font.SysFont('SimHei', 20)
    # ftext = font.render(content, True, (0, 0, 0))
    # pygame.image.save(ftext, r"C:\Users\j20687\Desktop\t.png")
    #
    # # 3 合成中文文字到条形码，生成新的条码
    # img = Image.open(r"C:\Users\j20687\Desktop\test.bmp")
    # img_w, img_h = img.size
    # icon = Image.open(r"C:\Users\j20687\Desktop\t.png")
    # icon_w, icon_h = icon.size
    # icon = icon.resize((icon_w, icon_h), Image.ANTIALIAS)
    # w = int((img_w - icon_w) / 2)
    # h = int((img_h - icon_h) / 2) + 40
    # icon = icon.convert("RGBA")
    # img.paste(icon, (w, h), icon)
    # img.save(r"C:\Users\j20687\Desktop\{}".format(text)+".png")

    doc = Document(doc_path)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i in range(per_count):
        para.add_run('  ').add_picture(pic_path, height=Pt(60), width=Pt(190))
    doc.save(doc_path)


if __name__ == "__main__":
    get_barcode_file([], 24)




