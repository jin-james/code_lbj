import traceback
import logging

import os
import time
import docx
import barcode
from barcode.writer import ImageWriter
from django.utils.translation import ugettext_lazy as _

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.xmlchemy import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL

from nile import models
from nile.api import serializers
from nile.api.endpoints.common import QuestionNumberBuilder
from nile.core.response import APIResponse
from nile.utils.ans_sheet_json import get_json
from nile.utils.polo import polo
from nile.utils.constants import CHOICE_QUESTIONS, FILLIN_QUESTIONS
from nile.utils.httputils import DATA_NOT_FOUND, MESSAGE, STORE_FILE_FAILED
from nile.utils.write2word import get_paper

logger = logging.getLogger(__name__)


def ichoice_sub_write(table, r, c, n, ques_no, isfillin=False):
    ques_no += 1
    p = table.cell(r, c).paragraphs[0]
    if isfillin:
        p.add_run().add_break()
        p.add_run(str(ques_no) + '题、________________________________________')
    else:
        p.add_run().add_break()
        p.add_run(str(ques_no) + '题、')
    for i in range(n):
        p.add_run().add_break()
    return ques_no


def set_col_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def hide_frame(table, rows, cols, ischoice=True):
    if ischoice:
        for r in table.rows:
            for c in r.cells:
                set_cell_border(
                    c,
                    top={"sz": 0, "val": "single", "color": "#FFFFFF"},
                    left={"sz": 0, "val": "single", "color": "#FFFFFF"},
                    bottom={"sz": 0, "color": "#FFFFFF", "val": "single"},
                    right={"sz": 0, "val": "single", "color": "#FFFFFF", },
                )
        set_cell_border(
            table.cell(0, 1),
            left={"sz": 5, "val": "single", "color": "#000000"},
            top={"sz": 5, "color": "#000000", "val": "single"},
        )
        set_cell_border(
            table.cell(0, 2),
            top={"sz": 5, "val": "single", "color": "#000000"},
        )
        set_cell_border(
            table.cell(0, 3),
            top={"sz": 5, "color": "#000000", "val": "single"},
            right={"sz": 5, "val": "single", "color": "#000000", },
        )
        set_cell_border(
            table.cell(rows - 1, 1),
            left={"sz": 5, "val": "single", "color": "#000000"},
            bottom={"sz": 5, "color": "#000000", "val": "single"},
        )
        set_cell_border(
            table.cell(rows - 1, 2),
            bottom={"sz": 5, "val": "single", "color": "#000000"},
        )
        set_cell_border(
            table.cell(rows - 1, 3),
            bottom={"sz": 5, "color": "#000000", "val": "single"},
            right={"sz": 5, "val": "single", "color": "#000000", },
        )
        for r in range(1, rows - 1):
            set_cell_border(
                table.cell(r, 1),
                left={"sz": 5, "val": "single", "color": "#000000", },
            )
            set_cell_border(
                table.cell(r, 3),
                right={"sz": 5, "val": "single", "color": "#000000", },
            )
    else:
        for r in range(rows):
            set_cell_border(
                table.cell(r, 0),
                top={"sz": 0, "val": "single", "color": "#FFFFFF"},
                bottom={"sz": 0, "color": "#FFFFFF", "val": "single"},
                left={"sz": 0, "val": "single", "color": "#FFFFFF"},
            )
            set_cell_border(
                table.cell(r, cols - 1),
                top={"sz": 0, "val": "single", "color": "#FFFFFF"},
                bottom={"sz": 0, "color": "#FFFFFF", "val": "single"},
                right={"sz": 0, "val": "single", "color": "#FFFFFF", },
            )


def add_options(p, option_no):
    p.add_run().add_picture(r'nile/static/mark_pic/sel-a.bmp', width=Pt(13.8))
    p.add_run(' ').add_picture(r'nile/static/mark_pic/sel-b.bmp', width=Pt(13.8))
    p.add_run(' ').add_picture(r'nile/static/mark_pic/sel-c.bmp', width=Pt(13.8))
    p.add_run(' ').add_picture(r'nile/static/mark_pic/sel-d.bmp', width=Pt(13.8))
    if int(option_no) == 5:
        p.add_run(' ').add_picture(r'nile/static/mark_pic/sel-e.bmp', width=Pt(13.8))
    if int(option_no) == 6:
        p.add_run(' ').add_picture(r'nile/static/mark_pic/sel-e.bmp', width=Pt(13.8))
        p.add_run(' ').add_picture(r'nile/static/mark_pic/sel-f.bmp', width=Pt(13.8))


def add_mark(cols, table, rows):
    cell1 = table.cell(0, 0)
    cell1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p1 = cell1.paragraphs[0]
    p1.add_run().add_picture(r'nile/static/mark_pic/Rectangle-c.png', width=Pt(6))
    p1.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell2 = table.cell(0, cols + 1)
    cell2.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p2 = cell2.paragraphs[0]
    p2.add_run().add_picture(r'nile/static/mark_pic/Rectangle-c.png', width=Pt(6))
    p2.alignment = WD_TABLE_ALIGNMENT.RIGHT
    if rows > 1:
        cell3 = table.cell(rows - 1, 0)
        cell3.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        p3 = cell3.paragraphs[0]
        p3.add_run().add_picture(r'nile/static/mark_pic/Rectangle-c.png', width=Pt(6))
        p3.alignment = WD_TABLE_ALIGNMENT.LEFT
        cell4 = table.cell(rows - 1, cols + 1)
        cell4.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        p4 = cell4.paragraphs[0]
        p4.add_run().add_picture(r'nile/static/mark_pic/Rectangle-c.png', width=Pt(6))
        p4.alignment = WD_TABLE_ALIGNMENT.RIGHT


def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def write_option(ques_no, choice_len, choice_group, document):
    width5 = (Cm(0.59), Cm(5.34), Cm(5.34), Cm(5.34), Cm(0.59))

    if choice_len:
        if choice_len > 30:
            if divmod(choice_len, 15)[1] != 0:
                rows = divmod(choice_len, 15)[0] + 1
            else:
                rows = divmod(choice_len, 15)[0]
        else:
            rows = 2
        cols = 3
        table_choice = document.add_table(rows=rows, cols=cols + 2, style='Table Grid')
        set_col_widths(table_choice, width5)
        hide_frame(table_choice, rows, cols + 2)
        c = 1
        r = 0
        p = table_choice.cell(r, c).paragraphs[0]
        p.add_run().add_break()
        for n in range(choice_len):
            option_no = len(choice_group[n].get('options', []))
            if ques_no + 1 < 10:
                p.add_run(' ' + str(ques_no + 1) + '、')
            else:
                p.add_run(str(ques_no + 1) + '、')

            add_options(p, option_no)
            ques_no += 1
            if divmod(ques_no, 15)[1] == 0 and (r + 1) * 15 < choice_len:
                r += 1
                c = 1
                p.add_run().add_break()  # 另起一行的时候在前一单元格末加一空行
                p = table_choice.cell(r, c).paragraphs[0]
                p.add_run().add_break()  # 另起一行的时候在前加一空行
            elif divmod(ques_no, 5)[1] == 0:
                c += 1
                p.add_run().add_break()  # 另起一列的时候在前一列末加一空行
                p = table_choice.cell(r, c).paragraphs[0]
                p.add_run().add_break()  # 另起一列的时候在前加一空行
            else:
                p.add_run().add_break()
        add_mark(cols, table_choice, rows)
    return ques_no


def sub_subjective(n, ques_no, page_num, document):
    width3 = (Cm(0.59), Cm(16.02), Cm(0.59))
    rows = 2
    cols = 1
    div1, div2 = divmod(n, 2)
    if div1 != 0 and div2 == 0:
        document.add_page_break()
        page_num += 1
        add_barcode(document, page_num)
    table_sub = document.add_table(rows=rows, cols=cols + 2, style='Table Grid')
    set_col_widths(table_sub, width3)
    add_mark(cols, table_sub, rows)
    table_sub.cell(1, 1).merge(table_sub.cell(0, 1))
    hide_frame(table_sub, rows, cols + 2, ischoice=False)
    ques_no = ichoice_sub_write(table_sub, 0, 1, 18, ques_no)
    return ques_no, page_num


def sub_fillin(fillin_len, ques_no, document):
    width3 = (Cm(0.59), Cm(16.02), Cm(0.59))
    if fillin_len:
        rows = fillin_len + 1 if fillin_len == 1 else fillin_len
        cols = 1
        table_fillin = document.add_table(rows=rows, cols=cols + 2, style='Table Grid')
        set_col_widths(table_fillin, width3)
        hide_frame(table_fillin, rows, cols + 2, ischoice=False)
        for r in range(1 if fillin_len == 1 else rows):
            ques_no = ichoice_sub_write(table_fillin, r, 1, 1, ques_no, isfillin=True)
        add_mark(cols, table_fillin, rows)
        if fillin_len == 1:
            table_fillin.cell(0, 1).merge(table_fillin.cell(1, 1))
    return ques_no


def choice0_60(fillin_len, choice_len, ques_no, document, page_num):
    num = 1
    if 45 >= choice_len >= 31:
        number = 2
        if fillin_len > number:
            ques_no = sub_fillin(number, ques_no, document)
            fillin_len = fillin_len - number
            num += 1
            add_barcode(document, page_num + num)
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document)
        else:
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document,
                                                                                    number=number)
    elif 30 >= choice_len >= 16:
        number = 4
        if fillin_len > number:
            ques_no = sub_fillin(number, ques_no, document)
            fillin_len = fillin_len - number
            num += 1
            document.add_page_break()
            add_barcode(document, page_num + num)
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document)
        else:
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document,
                                                                                number=number)
    elif 15 >= choice_len >= 1:
        number = 6
        if fillin_len > number:
            ques_no = sub_fillin(number, ques_no, document)
            fillin_len = fillin_len - number
            num += 1
            add_barcode(document, page_num + num)
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document)
        else:
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document,
                                                                                number=number)
    else:
        ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document,
                                                                                number=fillin_len)
    return ques_no, fill_no, num, page_num, is_fillin_breakpage


def choice60_(fillin_len, choice_len, ques_no, document, page_num):
    num = 1
    if 75 >= choice_len >= 61:
        number = 2
        if fillin_len > number:
            ques_no = sub_fillin(number, ques_no, document)
            fillin_len = fillin_len - number
            document.add_page_break()
            num += 1
            add_barcode(document, page_num + num)
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document)
        else:
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document,
                                                                                number=number)
    elif 60 >= choice_len >= 46 and fillin_len > 4:
        number = 4
        if fillin_len > number:
            ques_no = sub_fillin(number, ques_no, document)
            fillin_len = fillin_len - number
            document.add_page_break()
            num += 1
            add_barcode(document, page_num + num)
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document)
        else:
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document,
                                                                                number=number)
    elif 45 >= choice_len >= 31:
        number = 7
        if fillin_len > number:
            ques_no = sub_fillin(number, ques_no, document)
            fillin_len = fillin_len - number
            num += 1
            add_barcode(document, page_num + num)
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document)
        else:
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document,
                                                                                number=number)
    elif 30 >= choice_len >= 16:
        number = 9
        if fillin_len > number:
            ques_no = sub_fillin(number, ques_no, document)
            fillin_len = fillin_len - number
            document.add_page_break()
            num += 1
            add_barcode(document, page_num + num)
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document)
        else:
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document,
                                                                                number=number)
    elif 15 >= choice_len >= 1 and fillin_len > 11:
        number = 11
        if fillin_len > number:
            ques_no = sub_fillin(number, ques_no, document)
            fillin_len = fillin_len - number
            num += 1
            add_barcode(document, page_num + num)
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document)
        else:
            ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document,
                                                                                    number=number)
    else:
        ques_no, fill_no, num, is_fillin_breakpage = fillin_in_choice_judgement(fillin_len, ques_no, num, document,
                                                                                number=fillin_len)
    return ques_no, fill_no, num, page_num, is_fillin_breakpage


def fillin_in_choice_judgement(fillin_len, ques_no, page_num, document, number=0):
    if not number:
        d0, d1 = divmod(fillin_len, 14)
        for n in range(d0):
            ques_no = sub_fillin(14, ques_no, document)
            page_num += 1
        ques_no = sub_fillin(d1, ques_no, document)
        is_fillin_breakpage = True
        fill_no = d1
    else:
        ques_no = sub_fillin(fillin_len, ques_no, document)
        is_fillin_breakpage = False
        fill_no = 0
    return ques_no, fill_no, page_num, is_fillin_breakpage


def writeline(document, black_font):
    p = document.add_paragraph()
    # p.add_run().add_break()
    run = p.add_run('非选择题（请在各试题的答题区内作答）')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = black_font


def add_barcode(document, page_num):
    # print(barcode.PROVIDED_BARCODES)
    ITF = barcode.get_barcode_class('itf')
    itf = ITF('{}'.format(page_num), writer=ImageWriter())
    itf_path = r'/tmp/paper/{}'.format(page_num)
    itf.save(itf_path)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    itf_path = "{}.png".format(itf_path)
    run.add_picture(itf_path, height=Cm(0.5), width=Cm(4))
    os.remove(itf_path)


def write4answersheet(answer_sheet_data):
    '''
    answer_sheet_data = {
            "name": paper_data["name"],
            "testee": paper_data["testee"],
            "sheet": answer_sheet
        }
    '''
    sheet = answer_sheet_data.get('sheet', {})
    main_title = answer_sheet_data.get('name')
    answer_path = r'nile/static/mark_pic/answersheet.docx'
    os.makedirs("/tmp/paper", exist_ok=True)
    t0 = int(round(time.time() * 1000))
    tmp_path = '/tmp/paper/%d.docx' % t0
    os.system('cp %s %s' % (answer_path, tmp_path))
    document = Document(tmp_path)
    # 设置整个文档的默认字体
    microsoft_font = u'宋体'  # u 表示后面的字符串以 Unicode 格式进行编码
    black_font = u'黑体'
    number_font = 'Times New Roman'
    area = qn('w:eastAsia')
    document.styles['Normal'].font.name = microsoft_font
    document.styles['Normal'].font.size = Pt(10.5)
    document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
    document.styles['Normal']._element.rPr.rFonts.set(area, number_font)
    # 指定段落样式
    styles = document.styles
    # fls = time.time()
    # strr = 'lbj_Style%s' %fls  #自定义的样式的名称
    # strr = strr.replace('.', '')
    # strr = strr + ''.join(random.sample('zyxwvutsrqponmlkjihgfedcbaABCDEFGHIJKLMNOPQRST', 5))
    s = styles.add_style('lbj_Style', WD_STYLE_TYPE.PARAGRAPH)
    s.font.name = microsoft_font
    s.font.size = Pt(10.5)
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.line_spacing = Pt(0)  # 行距值
    s.paragraph_format.space_before = Pt(0)  # 段前距
    s.paragraph_format.space_after = Pt(0)  # 段后距
    s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE  # 行间距
    s._element.rPr.rFonts.set(area, number_font)  # 除中文外其它文字 使用的字体 ，备选项

    add_barcode(document, 1)
    run = document.add_paragraph()
    run.style = s
    p = run.add_run(main_title)
    font = p.font
    font.name = black_font
    font.color.rgb = RGBColor(0, 0, 0)
    font.size = Pt(15)
    run.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph()
    p.add_run().add_picture(r'nile/static/mark_pic/prefix.png')

    p = document.add_paragraph()
    # p.add_run().add_break()
    run = p.add_run('选择题（请用2B铅笔填涂）')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = black_font
    choice_group = sheet.get('choice_question', [])
    sub_mark_num_vec = []  # 每页竖长条表示块数量
    sub_mark_num = 0  # 竖长条表示块总数量
    sub_mark_choice = []  # 选择题黑块区域
    sub_mark_fillin = []  # 填空题黑块区域
    sub_mark_eassy = []  # 简答题黑块区域
    choice_page = 0  # 选择题所在页数
    fillin_page = 0  # 填空题所在页数
    eassy_page = 0  # 简答题所在页数
    page_num = 0  # 页数
    ques_no = 0  # 题号
    choice_answer_cnt = 4  # 选项列数
    choice_answer_pattern = []  # 每列选择题数目
    choice_col_num = 0  # 选择题每行最多列数
    choice_len = len(choice_group)
    fillin_group = sheet.get('fillin_question', [])
    fillin_len = len(fillin_group)
    fill_no = 0
    is_fillin_breakpage = False
    if 0 <= choice_len <= 5:
        choice_col_num = 1
        choice_answer_pattern.append(choice_len)
    elif 6 <= choice_len <= 10:
        choice_col_num = 2
        choice_answer_pattern.append(5)
        choice_answer_pattern.append(choice_len - 5)
    else:
        choice_col_num = 3
        for n in range(divmod(choice_len, 5)[0]):
            choice_answer_pattern.append(5)
        if divmod(choice_len, 5)[1]:
            choice_answer_pattern.append(divmod(choice_len, 5)[1])

    if choice_len > 150:
        ques_no = write_option(ques_no, 60, choice_group, document)
        document.add_page_break()
        page_num += 1
        add_barcode(document, page_num + 1)
        else_choice = choice_len - 60
        pre_no, rest = divmod(else_choice, 90)
        for n in range(pre_no):
            page_num += 1
            add_barcode(document, page_num + 1)
            ques_no = write_option(ques_no, 90, choice_group, document)
            document.add_page_break()
        ques_no = write_option(ques_no, rest, choice_group, document)
        writeline(document, black_font)
        choice_page = page_num + 1 if rest else page_num
        ques_no, fill_no, num, page_num, is_fillin_breakpage = choice60_(fillin_len, rest, ques_no, document, page_num)
        page_num += num
        fillin_page = page_num
    else:
        if choice_len > 60:
            else_choice = choice_len - 60
            ques_no = write_option(ques_no, 60, choice_group, document)
            document.add_page_break()
            page_num += 1
            add_barcode(document, page_num + 1)
            ques_no = write_option(ques_no, else_choice, choice_group, document)
            writeline(document, black_font)
            choice_page = page_num + 1 if else_choice else page_num
            ques_no, fill_no, num, page_num, is_fillin_breakpage = choice60_(fillin_len, else_choice, ques_no, document,
                                                                             page_num)
            page_num += num
            fillin_page = page_num
        else:
            ques_no = write_option(ques_no, choice_len, choice_group, document)
            writeline(document, black_font)
            choice_page = page_num + 1
            ques_no, fill_no, num, page_num, is_fillin_breakpage = choice0_60(fillin_len, choice_len, ques_no, document,
                                                                              page_num)
            page_num += num
            fillin_page = page_num
    subjective_question = sheet.get('subjective_question', [])
    subjec_len = len(subjective_question)
    rest_subjec_len = subjec_len
    nn = 0
    if is_fillin_breakpage and fillin_len:
        if fill_no == 0:
            nn = 2
            rest_subjec_len = subjec_len - nn
            for n in range(nn):
                ques_no, page_num = sub_subjective(n, ques_no, page_num, document)
            eassy_page = page_num
            document.add_page_break()
            page_num += 1
            add_barcode(document, page_num)
        elif 7 >= fill_no >= 1:
            nn = 1
            rest_subjec_len = subjec_len - nn
            for n in range(nn):
                ques_no, page_num = sub_subjective(n, ques_no, page_num, document)
            eassy_page = page_num
            document.add_page_break()
            page_num += 1
            add_barcode(document, page_num)
        else:
            document.add_page_break()
            page_num += 1
            add_barcode(document, page_num)
            eassy_page = page_num
    else:
        document.add_page_break()
        page_num += 1
        add_barcode(document, page_num)
        eassy_page = page_num
    div1, div2 = divmod(rest_subjec_len, 2)
    for n in range(rest_subjec_len):
        ques_no, page_num = sub_subjective(n, ques_no, page_num, document)
    sub_page_num = 1 if div2 else 0
    page_num += sub_page_num
    #  return sub_mark_num_vec except eassy
    if choice_len:
        if fillin_page >= choice_page:
            d1, d2 = divmod(choice_len, 90)
            if d2 > 45:
                for n in range(d1):
                    sub_mark_num_vec.append(4)
            else:
                if fill_no >= 0:
                    sub_mark_num_vec.append(8)
                    if divmod(fill_no, 14)[1] > 7:
                        for n in range(divmod(fill_no, 14)[0]):
                            sub_mark_num_vec.append(4)
                    else:
                        if fill_no:
                            sub_mark_num_vec.append(4 + nn * 4)
        else:
            if fillin_len:
                sub_mark_num_vec.append(8)
            else:
                sub_mark_num_vec.append(6)
    else:
        if fillin_len:
            sub_mark_num_vec.append(4 + nn * 4 if fill_no else nn * 4)
        else:
            sub_mark_num_vec.append(2)
    sub_mark_num_vec[0] = sub_mark_num_vec[0] + 2
    #  return sub_mark_num_vec with eassy
    for n in range(div1):
        sub_mark_num_vec.append(8)
    if div2 != 0:
        sub_mark_num_vec.append(4)
    #  return sub_mark for choice, fillin, eassy
    block_num = 0
    start = 3
    step = 3
    if choice_len and choice_page:
        for n in range(choice_page):
            block_num += 1
            sub_mark_choice.append(start + step * (block_num - 1) + block_num - 1)
            sub_mark_choice.append(start + step * block_num + block_num - 1)
    if fillin_page >= choice_page and divmod(choice_len, 90)[1] <= 45:
        for n in range(fillin_page):
            block_num += 1
            sub_mark_fillin.append(start + step * (block_num - 1) + block_num - 1)
            sub_mark_fillin.append(start + step * block_num + block_num - 1)
    else:
        for n in range(fillin_page - choice_page):
            block_num += 1
            sub_mark_fillin.append(start + step * (block_num - 1) + block_num - 1)
            sub_mark_fillin.append(start + step * block_num + block_num - 1)
    for n in range(subjec_len):
        block_num += 1
        sub_mark_eassy.append(start + step * (block_num - 1) + block_num - 1)
        sub_mark_eassy.append(start + step * block_num + block_num - 1)

    #  return sub_mark_num
    for num in sub_mark_num_vec:
        sub_mark_num += num

    t1 = int(round(time.time() * 1000))
    return_path = '/tmp/paper/%s.%s' % (main_title or str(t1), 'docx')
    document.save(return_path)
    sub_json_data = {}
    sub_json_data['sub_mark_choice'] = sub_mark_choice
    sub_json_data['sub_mark_fillin'] = sub_mark_fillin
    sub_json_data['sub_mark_eassy'] = sub_mark_eassy
    sub_json_data['sub_mark_num'] = sub_mark_num
    sub_json_data['sub_mark_num_vec'] = sub_mark_num_vec
    sub_json_data['page_num'] = page_num
    sub_json_data['choice_answer_cnt'] = choice_answer_cnt
    sub_json_data['choice_answer_pattern'] = choice_answer_pattern
    sub_json_data['choice_col_num'] = choice_col_num

    return return_path, sub_json_data


def get4answer_sheet_data(paper_data):
    answer_sheet = {}
    for g_idx, grp in enumerate(paper_data.get("question_group", [])):
        for eq in grp.get("exam_questions", []):
            question = eq.get("question", {})
            q_type_name = question.get("q_type", {}).get("name")
            options = question.get("options") or []
            if q_type_name in CHOICE_QUESTIONS:
                key = "choice_question"
            elif q_type_name in FILLIN_QUESTIONS:
                key = "fillin_question"
            else:
                key = "subjective_question"
            if not answer_sheet.get(key):
                answer_sheet[key] = []
            answer_sheet[key].append({
                "q_no": eq["q_no"],
                "options": [v["key"] for v in options if v.get("key")]
            })
    answer_sheet_data = {
        "name": paper_data["name"],
        "testee": paper_data["testee"],
        "sheet": answer_sheet
    }
    answer_path, sub_json_data = write4answersheet(answer_sheet_data)
    return answer_path, answer_sheet_data


def create_paper_file(paper, data, fmt, size, usage):
    try:
        # create paper file
        filepath = get_paper(data, fmt, size, usage)
        outfile = open(filepath, "rb")
        r = polo.upload_file(outfile)
        if not r:
            return APIResponse(code=STORE_FILE_FAILED, message=_("存储文件失败"))
        hash_code = r.hash
        models.ExamRelatedPaper.save_hash(paper, size, hash_code)
    except Exception as e:
        logger.error("upload paper file to polo error, detail: {}".format(e))


def create_answersheet(paper, data):
    try:
        # create answer sheet
        paper_data = QuestionNumberBuilder(paper_data=data).add_question_no
        answer_path, answer_sheet_data = get4answer_sheet_data(paper_data)
        outfile = open(answer_path, "rb")
        r = polo.upload_file(outfile)
        if not r:
            return APIResponse(code=STORE_FILE_FAILED, message=_("存储文件失败"))
        answer_hash_code = r.hash
        # save json_data and hash
        json_data = get_json(answer_sheet_data)
        models.ExamRelatedAnswersheet.save_answer(paper, models.AnswerSheetSizeSet.A4, answer_hash_code, json_data)
    except Exception as e:
        traceback.print_exc()
        logger.error("upload answer sheet file to polo error, detail: {}".format(e))



