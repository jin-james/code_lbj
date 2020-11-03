# -*- coding: utf-8 -*-
from __future__ import absolute_import, unicode_literals

import os
import time

import barcode
import docx
from barcode.writer import ImageWriter
from docx.shared import Cm


def add_barcode(info, code_type='code128', options=None):
    """生成条形码
    code(code, writer=None, add_checksum=True)
    code('123456',barcode.writer.ImageWriter(),False), 默认使用barcodr.writer.SVGWriter()
    """
    assert code_type in barcode.PROVIDED_BARCODES
    t1 = int(round(time.time() * 1000))
    code = barcode.get_barcode_class(code_type)
    info = code('{}'.format(info), writer=ImageWriter(), add_checksum=False)
    info_path = r'C:\Users\j20687\Desktop\barcode{}'.format(str(t1))
    info.save(info_path, options=options)
    info_path = info_path + '.' + str(options['format'].lower())
    return info_path


def get_testee_barcode(numbers):
    """
        'module_width'：默认值0.2，每个条码宽度（？），单位为毫米
        'module_height'：默认值15.0，条码高度，单位为毫米
        'quiet_zone'：默认值6.5，两端空白宽度，单位为毫米
        'font_size'：默认值10，文本字体大小，单位为磅
        'text_distance'：默认值5.0，文本和条码之间的距离，单位为毫米
        'background'：默认值'white'，背景色
        'foreground'：默认值'black'，前景色
        'text'：默认值''，显示文本，默认显示编码，也可以自行设定
        'write_text'：默认值True，是否显示文本，如果为True自动生成text的值，如果为False则不生成（如果此时手工设置了text的值，仍然会显示文本）。
        'center_text'：默认值True，是否居中显示文本
        'format'：默认值'PNG'，保存文件格式，默认为PNG，也可以设为JPEG、BMP等，只在使用ImageWriter时有效。
        'dpi'：默认值300，图片分辨率，，只在使用ImageWriter时有效。
    """
    path = add_barcode(str(numbers), code_type='code128', options={
        "format": 'BMP',
        "text_distance": 1,
        "quiet_zone": 1,
        "font_size": 10,
        "module_height": 4,
        "text": u'金林波'
    })
    document = docx.Document(r'C:\Users\j20687\Desktop\校本资源相关\答题卡\testee_barcode.docx')
    p = document.add_paragraph()
    p.add_run().add_picture(r'{}'.format(path), height=Cm(1.8), width=Cm(6.6))
    document.save(r'C:\Users\j20687\Desktop\testee_barcode.docx')
    # os.remove(path)


if __name__ == '__main__':
    get_testee_barcode('j2344567889')
