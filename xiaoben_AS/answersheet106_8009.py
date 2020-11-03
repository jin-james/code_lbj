import logging
import os
import re
import time
import traceback

import barcode
import requests
from barcode.writer import ImageWriter
from django.utils.translation import ugettext_lazy as _
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml.xmlchemy import OxmlElement
from docx.shared import Pt, RGBColor, Cm

from nile import models
from nile.api.common import QuestionNumberBuilder
from nile.core.response import APIResponse
from nile.utils.ans_sheet_json import get_json
from nile.utils.constants import CHOICE_QUESTIONS, FILL_IN_QUESTIONS
from nile.utils.fileutils import remove_file
from nile.utils.httputils import STORE_FILE_FAILED
from nile.utils.polo import polo
from nile.utils.write2word import get_paper

logger = logging.getLogger(__name__)
imgtmp_base_path = '/tmp/img'


def ichoice_sub_write(table, r, c, n, ques_no, pic_in_sub, isfillin=False):
    cell = table.cell(r, c)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p = cell.paragraphs[0]
    if isfillin:
        p.add_run().add_break()
        p.add_run(str(ques_no) + '题、________________________________________')
    else:
        if not pic_in_sub:
            p.add_run().add_break()
        p.add_run(str(ques_no) + '题、')
    if pic_in_sub:
        add_pic(table, pic_in_sub)
        n -= 3
    for i in range(n):
        p.add_run().add_break()
    ques_no += 1
    return ques_no


def set_col_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def hide_frame(table, rows, cols, ischoice=True):
    if ischoice:
        for r in table.rows:
            for c in r.cells:
                set_cell_border(
                    c,
                    top={"sz": 0, "val": "single", "color": "#FFFFFF"},
                    left={"sz": 0, "val": "single", "color": "#FFFFFF"},
                    bottom={"sz": 0, "color": "#FFFFFF", "val": "single"},
                    right={"sz": 0, "val": "single", "color": "#FFFFFF", },
                )
        set_cell_border(
            table.cell(0, 1),
            left={"sz": 5, "val": "single", "color": "#000000"},
            top={"sz": 5, "color": "#000000", "val": "single"},
        )
        set_cell_border(
            table.cell(0, 2),
            top={"sz": 5, "val": "single", "color": "#000000"},
        )
        set_cell_border(
            table.cell(0, 3),
            top={"sz": 5, "color": "#000000", "val": "single"},
            right={"sz": 5, "val": "single", "color": "#000000", },
        )
        set_cell_border(
            table.cell(rows - 1, 1),
            left={"sz": 5, "val": "single", "color": "#000000"},
            bottom={"sz": 5, "color": "#000000", "val": "single"},
        )
        set_cell_border(
            table.cell(rows - 1, 2),
            bottom={"sz": 5, "val": "single", "color": "#000000"},
        )
        set_cell_border(
            table.cell(rows - 1, 3),
            bottom={"sz": 5, "color": "#000000", "val": "single"},
            right={"sz": 5, "val": "single", "color": "#000000", },
        )
        for r in range(1, rows - 1):
            set_cell_border(
                table.cell(r, 1),
                left={"sz": 5, "val": "single", "color": "#000000", },
            )
            set_cell_border(
                table.cell(r, 3),
                right={"sz": 5, "val": "single", "color": "#000000", },
            )
    else:
        for r in range(rows):
            set_cell_border(
                table.cell(r, 0),
                top={"sz": 0, "val": "single", "color": "#FFFFFF"},
                bottom={"sz": 0, "color": "#FFFFFF", "val": "single"},
                left={"sz": 0, "val": "single", "color": "#FFFFFF"},
            )
            set_cell_border(
                table.cell(r, cols - 1),
                top={"sz": 0, "val": "single", "color": "#FFFFFF"},
                bottom={"sz": 0, "color": "#FFFFFF", "val": "single"},
                right={"sz": 0, "val": "single", "color": "#FFFFFF", },
            )


def add_options(p, option_no):
    p.add_run().add_picture(r'nile/static/mark_pic/sel-a_.bmp', width=Pt(12.5))
    p.add_run(' ').add_picture(r'nile/static/mark_pic/sel-b_.bmp', width=Pt(12.5))
    p.add_run(' ').add_picture(r'nile/static/mark_pic/sel-c_.bmp', width=Pt(12.5))
    p.add_run(' ').add_picture(r'nile/static/mark_pic/sel-d_.bmp', width=Pt(12.5))
    if int(option_no) == 5:
        p.add_run(' ').add_picture(r'nile/static/mark_pic/sel-e_.bmp', width=Pt(12.5))
    if int(option_no) == 6:
        p.add_run(' ').add_picture(r'nile/static/mark_pic/sel-e_.bmp', width=Pt(12.5))
        p.add_run(' ').add_picture(r'nile/static/mark_pic/sel-f_.bmp', width=Pt(12.5))
    if int(option_no) == 7:
        p.add_run(' ').add_picture(r'nile/static/mark_pic/sel-e_.bmp', width=Pt(12.5))
        p.add_run(' ').add_picture(r'nile/static/mark_pic/sel-f_.bmp', width=Pt(12.5))
        p.add_run(' ').add_picture(r'nile/static/mark_pic/sel-g_.bmp', width=Pt(12.5))


def add_mark(cols, table):
    cell1 = table.cell(0, 0)
    cell1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p1 = cell1.paragraphs[0]
    p1.add_run().add_picture(r'nile/static/mark_pic/Rectangle-c.png', width=Pt(6))
    p1.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell2 = table.cell(0, cols + 1)
    cell2.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p2 = cell2.paragraphs[0]
    p2.add_run().add_picture(r'nile/static/mark_pic/Rectangle-c.png', width=Pt(6))
    p2.alignment = WD_TABLE_ALIGNMENT.RIGHT


def add_pic(table_sub, pic_in_sub):
    # table_sub.alignment = WD_TABLE_ALIGNMENT.CENTER
    para = table_sub.cell(0, 1).paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    re_image = re.compile("<img.*?>", re.S | re.I)
    re_src = re.compile("src=\"(.*?)\"", re.S | re.I)
    re_height = re.compile("height=\"(.*?)\"", re.S | re.I)
    line = [v for v in pic_in_sub if v]
    line = line[0] if line else ""
    image = re_image.findall(line)
    for img in image:
        string = re_src.search(img)
        string = string[1] if string else ""
        hei_str = re_height.search(img)
        hei_str = hei_str[1] if hei_str else ""
        content = requests.get(string).content
        subfix = string[-32:]
        path = imgtmp_base_path + '/' + subfix + '{}'.format('.png')
        height_v = float(hei_str.split('p')[0])
        if os.path.exists(path):
            try:
                para.add_run().add_picture(path, height=Pt(height_v * 0.7))
                continue
            except:
                remove_file(path)
        with open(path, 'wb')as f:
            f.write(content)
        para.add_run().add_picture(path, height=Pt(height_v * 0.7))
    for i in range(len(pic_in_sub)):
        tmp = pic_in_sub[i]
        if not tmp:
            pic_in_sub[i] = []  # 剔除已经添加的图片数据
            break


def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


# def write_option(ques_no, choice_len, option_len, choice_ques_no, document):
#     width5 = (Cm(0.59), Cm(5.34), Cm(5.34), Cm(5.34), Cm(0.59))
#     if choice_len:
#         # if choice_len > 30:
#         #     if divmod(choice_len, 15)[1] != 0:
#         #         rows = divmod(choice_len, 15)[0] + 1
#         #     else:
#         #         rows = divmod(choice_len, 15)[0]
#         # else:
#         #     rows = 2
#         cols = 3
#         table_choice = document.add_table(rows=rows, cols=cols + 2, style='Table Grid')
#         set_col_widths(table_choice, width5)
#         hide_frame(table_choice, rows, cols + 2)
#         c = 1
#         r = 0
#         no = ques_no
#         p = table_choice.cell(r, c).paragraphs[0]
#         p.add_run().add_break()
#         for n in range(choice_len):
#             option_no = option_len[n]
#             choice_no = choice_ques_no[n]
#             no = choice_no if choice_no > no else no
#             # if break_no and break_len:
#             #     if no < break_len:
#             #         if no + 1 < 10:
#             #             p.add_run(' ' + str(no + 1) + '、')
#             #         else:
#             #             p.add_run(str(no + 1) + '、')
#             #         no += 1
#             #     else:
#             #         if break_no + 1 < 10:
#             #             p.add_run(' ' + str(break_no + 1) + '、')
#             #         else:
#             #             p.add_run(str(break_no + 1) + '、')
#             #         break_no += 1
#             # else:
#
#             if choice_no < 10:
#                 p.add_run(' ' + str(choice_no) + '、')
#             else:
#                 p.add_run(str(choice_no) + '、')
#             # for n in range(choice_len):
#             #     option_no = option_len[n]
#             #     if ques_no + 1 < 10:
#             #         p.add_run(' ' + str(ques_no + 1) + '、')
#             #     else:
#             #         p.add_run(str(ques_no + 1) + '、')
#             add_options(p, option_no)
#             ques_no += 1
#             if divmod(ques_no, 15)[1] == 0 and (r + 1) * 15 < choice_len:
#                 r += 1
#                 c = 1
#                 p.add_run().add_break()  # 另起一行的时候在前一单元格末加一空行
#                 p = table_choice.cell(r, c).paragraphs[0]
#                 p.add_run().add_break()  # 另起一行的时候在前加一空行
#             elif divmod(ques_no, 5)[1] == 0:
#                 if ques_no != 15:
#                     c += 1
#                     p.add_run().add_break()  # 另起一列的时候在前一列末加一空行
#                     p = table_choice.cell(r, c).paragraphs[0]
#                     p.add_run().add_break()  # 另起一列的时候在前加一空行
#             else:
#                 p.add_run().add_break()
#         add_mark(cols, table_choice)
def write_option(ques_no, choice_num, choice_opt, rows, document):
    width5 = (Cm(0.59), Cm(5.34), Cm(5.34), Cm(5.34), Cm(0.59))
    choice_len = len(choice_num)
    if choice_len:
        # if choice_len > 30:
        #     if divmod(choice_len, 15)[1] != 0:
        #         rows = divmod(choice_len, 15)[0] + 1
        #     else:
        #         rows = divmod(choice_len, 15)[0]
        # else:
        #     rows = 2
        cols = 3
        table_choice = document.add_table(rows=rows, cols=cols + 2, style='Table Grid')
        set_col_widths(table_choice, width5)
        hide_frame(table_choice, rows, cols + 2)
        c = 1
        r = 0
        p = table_choice.cell(r, c).paragraphs[0]
        p.add_run().add_break()
        for n in range(choice_len):
            option_no = choice_opt[n]
            # count = choice_count[n]
            if n > 0 and option_no != choice_opt[n-1]:
                # div1 = divmod(count, 15)[1]
                # if 0 < div1 < 5:
                #     for i in range(5 - div1):
                #         p.add_run().add_break()
                ques_no = 0
                r += 1
                c = 1
                p = table_choice.cell(r, c).paragraphs[0]
                p.add_run().add_break()
            choice_no = choice_num[n]
            if choice_no < 10:
                p.add_run(' ' + str(choice_no) + '.')
            else:
                p.add_run(str(choice_no) + '.')
            add_options(p, option_no)
            ques_no += 1
            if divmod(ques_no, 15)[1] == 0 and (r + 1) * 15 < choice_len:
                r += 1
                c = 1
                p.add_run().add_break()  # 另起一行的时候在前一单元格末加一空行
                p = table_choice.cell(r, c).paragraphs[0]
                p.add_run().add_break()  # 另起一行的时候在前加一空行
            elif divmod(ques_no, 5)[1] == 0:
                if ques_no != 15:
                    c += 1
                    p.add_run().add_break()  # 另起一列的时候在前一列末加一空行
                    p = table_choice.cell(r, c).paragraphs[0]
                    p.add_run().add_break()  # 另起一列的时候在前加一空行
            else:
                p.add_run().add_break()
        add_mark(cols, table_choice)


def sub_subjective(ques_no, paper_num, document, pic_in_sub):
    width3 = (Cm(0.59), Cm(16.02), Cm(0.59))
    rows = 2
    cols = 1
    table_sub = document.add_table(rows=rows, cols=cols + 2, style='Table Grid')
    set_col_widths(table_sub, width3)
    add_mark(cols, table_sub)
    table_sub.cell(1, 1).merge(table_sub.cell(0, 1))
    hide_frame(table_sub, rows, cols + 2, ischoice=False)
    ques_no = ichoice_sub_write(table_sub, 0, 1, 18, ques_no, pic_in_sub)
    return ques_no, paper_num


def sub_fillin(fillin_len, ques_no, document):
    width3 = (Cm(0.59), Cm(16.02), Cm(0.59))
    if fillin_len:
        rows = fillin_len + 1 if fillin_len == 1 else fillin_len
        cols = 1
        table_fillin = document.add_table(rows=rows, cols=cols + 2, style='Table Grid')
        set_col_widths(table_fillin, width3)
        hide_frame(table_fillin, rows, cols + 2, ischoice=False)
        for r in range(1 if fillin_len == 1 else rows):
            ques_no = ichoice_sub_write(table_fillin, r, 1, 1, ques_no, [], isfillin=True)
        add_mark(cols, table_fillin)
        if fillin_len == 1:
            table_fillin.cell(0, 1).merge(table_fillin.cell(1, 1))
    return ques_no


def choice0_60(fillin_len, ques_no, document, paper_num):
    num = 1
    ques_no, num = fillin_in_choice_judgement(fillin_len, ques_no, num, document)
    return num, paper_num


def fillin_in_choice_judgement(fillin_len, ques_no, paper_num, document):
    ques_no = sub_fillin(fillin_len, ques_no, document)
    return ques_no, paper_num,


def writeline(document, black_font):
    p = document.add_paragraph()
    # p.add_run().add_break()
    run = p.add_run('非选择题（请在各试题的答题区内作答）')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = black_font


def add_barcode(document, paper_num, paper_num_count_flag=1, code_type='itf'):
    # print(barcode.PROVIDED_BARCODES)
    if paper_num_count_flag:
        code = barcode.get_barcode_class(code_type)
        itf = code('{}'.format(paper_num), writer=ImageWriter())
        itf_path = r'/tmp/paper/{}'.format(paper_num)
        itf.save(itf_path)
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        itf_path = "{}.png".format(itf_path)
        run.add_picture(itf_path, height=Cm(0.5), width=Cm(4))
        os.remove(itf_path)


def write4answersheet(answer_sheet_data):
    '''
    answer_sheet_data = {
        "name": paper_data["name"],
        "main_title": "{}-答题卡".format(paper_data["main_title"]),
        "testee": paper_data["testee"],
        "sheet": answer_sheet
    }
    answer_sheet[key].append({
                "q_no": ques_no,
                "pic_in_card": pic_in_card,
                "options": [v["key"] for v in options if v.get("key")],
                "subques_len": len(subquestions) if subquestions else 0
            })
    '''
    sheet = answer_sheet_data.get('sheet', {})
    sheet_name = answer_sheet_data.get('name')
    main_title = answer_sheet_data.get('main_title')
    answer_path = r'nile/static/mark_pic/answersheet.docx'
    os.makedirs("/tmp/paper", exist_ok=True)
    t0 = int(round(time.time() * 1000))
    tmp_path = '/tmp/paper/%d.docx' % t0
    os.system('cp %s %s' % (answer_path, tmp_path))
    document = Document(tmp_path)
    # 设置整个文档的默认字体
    microsoft_font = u'宋体'  # u 表示后面的字符串以 Unicode 格式进行编码
    black_font = u'黑体'
    number_font = 'Times New Roman'
    area = qn('w:eastAsia')
    document.styles['Normal'].font.name = microsoft_font
    document.styles['Normal'].font.size = Pt(10.5)
    document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
    document.styles['Normal']._element.rPr.rFonts.set(area, number_font)
    # 指定段落样式
    styles = document.styles
    # fls = time.time()
    # strr = 'lbj_Style%s' %fls  #自定义的样式的名称
    # strr = strr.replace('.', '')
    # strr = strr + ''.join(random.sample('zyxwvutsrqponmlkjihgfedcbaABCDEFGHIJKLMNOPQRST', 5))
    s = styles.add_style('lbj_Style', WD_STYLE_TYPE.PARAGRAPH)
    s.font.name = microsoft_font
    s.font.size = Pt(10.5)
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.line_spacing = Pt(0)  # 行距值
    s.paragraph_format.space_before = Pt(0)  # 段前距
    s.paragraph_format.space_after = Pt(0)  # 段后距
    s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE  # 行间距
    s._element.rPr.rFonts.set(area, number_font)  # 除中文外其它文字 使用的字体 ，备选项

    paper_num_count_flag = 0  # 页面是否有页码条形码
    run = document.paragraphs[0]
    run.style = s
    p = run.add_run(main_title)
    font = p.font
    font.name = black_font
    font.color.rgb = RGBColor(0, 0, 0)
    font.size = Pt(15)
    run.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph()
    p.add_run().add_picture(r'nile/static/mark_pic/prefix_2.png')

    p = document.add_paragraph()
    # p.add_run().add_break()
    run = p.add_run('选择题（请用2B铅笔填涂）')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = black_font
    # sub_mark_num_vec = []  # 每页竖长条表示块数量
    sub_mark_num = 0  # 竖长条表示块总数量
    sub_mark_choice = []  # 选择题黑块区域
    sub_mark_fillin = []  # 填空题黑块区域
    sub_mark_eassy = []  # 简答题黑块区域
    qus_range_choice = []  # 选择题题号范围
    qus_range_fillin = []  # 填空题题号范围
    qus_range_eassy = []  # 简答题题号范围
    paper_num = 0  # 页数
    ques_no = 0  # 题号
    choice_answer_cnt = []  # 选项列数
    choice_line_number = []  # 每行的题目数量
    choice_answer_pattern = []  # 每列选择题数目
    choice_col_num = []  # 选择题每行的列数
    option_len = []
    choice_ques_no, mm, combine_ques_index = [], [], []
    # break_page_ques_no = []  # 如果选择题分页时的题号
    count_ques = 0
    for key, value in sheet.items():
        if key == 'choice_question_all':
            for k in range(len(value)):
                q_no = value[k].get("q_no")
                opt_len = len(value[k].get("options"))
                mm.append(q_no)
                if not isinstance(q_no, list):
                    option_len.append(opt_len)
                else:
                    for qq in q_no:
                        option_len.append(opt_len)
    for i in range(len(mm)):
        if isinstance(mm[i], list):
            count_ques += len(mm[i])
            combine_ques_index.append(i)
            for j in mm[i]:
                choice_ques_no.append(j)
        else:
            count_ques += 1
            choice_ques_no.append(mm[i])
    # 选项个数不一样
    group_choice_count = {}
    group_choice_opt = {}  # 各选项数对应的选项数集合
    group_choice_num = {}  # 各选项数对应的题号集合
    for opt_i in range(len(option_len)):
        pre = option_len[opt_i]
        if not group_choice_count.get(pre):
            group_choice_count[pre] = 0
            group_choice_num[pre] = []
            group_choice_opt[pre] = []
        count = group_choice_count[pre]
        group_choice_count[pre] = count + 1
        group_choice_num[pre].append(choice_ques_no[opt_i])
        group_choice_opt[pre].append(option_len[opt_i])
    choice_group = sheet.get('choice_question', [])
    choice_len = len(choice_group)
    choice_len_sum = len(choice_ques_no)  # 选择题总数，包括组合题中的小选择题
    fillin_group = sheet.get('fillin_question', [])
    fillin_len = len(fillin_group)
    fs_num = fillin_group[0].get('q_no') if fillin_group else 0  # 填空题开始的题号
    subjective_question = sheet.get('subjective_question', [])
    subjec_len = len(subjective_question)
    ss_num = subjective_question[0].get('q_no') if subjective_question else 0  # 主观题开始的题号
    combine_question = sheet.get('combine_question', [])

    pic_in_sub = [[] for x in range(subjec_len)]
    pic_num = 0
    for ques in subjective_question:
        pic_in_card = ques.get('pic_in_card', "")
        pic_in_sub[pic_num] = pic_in_card if pic_in_card else []
        pic_num += 1
    tmp = []
    for i in range(len(mm)):
        if isinstance(mm[i], list):
            if tmp:
                qus_range_choice.append(tmp)
                tmp = []
            qus_range_choice.append(mm[i])
        else:
            tmp.append(mm[i])

    if not combine_question:
        qus_range_choice.append(tmp)
    else:
        if not isinstance(mm[-1], list):
            qus_range_choice.append(tmp)

    if fillin_group:
        qus_range_fillin.append(fs_num)
        qus_range_fillin.append(fs_num + fillin_len - 1)
    if subjective_question:
        qus_range_eassy.append(ss_num)
        qus_range_eassy.append(ss_num + subjec_len - 1)

    # row_n, row_rest = divmod(choice_len_sum, 15)
    # choice_row_num = row_n + 1 if row_rest else row_n  # 选择题行数
    # choice_row_num = [choice_row_num] if choice_row_num <= 4 else [4, choice_row_num-4]
    choice_row_num = 0
    rest = 0
    break_num = 0
    ap_n = 0  # choice_answer_pattern分页时的块数
    for key, value in group_choice_count.items():
        row_n, row_rest = divmod(value, 15)
        row_n5, row_rest5 = divmod(value, 5)
        for i in range(row_n):
            choice_line_number.append(5)
            choice_col_num.append(3)
        if row_rest:
            if row_rest > 10:
                choice_col_num.append(3)
            elif row_rest > 5:
                choice_line_number.append(5)
                choice_col_num.append(2)
            else:
                choice_line_number.append(row_rest)
                choice_col_num.append(1)
        for i in range(row_n5):
            choice_answer_pattern.append(5)
        if row_rest5:
            choice_answer_pattern.append(row_rest5)
        count_ed = row_n + 1 if row_rest else row_n
        choice_row_num += count_ed
        for i in range(count_ed):
            choice_answer_cnt.append(key)
        if choice_row_num == 4:
            rest = row_rest
        if choice_row_num > 4:
            if not break_num:
                if not rest:
                    break_num = group_choice_num[key][value - row_rest + 1]
                else:
                    break_num = group_choice_num[key][0]
    choice_row_num = [choice_row_num] if choice_row_num <= 4 else [4, choice_row_num - 4]
    choice_answer_cnt = [choice_answer_cnt] if not break_num else [choice_answer_cnt[:4], choice_answer_cnt[4:]]
    choice_col_num = [choice_col_num] if not break_num else [choice_col_num[:4], choice_col_num[4:]]
    choice_line_number = [choice_line_number] if not break_num else [choice_line_number[:4], choice_line_number[4:]]
    logger.info("group_choice_count========================={}".format(group_choice_count))
    logger.info("choice_answer_cnt========================={}".format(choice_answer_cnt))
    logger.info("choice_line_number========================={}".format(choice_line_number))

    group_choice_num_ = sorted(group_choice_num.items(), key=lambda d: d[0])
    group_choice_opt_ = sorted(group_choice_opt.items(), key=lambda d: d[0])
    group_choice_count_ = sorted(group_choice_count.items(), key=lambda d: d[0])
    # logger.info("group_choice_num_========================={}".format(group_choice_num_))
    choice_num = []
    for key, value in group_choice_num_:
        if isinstance(value, list):
            for i in value:
                choice_num.append(i)
        else:
            choice_num.append(value)
    choice_opt = []
    for key, value in group_choice_opt_:
        if isinstance(value, list):
            for i in value:
                choice_opt.append(i)
        else:
            choice_opt.append(value)

    row0 = choice_row_num[0]
    if break_num:
        for nn in choice_col_num[0:4]:
            ap_n += nn
        choice_answer_pattern = [choice_answer_pattern[0:ap_n], choice_answer_pattern[ap_n:]]
        row1 = choice_row_num[1]
        index = choice_ques_no.index(break_num)
        write_option(ques_no, choice_num[0:index], choice_opt[0:index], row0, document)
        document.add_page_break()
        write_option(ques_no, choice_num[index:], choice_opt[index:], row1, document)
    else:
        choice_answer_pattern = [choice_answer_pattern]
        write_option(ques_no, choice_num, choice_opt, row0, document)
    # write_option(ques_no, choice_len_sum, option_len, choice_ques_no, document)
    writeline(document, black_font)
    num, paper_num = choice0_60(fillin_len, fs_num, document, paper_num)
    paper_num += num

    if ss_num:
        sub_subjective(ss_num, paper_num, document, pic_in_sub)

    block_num = 0
    start = 0
    step = 3
    if choice_len_sum:
        for n in range(len(choice_row_num)):
            block_num += 1
            sub_mark_choice.append(start + step * (block_num - 1) - (block_num - 1))
            sub_mark_choice.append(start + step * block_num - (block_num - 1))
    if fillin_len:
        block_num += 1
        sub_mark_fillin.append(start + step * (block_num - 1) - (block_num - 1))
        sub_mark_fillin.append(start + step * block_num - (block_num - 1))
        if not subjec_len:
            index = sub_mark_fillin[-2]
            sub_mark_fillin.pop(-1)
            sub_mark_fillin.append(index + 1)
        else:
            for n in range(subjec_len):
                block_num += 1
                sub_mark_eassy.append(start + step * (block_num - 1) - (block_num - 1))
                sub_mark_eassy.append(start + step * block_num - (block_num - 1))
            if subjec_len == 1:
                index = sub_mark_eassy[-2]
                sub_mark_eassy.pop(-1)
                sub_mark_eassy.append(index + 1)
    else:
        if subjec_len:
            for n in range(subjec_len):
                block_num += 1
                sub_mark_eassy.append(start + step * (block_num - 1) - (block_num - 1))
                sub_mark_eassy.append(start + step * block_num - (block_num - 1))
            if subjec_len == 1:
                index = sub_mark_eassy[-2]
                sub_mark_eassy.pop(-1)
                sub_mark_eassy.append(index + 1)
        else:
            if sub_mark_choice:
                index = sub_mark_choice[-2]
                sub_mark_choice.pop(-1)
                sub_mark_choice.append(index + 1)

    #  return sub_mark_num
    for i in range(len(choice_row_num)):
        sub_mark_num += 2
    if fillin_len:
        sub_mark_num += 2
    for num in range(subjec_len):
        sub_mark_num += 2

    t1 = int(round(time.time() * 1000))
    doc_file_name = "{}-答题卡".format(sheet_name or str(t1))
    return_path = '/tmp/paper/%s.%s' % (doc_file_name, 'docx')
    document.save(return_path)
    sub_json_data = {}
    tmp_sub_mark_choice = []
    for i in range(len(choice_row_num)):
        tmp_sub_mark_choice.append([sub_mark_choice[i * 2], sub_mark_choice[i * 2 + 1]])
    logger.info("sub_mark_choice========================={}".format(sub_mark_choice))
    sub_json_data['sub_mark_choice'] = tmp_sub_mark_choice
    sub_json_data['sub_mark_fillin'] = sub_mark_fillin
    sub_json_data['sub_mark_eassy'] = sub_mark_eassy
    sub_json_data['sub_mark_num'] = sub_mark_num
    # sub_json_data['sub_mark_num_vec'] = sub_mark_num_vec
    sub_json_data['page_num'] = paper_num
    # sub_json_data['paper_num_count_flag'] = paper_num_count_flag
    sub_json_data['choice_answer_cnt'] = choice_answer_cnt
    sub_json_data['choice_line_number'] = choice_line_number
    sub_json_data['choice_answer_pattern'] = choice_answer_pattern
    sub_json_data['choice_col_num'] = choice_col_num
    sub_json_data['choice_row_num'] = choice_row_num
    sub_json_data['qus_range_choice'] = qus_range_choice
    sub_json_data['qus_range_fillin'] = qus_range_fillin
    sub_json_data['qus_range_eassy'] = qus_range_eassy
    logger.info("sub_json_data========================={}".format(sub_json_data))
    json_json = get_json(sub_json_data)
    logger.info("json_json========================={}".format(json_json))
    return return_path, sub_json_data


def get4answer_sheet_data(paper_data):
    answer_sheet = {}
    ques_no = 0
    for g_idx, grp in enumerate(paper_data.get("question_group", [])):
        for eq in grp.get("exam_questions", []):
            question = eq.get("question", {})
            q_type_name = question.get("q_type", {}).get("name")
            pic_in_card = question.get("pic_in_card", "")
            options = question.get("options") or []
            subquestions = question.get("subs") or []
            if q_type_name in CHOICE_QUESTIONS:
                key = "choice_question"
            elif subquestions:
                key = "combine_question"
            elif q_type_name in FILL_IN_QUESTIONS:
                key = "fillin_question"
            else:
                key = "subjective_question"
            if not answer_sheet.get(key):
                answer_sheet[key] = []
            answer_sheet[key].append({
                "q_no": eq.get("q_no"),
                "pic_in_card": pic_in_card,
                "options": [v["key"] for v in options if v.get("key")],
                "subques_len": len(subquestions) if subquestions else 0
            })
            if q_type_name in CHOICE_QUESTIONS or subquestions:
                key = "choice_question_all"
                if not answer_sheet.get(key):
                    answer_sheet[key] = []
                answer_sheet[key].append({
                    "q_no": eq.get("q_no"),
                    "pic_in_card": pic_in_card,
                    "options": [v["key"] for v in options if v.get("key")],
                    "subques_len": len(subquestions) if subquestions else 0
                })
            if key == "combine_question":
                ques_no += len(subquestions)
            else:
                ques_no += 1
    answer_sheet_data = {
        "name": paper_data["name"],
        "main_title": "{}-答题卡".format(paper_data["main_title"]),
        "testee": paper_data["testee"],
        "sheet": answer_sheet
    }
    answer_path, sub_json_data = write4answersheet(answer_sheet_data)
    return answer_path, answer_sheet_data, sub_json_data


def create_paper_file(paper, data, fmt, size, usage):
    try:
        # create paper file
        filepath = get_paper(data, fmt, size, usage)
        with open(filepath, "rb") as outfile:
            r = polo.upload_file(outfile)
        if not r:
            return APIResponse(code=STORE_FILE_FAILED, message=_("存储文件失败"))
        hash_code = r.hash
        models.ExamRelatedPaper.save_hash(paper=paper, type=size, hash_code=hash_code, usage=usage, fmt=fmt)
    except Exception as e:
        logger.error("upload paper file to polo error, detail: {}".format(e))


def create_answersheet(paper, data, size=None):
    try:
        size = size or models.AnswerSheetSizeSet.A4
        # create answer sheet
        paper_data = QuestionNumberBuilder(paper_data=data).attach_ques_no
        answer_path, answer_sheet_data, sub_json_data = get4answer_sheet_data(paper_data)
        with open(answer_path, "rb") as outfile:
            r = polo.upload_file(outfile)
        if not r:
            return APIResponse(code=STORE_FILE_FAILED, message=_("存储文件失败"))
        answer_hash_code = r.hash
        # save json_data and hash
        json_data = get_json(sub_json_data)
        models.ExamRelatedAnswersheet.save_answer(paper, size, answer_hash_code, json_data)
    except Exception as e:
        traceback.print_exc()
        logger.error("upload answer sheet file to polo error, detail: {}".format(e))
